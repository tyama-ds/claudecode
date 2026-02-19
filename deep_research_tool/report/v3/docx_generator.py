"""
DOCX-native Report Generator V3.

Builds Word documents directly via python-docx official API.
No Markdown intermediate representation — LLM-generated chapter content
is parsed and written straight into Document objects.

Inherits V2's consistency features (glossary, context, two-phase generation)
for the content generation phase.
"""

import io
import logging
import re
import time
from datetime import datetime
from pathlib import Path
from typing import List, Dict, Optional, Any, Callable, Tuple

from ..v2.generator import (
    ReportGeneratorV2,
    GenerationResult,
    ChapterContent,
    ReportFormatError,
    _sanitize_for_xml,
    _section_sort_key,
)
from ..v2.context import ReportContext, WritingStyle, TargetAudience
from ..figure_table_generator import FigureTableCollection, Figure, TableData
from ...utils.helpers import ResearchWarnings

logger = logging.getLogger(__name__)


class DocxReportGeneratorV3(ReportGeneratorV2):
    """
    DOCX-native report generator (V3).

    Extends ReportGeneratorV2 for content generation (glossary, context,
    consistency checking) but replaces the save pipeline with direct
    python-docx API calls.

    Key differences from V2:
    - No Markdown intermediate step for DOCX output
    - All document elements built via python-docx official API
    - Figures/tables/charts inserted with doc.add_picture() / doc.add_table()
    - Markdown output still available as fallback
    """

    def __init__(
        self,
        llm_client,
        language: str = "ja",
        writing_style: WritingStyle = WritingStyle.BUSINESS,
        target_audience: TargetAudience = TargetAudience.BUSINESS,
        technical_level: int = 3,
        enable_consistency_check: bool = True,
        enable_two_phase: bool = True,
        progress_callback: Callable[[str, int, int], None] = None,
        target_pages: Optional[int] = None,
        target_characters: Optional[int] = None,
    ):
        super().__init__(
            llm_client=llm_client,
            language=language,
            writing_style=writing_style,
            target_audience=target_audience,
            technical_level=technical_level,
            enable_consistency_check=enable_consistency_check,
            enable_two_phase=enable_two_phase,
            progress_callback=progress_callback,
            target_pages=target_pages,
            target_characters=target_characters,
        )

    # ------------------------------------------------------------------
    # Public API
    # ------------------------------------------------------------------

    def generate_and_save(
        self,
        result: GenerationResult,
        output_dir: Path,
        filename: str,
        evidence_locker=None,
        figure_collection: Optional[FigureTableCollection] = None,
        include_glossary: bool = True,
        fermi_markdown: str = "",
        warnings_text: str = "",
    ) -> Path:
        """
        Build a DOCX file directly from GenerationResult using python-docx.

        Args:
            result: GenerationResult from generate_report()
            output_dir: Directory to save the report
            filename: Base filename (without extension)
            evidence_locker: EvidenceLocker for references section
            figure_collection: Optional FigureTableCollection for figures/tables/charts
            include_glossary: Include glossary section
            fermi_markdown: Fermi estimation text to append
            warnings_text: Warnings text to append

        Returns:
            Path to the saved DOCX file.

        Raises:
            ReportFormatError: When DOCX generation fails.
        """
        try:
            from docx import Document
        except ImportError:
            raise ReportFormatError(
                "python-docx is not installed. Cannot generate DOCX. "
                "Install with: pip install python-docx"
            )

        output_dir = Path(output_dir)
        output_dir.mkdir(parents=True, exist_ok=True)

        doc = Document()
        self._setup_document_styles(doc)

        # Build evidence mapping for citation renumbering
        evidence_list = []
        url_to_ref = {}
        if evidence_locker is not None:
            evidence_list = evidence_locker.get_all_evidence()
            for i, evidence in enumerate(evidence_list, 1):
                url_to_ref[evidence.url] = i

        # === Title ===
        self._add_title_page(doc, result)

        # === Table of Contents ===
        self._add_toc(doc, result)

        # === Chapters ===
        sorted_chapters = sorted(result.chapters.items(), key=_section_sort_key)
        for section_num, chapter in sorted_chapters:
            section_figures = None
            if figure_collection:
                section_figures = self._collect_section_figures(
                    figure_collection, section_num
                )
            self._add_chapter(
                doc, section_num, chapter, url_to_ref, section_figures
            )

        # === Glossary ===
        if include_glossary and result.context.glossary:
            self._add_glossary(doc, result.context)

        # === References ===
        if evidence_list:
            self._add_references(doc, evidence_list)

        # === Fermi estimation (plain text appendix) ===
        if fermi_markdown:
            self._add_text_section(doc, fermi_markdown)

        # === Warnings (plain text appendix) ===
        if warnings_text:
            self._add_text_section(doc, warnings_text)

        # === Orphaned figures (not matched to any section) ===
        if figure_collection:
            matched_sections = {ch[0] for ch in sorted_chapters}
            self._add_orphaned_figures(doc, figure_collection, matched_sections)

        # === Save ===
        filepath = output_dir / f"{filename}.docx"
        self._safe_save(doc, filepath)
        return filepath

    # ------------------------------------------------------------------
    # Document Structure Builders
    # ------------------------------------------------------------------

    def _setup_document_styles(self, doc) -> None:
        """Configure base document styles."""
        from docx.shared import Pt, Cm

        style = doc.styles["Normal"]
        font = style.font
        font.size = Pt(10.5)
        # Use a widely available font; CJK fallback is handled by Word itself
        font.name = "Arial"

        # Paragraph spacing
        paragraph_format = style.paragraph_format
        paragraph_format.space_after = Pt(6)
        paragraph_format.line_spacing = 1.5

        # Page margins
        for section in doc.sections:
            section.top_margin = Cm(2.5)
            section.bottom_margin = Cm(2.5)
            section.left_margin = Cm(2.5)
            section.right_margin = Cm(2.5)

    def _add_title_page(self, doc, result: GenerationResult) -> None:
        """Add report title and generation date."""
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        title = _sanitize_for_xml(result.context.report_title)
        heading = doc.add_heading(title, level=0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Generation date
        date_para = doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = date_para.add_run(
            _sanitize_for_xml(f"Generated: {datetime.now().strftime('%Y-%m-%d')}")
        )
        run.italic = True
        run.font.size = Pt(9)

        doc.add_paragraph()  # spacer

    def _add_toc(self, doc, result: GenerationResult) -> None:
        """Add table of contents."""
        from docx.shared import Pt

        toc_title = "目次" if self.language == "ja" else "Table of Contents"
        doc.add_heading(_sanitize_for_xml(toc_title), level=1)

        for section_num, chapter in sorted(result.chapters.items(), key=_section_sort_key):
            text = _sanitize_for_xml(f"{section_num}. {chapter.section_title}")
            para = doc.add_paragraph()
            run = para.add_run(text)
            run.font.size = Pt(10)

        doc.add_paragraph()  # spacer

    def _add_chapter(
        self,
        doc,
        section_num: str,
        chapter: ChapterContent,
        url_to_ref: Dict[str, int],
        section_figures: Optional[Dict[str, list]] = None,
    ) -> None:
        """
        Add a single chapter to the document.

        Parses the chapter content (which may contain lightweight markdown
        from LLM output) and writes it via python-docx API.
        """
        # Chapter heading
        heading_text = _sanitize_for_xml(
            f"{section_num}. {chapter.section_title}"
        )
        # Determine heading level: "1" -> level 1, "1.1" -> level 2
        heading_level = min(len(section_num.split(".")), 3)
        doc.add_heading(heading_text, level=heading_level)

        # Renumber citations in content
        content = chapter.content
        if url_to_ref and chapter.citations:
            content = self._renumber_citations(content, chapter.citations, url_to_ref)

        # Parse content and add to doc
        self._add_content_to_doc(doc, content)

        # Insert figures/tables/charts for this section
        if section_figures:
            self._insert_section_figures(doc, section_figures)

    def _add_content_to_doc(self, doc, content: str) -> None:
        """
        Parse LLM-generated content and write to doc via python-docx API.

        Handles: headings, bullet lists, numbered lists, markdown tables,
        inline bold/italic, and regular paragraphs.
        """
        if not content:
            return

        lines = content.split("\n")
        i = 0

        while i < len(lines):
            line = lines[i]
            stripped = line.strip()

            # Empty line — skip
            if not stripped:
                i += 1
                continue

            stripped = _sanitize_for_xml(stripped)

            # Sub-heading within chapter (### or deeper)
            if stripped.startswith("#"):
                level, text = self._parse_heading(stripped)
                text = _sanitize_for_xml(text)
                # Clamp to valid range (heading within chapter: level 2-4)
                level = max(2, min(level + 1, 4))
                try:
                    doc.add_heading(text, level=level)
                except Exception:
                    doc.add_paragraph(text)
                i += 1
                continue

            # Markdown table
            if stripped.startswith("|") and "|" in stripped[1:]:
                table_lines = []
                while i < len(lines) and lines[i].strip().startswith("|"):
                    table_lines.append(_sanitize_for_xml(lines[i].strip()))
                    i += 1
                if table_lines:
                    self._add_markdown_table(doc, table_lines)
                continue

            # Unordered list
            if re.match(r"^[-*+]\s", stripped):
                text = _sanitize_for_xml(re.sub(r"^[-*+]\s+", "", stripped))
                try:
                    para = doc.add_paragraph(style="List Bullet")
                except KeyError:
                    para = doc.add_paragraph()
                    text = "- " + text
                self._add_formatted_runs(para, text)
                i += 1
                continue

            # Ordered list
            if re.match(r"^\d+\.\s", stripped):
                text = _sanitize_for_xml(re.sub(r"^\d+\.\s+", "", stripped))
                try:
                    para = doc.add_paragraph(style="List Number")
                except KeyError:
                    para = doc.add_paragraph()
                    text = stripped
                self._add_formatted_runs(para, text)
                i += 1
                continue

            # Horizontal rule
            if re.match(r"^-{3,}$|^\*{3,}$|^_{3,}$", stripped):
                # Add a thin horizontal line via a paragraph border
                doc.add_paragraph("─" * 60)
                i += 1
                continue

            # Regular paragraph — collect consecutive lines
            para_lines = []
            start_i = i
            while i < len(lines):
                l = lines[i].strip()
                if not l:
                    break
                if l.startswith("#") or re.match(r"^[-*+]\s", l) or re.match(r"^\d+\.\s", l):
                    break
                if l.startswith("|") and "|" in l[1:]:
                    break
                if re.match(r"^-{3,}$|^\*{3,}$|^_{3,}$", l):
                    break
                para_lines.append(_sanitize_for_xml(l))
                i += 1

            if i == start_i:
                i += 1

            if para_lines:
                para = doc.add_paragraph()
                self._add_formatted_runs(para, " ".join(para_lines))

    def _add_markdown_table(self, doc, table_lines: List[str]) -> None:
        """Convert markdown table lines to a python-docx Table."""
        from docx.shared import Pt

        rows_data = []
        for line in table_lines:
            cells = [c.strip() for c in line.strip("|").split("|")]
            # Skip separator row (e.g., |---|---|)
            if all(re.match(r"^[-:]+$", c.strip()) for c in cells if c.strip()):
                continue
            rows_data.append(cells)

        if not rows_data:
            return

        num_cols = max(len(row) for row in rows_data)
        table = doc.add_table(rows=len(rows_data), cols=num_cols)
        try:
            table.style = "Table Grid"
        except KeyError:
            pass  # style not available — use default

        for row_idx, row_data in enumerate(rows_data):
            for col_idx, cell_text in enumerate(row_data):
                if col_idx < num_cols:
                    cell = table.rows[row_idx].cells[col_idx]
                    cell.text = _sanitize_for_xml(cell_text.strip())
                    # Bold header row
                    if row_idx == 0:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.bold = True
                                run.font.size = Pt(9)

        doc.add_paragraph()  # spacing after table

    # ------------------------------------------------------------------
    # Figure / Table / Chart Insertion (python-docx official API)
    # ------------------------------------------------------------------

    def _collect_section_figures(
        self, collection: FigureTableCollection, section_num: str
    ) -> Dict[str, list]:
        """Collect figures, tables, and charts for a given section."""
        return {
            "figures": collection.get_figures_for_section(section_num),
            "tables": collection.get_tables_for_section(section_num),
            "charts": collection.get_charts_for_section(section_num),
        }

    def _insert_section_figures(self, doc, section_figures: Dict[str, list]) -> None:
        """Insert all figures/tables/charts for a section at the current position."""
        # Charts first (generated visualizations)
        for chart in section_figures.get("charts", []):
            self._add_figure_to_doc(doc, chart)

        # Tables
        for table_data in section_figures.get("tables", []):
            self._add_table_data_to_doc(doc, table_data)

        # Figures (images from web sources)
        for figure in section_figures.get("figures", []):
            self._add_figure_to_doc(doc, figure)

    def _add_figure_to_doc(self, doc, figure: Figure) -> None:
        """
        Add a Figure (image/chart) to the document using doc.add_picture().

        Falls back to a text placeholder if the image file is missing.
        """
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        img_path = figure.image_path
        image_added = False

        if img_path:
            path = Path(img_path) if not isinstance(img_path, Path) else img_path
            if path.exists():
                try:
                    para = doc.add_paragraph()
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = para.add_run()
                    run.add_picture(str(path), width=Inches(5.0))
                    image_added = True
                except Exception as e:
                    logger.warning(f"[V3] Failed to add image {path}: {e}")

        if not image_added and figure.image_data:
            # Image data in memory (bytes)
            try:
                para = doc.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = para.add_run()
                run.add_picture(io.BytesIO(figure.image_data), width=Inches(5.0))
                image_added = True
            except Exception as e:
                logger.warning(f"[V3] Failed to add image from bytes: {e}")

        if not image_added:
            # Placeholder text
            alt = figure.alt_text or figure.title or "Image"
            doc.add_paragraph(
                _sanitize_for_xml(f"[Image: {alt}]")
            )

        # Caption
        if figure.caption:
            caption_para = doc.add_paragraph()
            caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            caption_run = caption_para.add_run(
                _sanitize_for_xml(figure.caption)
            )
            caption_run.italic = True
            caption_run.font.size = Pt(9)

    def _add_table_data_to_doc(self, doc, table_data: TableData) -> None:
        """
        Add a TableData object to the document using doc.add_table().
        """
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        if not table_data.headers:
            return

        # Table title
        if table_data.title:
            title_para = doc.add_paragraph()
            title_run = title_para.add_run(
                _sanitize_for_xml(table_data.title)
            )
            title_run.bold = True
            title_run.font.size = Pt(10)

        num_cols = len(table_data.headers)
        num_rows = len(table_data.rows) + 1  # +1 for header row
        table = doc.add_table(rows=num_rows, cols=num_cols)

        try:
            table.style = "Light Grid"
        except KeyError:
            try:
                table.style = "Table Grid"
            except KeyError:
                pass

        # Header row
        for col_idx, header in enumerate(table_data.headers):
            cell = table.rows[0].cells[col_idx]
            cell.text = _sanitize_for_xml(str(header))
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
                    run.font.size = Pt(9)

        # Data rows
        for row_idx, row in enumerate(table_data.rows):
            for col_idx, cell_value in enumerate(row):
                if col_idx < num_cols:
                    cell = table.rows[row_idx + 1].cells[col_idx]
                    cell.text = _sanitize_for_xml(str(cell_value))
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.size = Pt(9)

        # Caption
        if table_data.caption and table_data.caption != table_data.title:
            caption_para = doc.add_paragraph()
            caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            caption_run = caption_para.add_run(
                _sanitize_for_xml(table_data.caption)
            )
            caption_run.italic = True
            caption_run.font.size = Pt(9)

        doc.add_paragraph()  # spacing

    # ------------------------------------------------------------------
    # Glossary / References / Appendix
    # ------------------------------------------------------------------

    def _add_glossary(self, doc, context: ReportContext) -> None:
        """Add glossary section."""
        from docx.shared import Pt

        title = "用語集" if self.language == "ja" else "Glossary"
        doc.add_heading(_sanitize_for_xml(title), level=1)

        for key in sorted(context.glossary.keys()):
            entry = context.glossary[key]
            term_text = _sanitize_for_xml(entry.term)
            definition = _sanitize_for_xml(entry.definition) if entry.definition else ""

            para = doc.add_paragraph()
            # Term in bold
            term_run = para.add_run(f"{term_text}")
            term_run.bold = True
            term_run.font.size = Pt(10)

            if definition:
                para.add_run(f" — {definition}").font.size = Pt(10)

            # Aliases
            if entry.aliases:
                aliases_text = _sanitize_for_xml(", ".join(entry.aliases))
                alias_para = doc.add_paragraph()
                alias_run = alias_para.add_run(f"  ({aliases_text})")
                alias_run.italic = True
                alias_run.font.size = Pt(9)

    def _add_references(self, doc, evidence_list: list) -> None:
        """Add references section."""
        from docx.shared import Pt

        title = "参考文献" if self.language == "ja" else "References"
        doc.add_heading(_sanitize_for_xml(title), level=1)

        for i, evidence in enumerate(evidence_list, 1):
            badge = self._get_quality_badge(evidence.quality_category)
            citation = _sanitize_for_xml(evidence.citation_text)
            text = f"{i}. {badge} {citation}" if badge else f"{i}. {citation}"

            para = doc.add_paragraph()
            run = para.add_run(_sanitize_for_xml(text))
            run.font.size = Pt(9)

    def _add_text_section(self, doc, text: str) -> None:
        """
        Add a plain text section (for Fermi estimation / warnings appendix).
        Parses lightweight markdown (headings, lists, paragraphs).
        """
        if not text:
            return
        self._add_content_to_doc(doc, text)

    def _add_orphaned_figures(
        self,
        doc,
        collection: FigureTableCollection,
        matched_sections: set,
    ) -> None:
        """Add figures/charts not matched to any section at the end of the document."""
        orphan_figures = [
            f for f in collection.figures if f.section_id not in matched_sections
        ]
        orphan_charts = [
            c for c in collection.charts if c.section_id not in matched_sections
        ]

        has_orphans = any(
            self._figure_has_image(item)
            for item in orphan_figures + orphan_charts
        )

        if not has_orphans:
            return

        title = "その他の図表" if self.language == "ja" else "Additional Figures"
        doc.add_heading(_sanitize_for_xml(title), level=1)

        for fig in orphan_figures:
            if self._figure_has_image(fig):
                self._add_figure_to_doc(doc, fig)
        for chart in orphan_charts:
            if self._figure_has_image(chart):
                self._add_figure_to_doc(doc, chart)

    @staticmethod
    def _figure_has_image(figure: Figure) -> bool:
        """Check if a figure has a valid image file or data."""
        if figure.image_data:
            return True
        if figure.image_path:
            path = Path(figure.image_path) if not isinstance(figure.image_path, Path) else figure.image_path
            return path.exists()
        return False

    # ------------------------------------------------------------------
    # Save with validation
    # ------------------------------------------------------------------

    def _safe_save(self, doc, filepath: Path) -> None:
        """
        Save document with BytesIO pre-validation.

        Sanitizes all text elements on failure and retries once.
        """
        try:
            buf = io.BytesIO()
            doc.save(buf)
            buf.seek(0)
            with open(filepath, "wb") as f:
                f.write(buf.read())
        except Exception as first_error:
            logger.warning(f"[V3] First save attempt failed: {first_error}. Retrying with sanitization.")
            try:
                self._sanitize_all_text(doc)
                buf2 = io.BytesIO()
                doc.save(buf2)
                buf2.seek(0)
                with open(filepath, "wb") as f:
                    f.write(buf2.read())
                logger.info("[V3] Retry save succeeded after sanitization.")
            except Exception as retry_error:
                raise ReportFormatError(
                    f"DOCX save failed even after sanitization. "
                    f"Original: {first_error} / Retry: {retry_error}",
                    original_error=retry_error,
                )

    @staticmethod
    def _sanitize_all_text(doc) -> None:
        """Sanitize all text in document paragraphs and tables."""
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                if run.text:
                    run.text = _sanitize_for_xml(run.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            if run.text:
                                run.text = _sanitize_for_xml(run.text)

    # ------------------------------------------------------------------
    # Overrides: disable V2's markdown-based save methods
    # ------------------------------------------------------------------

    def save_report(
        self,
        markdown_content: str,
        output_dir: Path,
        filename: str,
        format: str = "docx",
        strict_format: bool = False,
    ) -> Path:
        """
        Override V2's save_report for backward compatibility.

        For V3 the recommended path is generate_and_save(), but if called
        via the V2 interface (e.g. markdown_content provided), fall back
        to V2 behavior for non-DOCX formats and raise an error for DOCX
        suggesting generate_and_save() instead.
        """
        format_lower = format.lower() if isinstance(format, str) else format.value.lower()

        if format_lower != "docx":
            # For non-DOCX, delegate to V2's implementation
            return super().save_report(
                markdown_content, output_dir, filename, format, strict_format
            )

        # For DOCX via old interface, use V2's markdown-based conversion
        # (generate_and_save is the preferred V3 path)
        logger.info(
            "[V3] save_report() called with DOCX format. "
            "For best results, use generate_and_save() instead."
        )
        return super().save_report(
            markdown_content, output_dir, filename, format, strict_format
        )
