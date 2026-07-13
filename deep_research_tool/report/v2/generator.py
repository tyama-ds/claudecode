"""
Report Generator V2 - Enhanced report generation with consistency features.

Version 2.0 adds:
- Global context maintenance across chapters
- Terminology consistency via glossary
- Previous chapter summaries for continuity
- Fact tracking to prevent contradictions
- Post-generation consistency checking
- Two-phase generation (draft + refinement)
"""

import json
import re
import logging
import time
from ...utils.helpers import ResearchWarnings
from dataclasses import dataclass, field
from pathlib import Path
from typing import List, Dict, Optional, Any, Callable, Tuple
from datetime import datetime

from .context import ReportContext, WritingStyle, TargetAudience, ChapterSummary
from .consistency import ConsistencyChecker, ConsistencyReport
from .glossary import GlossaryManager
from ...utils.helpers import split_prose_and_meta

# Delimiter between chapter prose and metadata JSON in generation responses
CHAPTER_META_DELIMITER = "===CHAPTER_META==="

# Paragraph-leading label prefixes that make prose read like bullet notes
# ("結論：〜", "要点：〜"). Some models emit these compulsively (especially
# when a style instruction mentions conclusions-first); strip the label and
# keep the sentence. Headings / tables / bullets / quotes are left untouched.
_LABEL_PREFIX_RE = re.compile(
    r'^(?:\*\*)?(?:結論|要点|ポイント|概要|まとめ|示唆|考察|分析)(?:\*\*)?[：:]\s*'
)


def strip_label_prefixes(text: str) -> str:
    """Remove note-style label prefixes (結論： etc.) from paragraph starts."""
    out = []
    for line in text.split("\n"):
        stripped = line.lstrip()
        is_block = stripped.startswith(("#", "|", ">", "- ", "* ")) or \
            re.match(r"^\d+[.)]\s", stripped)
        if not is_block:
            indent = line[:len(line) - len(stripped)]
            new = _LABEL_PREFIX_RE.sub("", stripped)
            if new != stripped and new:
                line = indent + new
        out.append(line)
    return "\n".join(out)

logger = logging.getLogger(__name__)


def _section_sort_key(item) -> Tuple:
    """
    Natural sort key for section numbers like "1", "1.1", "2", "10", "10.1".

    Converts section number strings into tuples of integers for proper
    numerical ordering:
        "1"    -> (1,)
        "1.1"  -> (1, 1)
        "2"    -> (2,)
        "10"   -> (10,)
        "10.1" -> (10, 1)

    This ensures "2" comes before "10" (numerical order), not after "1.9"
    and before "10" (lexicographic order).
    """
    key = item[0] if isinstance(item, tuple) else item
    parts = []
    for part in str(key).split('.'):
        try:
            parts.append(int(part))
        except ValueError:
            # Non-numeric parts (e.g., "A", "_executive") sort after numbers
            parts.append(float('inf'))
    return tuple(parts)


class ReportFormatError(Exception):
    """Raised when report cannot be saved in the requested format (strict_format mode)."""

    def __init__(self, message: str, debug_md_path: str = None, original_error: Exception = None):
        super().__init__(message)
        self.debug_md_path = debug_md_path
        self.original_error = original_error


# Pre-compiled regex for XML sanitization (covers ALL illegal XML 1.0 characters)
# XML 1.0 legal chars: #x9 | #xA | #xD | [#x20-#xD7FF] | [#xE000-#xFFFD] | [#x10000-#x10FFFF]
# Everything else is illegal and must be removed.
_XML_ILLEGAL_CHARS_RE = re.compile(
    '['
    '\x00-\x08'          # C0 control chars (except TAB \x09)
    '\x0b\x0c'           # VT, FF (except LF \x0a, CR \x0d)
    '\x0e-\x1f'          # C0 control chars (cont.)
    '\x7f'               # DEL
    '\x80-\x84'          # C1 control chars
    '\x86-\x9f'          # C1 control chars (except NEL \x85)
    '\ud800-\udfff'      # Surrogate pairs (isolated)
    '\ufdd0-\ufdef'      # Non-characters
    '\ufffe\uffff'        # Non-characters (BOM reverse, etc.)
    ']',
)


def _sanitize_for_xml(text: str) -> str:
    """
    Remove characters illegal in XML 1.0 from text.

    DOCX files are XML internally. XML 1.0 spec defines legal chars as:
        #x9 | #xA | #xD | [#x20-#xD7FF] | [#xE000-#xFFFD] | [#x10000-#x10FFFF]

    All other characters (control chars, surrogates, non-characters, DEL, C1 range)
    are stripped to prevent doc.save() XML serialization errors.
    """
    if not text:
        return text
    return _XML_ILLEGAL_CHARS_RE.sub('', text)


@dataclass
class ChapterContent:
    """Generated chapter content with metadata."""
    section_number: str
    section_title: str
    content: str
    word_count: int = 0
    key_points: List[str] = field(default_factory=list)
    terms_used: List[str] = field(default_factory=list)
    facts_stated: List[str] = field(default_factory=list)
    citations: List[str] = field(default_factory=list)
    is_draft: bool = True
    generated_at: str = field(default_factory=lambda: datetime.now().isoformat())


@dataclass
class GenerationResult:
    """Result of report generation."""
    chapters: Dict[str, ChapterContent]
    context: ReportContext
    consistency_report: Optional[ConsistencyReport] = None
    total_word_count: int = 0
    generation_time_seconds: float = 0.0


class ReportGeneratorV2:
    """
    Enhanced report generator with consistency features.

    This is the V2 implementation that ensures consistency across all chapters
    by maintaining global context, terminology, and fact tracking.

    Usage:
        generator = ReportGeneratorV2(
            llm_client=llm,
            writing_style=WritingStyle.BUSINESS,
            target_audience=TargetAudience.BUSINESS,
        )

        result = generator.generate_report(
            research_topic="炭素繊維の市場調査",
            research_plan=plan,
            section_contents=section_contents,
        )
    """

    # Default per-chapter length target (chars) when neither target_pages
    # nor target_characters is specified
    DEFAULT_CHAPTER_CHARS_JA = 1500
    DEFAULT_CHAPTER_CHARS_EN = 2500

    # Short-chapter guard: a chapter below target * ratio triggers one
    # expansion retry; if it is still below, a ResearchWarning is recorded
    SHORT_CHAPTER_RATIO = 0.4

    def __init__(
        self,
        llm_client,
        language: str = "ja",
        writing_style: WritingStyle = WritingStyle.BUSINESS,
        target_audience: TargetAudience = TargetAudience.BUSINESS,
        technical_level: int = 3,
        enable_consistency_check: bool = True,
        enable_two_phase: bool = True,
        enable_polish: bool = True,
        progress_callback: Callable[[str, int, int], None] = None,
        target_pages: Optional[int] = None,
        target_characters: Optional[int] = None,
    ):
        """
        Initialize ReportGeneratorV2.

        Args:
            llm_client: LLM API client
            language: Target language
            writing_style: Writing style for the report
            target_audience: Target audience
            technical_level: Technical level (1-5)
            enable_consistency_check: Run consistency check after generation
            enable_two_phase: Enable two-phase generation (draft + refinement)
            enable_polish: Run a final naturalness polish pass over all chapters
            progress_callback: Progress callback function
            target_pages: Target page count (approximate, 1 page ≈ 1500 chars ja / 500 words en)
            target_characters: Target character count (overrides target_pages if set)
        """
        self.llm = llm_client
        self.language = language
        self._chapter_target_chars = None
        self.writing_style = writing_style
        self.target_audience = target_audience
        self.technical_level = technical_level
        self.enable_consistency_check = enable_consistency_check
        self.enable_two_phase = enable_two_phase
        self.enable_polish = enable_polish
        self.progress_callback = progress_callback
        self.target_pages = target_pages
        self.target_characters = target_characters

        # Initialize components
        self.glossary_manager = GlossaryManager(llm_client, language)
        self.consistency_checker = ConsistencyChecker(llm_client, language)

    def generate_report(
        self,
        research_topic: str,
        research_plan: Any,
        section_contents: Dict[str, Any],
        report_title: str = "",
    ) -> GenerationResult:
        """
        Generate a complete report with consistency features.

        Args:
            research_topic: Original research topic
            research_plan: ResearchPlan object
            section_contents: Dict of section_number -> extracted content
            report_title: Optional report title

        Returns:
            GenerationResult with all chapters and metadata
        """
        import time
        start_time = time.time()

        # Initialize context
        context = ReportContext(
            research_topic=research_topic,
            report_title=report_title or research_topic,
            language=self.language,
            writing_style=self.writing_style,
            target_audience=self.target_audience,
            technical_level=self.technical_level,
        )

        # Phase 0: Initialize glossary
        self._report_progress("Initializing glossary...", 0, 100)
        initial_glossary = self.glossary_manager.create_initial_glossary(
            research_topic=research_topic,
            research_plan=research_plan,
        )
        for key, entry in initial_glossary.items():
            context.add_glossary_term(
                term=entry.get("term", key),
                definition=entry.get("definition", ""),
                aliases=entry.get("aliases", []),
                preferred_form=entry.get("preferred_form", ""),
            )

        # Get sections from plan
        sections = self._get_sections_from_plan(research_plan)
        total_sections = len(sections)

        # Calculate per-chapter target length
        self._chapter_target_chars = None
        if total_sections > 0:
            if self.target_characters:
                self._chapter_target_chars = self.target_characters // total_sections
            elif self.target_pages:
                chars_per_page = 1500 if self.language == "ja" else 2500
                self._chapter_target_chars = (self.target_pages * chars_per_page) // total_sections

        # Default per-chapter target when no explicit target is given.
        # Without a length instruction the LLM tends to write very short
        # chapters (a few hundred characters), so always give one.
        if self._chapter_target_chars is None:
            self._chapter_target_chars = (
                self.DEFAULT_CHAPTER_CHARS_JA if self.language == "ja"
                else self.DEFAULT_CHAPTER_CHARS_EN
            )

        # Phase 1: Generate drafts
        chapters: Dict[str, ChapterContent] = {}

        for i, section in enumerate(sections):
            section_num = section.get("section", str(i + 1))
            section_title = section.get("title", f"Section {i + 1}")
            section_desc = section.get("description", "")

            self._report_progress(
                f"Generating chapter {section_num}: {section_title}",
                int(10 + (i / total_sections) * 60), 100
            )

            # Get content for this section
            content_data = section_contents.get(section_num, {})

            # Generate chapter with context
            chapter = self._generate_chapter(
                section_number=section_num,
                section_title=section_title,
                section_description=section_desc,
                content_data=content_data,
                context=context,
            )

            chapters[section_num] = chapter

            # Update context with chapter summary
            summary = self._extract_chapter_summary(chapter)
            context.add_chapter_summary(
                section_number=section_num,
                section_title=section_title,
                summary=summary,
                key_points=chapter.key_points,
                terms_introduced=chapter.terms_used,
                facts_established=chapter.facts_stated,
                word_count=chapter.word_count,
            )

            # Extract and add any new terms
            new_terms = self.glossary_manager.extract_terms_from_content(
                chapter.content, section_num
            )
            for term_candidate in new_terms:
                if term_candidate.term.lower() not in context.glossary:
                    context.add_glossary_term(
                        term=term_candidate.term,
                        definition="",
                        aliases=[term_candidate.expanded_form] if term_candidate.expanded_form else [],
                        section=section_num,
                    )

        # Phase 2: Consistency check and refinement
        consistency_report = None
        if self.enable_consistency_check:
            self._report_progress("Checking consistency...", 75, 100)
            chapter_texts = {k: v.content for k, v in chapters.items()}
            consistency_report = self.consistency_checker.check_all(chapter_texts, context)

            if self.enable_two_phase and not consistency_report.is_consistent:
                self._report_progress("Refining for consistency...", 85, 100)
                chapters = self._refine_chapters(chapters, consistency_report, context)

        # Phase 3: Naturalness polish over all chapters
        if self.enable_polish:
            self._report_progress("Polishing chapters for naturalness...", 90, 100)
            chapters = self._polish_chapters(chapters, context)

        # Calculate totals
        total_word_count = sum(c.word_count for c in chapters.values())
        context.total_word_count = total_word_count

        self._report_progress("Generation complete", 100, 100)

        return GenerationResult(
            chapters=chapters,
            context=context,
            consistency_report=consistency_report,
            total_word_count=total_word_count,
            generation_time_seconds=time.time() - start_time,
        )

    def _generate_chapter(
        self,
        section_number: str,
        section_title: str,
        section_description: str,
        content_data: Any,
        context: ReportContext,
    ) -> ChapterContent:
        """Generate a single chapter with context awareness."""

        # Build the prompt with full context
        context_prompt = context.get_full_context_prompt()

        # Format content data — include both summary and raw content per source.
        # The Researcher stores section_contents entries with keys
        # title/content/summary/sources; extracted_content is preferred when
        # a caller provides per-source material.
        no_info = "情報なし" if self.language == "ja" else "No information"
        if isinstance(content_data, dict):
            sources = content_data.get("sources", [])
            extracted = content_data.get("extracted_content", [])
            if extracted:
                source_blocks = []
                for i, e in enumerate(extracted[:10]):
                    title = e.get("title", "N/A")
                    summary = e.get("content", "")
                    raw = e.get("raw_content", "")
                    key_pts = e.get("key_points", [])
                    block = f"[SOURCE {i+1}] {title}\n"
                    if summary:
                        block += f"【要約】\n{summary}\n"
                    if key_pts:
                        block += f"【要点】{', '.join(key_pts[:5])}\n"
                    if raw:
                        block += f"【原文】\n{raw}\n"
                    source_blocks.append(block)
                content_summary = "\n---\n".join(source_blocks)
            else:
                parts = []
                body = content_data.get("content", "")
                if body:
                    parts.append(body[:4000])
                summary = content_data.get("summary", "")
                if summary:
                    label = "要約: " if self.language == "ja" else "Summary: "
                    parts.append(label + summary[:600])
                if sources:
                    label = "情報源:" if self.language == "ja" else "Sources:"
                    parts.append(label + "\n" + "\n".join(
                        f"- {s}" for s in sources[:10]
                    ))
                content_summary = "\n\n".join(parts) if parts else no_info
        else:
            sources = []
            content_summary = str(content_data)[:4000] if content_data else no_info

        length_instruction_ja = ""
        length_instruction_en = ""
        if self._chapter_target_chars:
            length_instruction_ja = f"\n8. このセクションは約{self._chapter_target_chars}文字を目安に執筆する"
            length_instruction_en = f"\n8. Target approximately {self._chapter_target_chars} characters for this section"

        if self.language == "ja":
            prompt = f"""{context_prompt}

【現在執筆中のセクション】
セクション番号: {section_number}
タイトル: {section_title}
説明: {section_description}

【収集された情報】
{content_summary}

【執筆指示】
上記の情報を基に、このセクションの内容を執筆してください。

重要な注意点：
1. 用語統一ルールに従い、一貫した用語を使用する
2. 前章までの内容と矛盾しないようにする
3. 文体・スタイル指示に従う
4. 事実に基づいた記述を行い、推測は「〜と考えられる」などの表現で明示する
5. 必要に応じて前章を参照する
6. 情報の出典を示すため、該当箇所に [SOURCE N] の形式で引用番号を付与する（Nは上記情報のSOURCE番号）
7. 本文は散文パラグラフで記述する。箇条書き（-、*、1.）や表は、比較一覧・手順・仕様など列挙が本質的に適切な場合のみ使用し、通常の説明・分析・考察は必ず文章で書く{length_instruction_ja}

【出力形式】
まず本文をマークダウンでそのまま書く（章見出し「## {section_number}. {section_title}」から始める）。
JSONやコードブロックで囲まない。

本文の後、最後に次の区切り記号とメタ情報を必ず出力する:

{CHAPTER_META_DELIMITER}
{{"key_points": ["要点1", "要点2", "要点3"], "terms_used": ["使用した専門用語1"], "facts_stated": ["記述した重要な事実1"]}}"""
        else:
            prompt = f"""{context_prompt}

[CURRENT SECTION]
Section: {section_number}
Title: {section_title}
Description: {section_description}

[GATHERED INFORMATION]
{content_summary}

[WRITING INSTRUCTIONS]
Write this section based on the above information.

Important notes:
1. Follow terminology consistency rules
2. Do not contradict previous chapters
3. Follow style instructions
4. Base writing on facts; phrase speculation naturally as your assessment
5. Reference previous chapters when appropriate
6. Include citation markers [SOURCE N] in the text where information from sources is used (N corresponds to the SOURCE number above)
7. Write in continuous prose paragraphs. Use bullet lists or tables only for comparisons, specifications, or steps where enumeration is inherently appropriate{length_instruction_en}

[OUTPUT FORMAT]
Write the body directly as markdown (start with the heading "## {section_number}. {section_title}").
Do not wrap it in JSON or a code block.

After the body, output this delimiter followed by metadata:

{CHAPTER_META_DELIMITER}
{{"key_points": ["point1", "point2", "point3"], "terms_used": ["technical term1"], "facts_stated": ["fact1"]}}"""

        try:
            response = self.llm.generate(prompt)

            # Guard against None/empty response (OpenAI API can return content=None)
            if not response or not response.content:
                _msg = (f"LLM returned empty response for section {section_number} "
                        f"'{section_title}'. This chapter has placeholder content only.")
                ResearchWarnings.get_instance().add(
                    ResearchWarnings.CRITICAL, "ReportGeneratorV2", _msg)
                if self.language == "ja":
                    placeholder = (f"## {section_number}. {section_title}\n\n"
                                   f"**[警告]** このセクションの内容を生成できませんでした"
                                   f"（LLM応答が空でした）。")
                else:
                    placeholder = (f"## {section_number}. {section_title}\n\n"
                                   f"**[WARNING]** Content generation failed for this section "
                                   f"(LLM returned empty response).")
                return ChapterContent(
                    section_number=section_number,
                    section_title=section_title,
                    content=placeholder,
                    word_count=0,
                    is_draft=True,
                )

            body, meta = split_prose_and_meta(response.content, CHAPTER_META_DELIMITER)

            if not body or not body.strip():
                raise ValueError("Empty chapter body in LLM response")

            # Short-chapter guard: a draft far below the target usually means
            # the LLM under-delivered, not that evidence is missing. Retry
            # once with an explicit expansion prompt before accepting it.
            min_chars = (
                int(self._chapter_target_chars * self.SHORT_CHAPTER_RATIO)
                if self._chapter_target_chars else 800
            )
            if len(body.strip()) < min_chars:
                body, meta = self._expand_short_chapter(
                    section_number=section_number,
                    section_title=section_title,
                    draft_body=body,
                    draft_meta=meta,
                    content_summary=content_summary,
                    context=context,
                    min_chars=min_chars,
                )

            # Note-style label prefixes (結論： etc.) break the prose flow
            body = strip_label_prefixes(body)

            return ChapterContent(
                section_number=section_number,
                section_title=section_title,
                content=body,
                word_count=len(body),
                key_points=meta.get("key_points", []),
                terms_used=meta.get("terms_used", []),
                facts_stated=meta.get("facts_stated", []),
                citations=sources,
                is_draft=True,
            )

        except Exception as e:
            print(f"[ReportGeneratorV2] Chapter generation failed for section {section_number}: {e}")
            return ChapterContent(
                section_number=section_number,
                section_title=section_title,
                content=f"## {section_number}. {section_title}\n\n（生成エラー: {e}）",
                word_count=0,
                is_draft=True,
            )

    def _expand_short_chapter(
        self,
        section_number: str,
        section_title: str,
        draft_body: str,
        draft_meta: Dict[str, Any],
        content_summary: str,
        context: ReportContext,
        min_chars: int,
    ):
        """Retry a too-short chapter once with an explicit expansion prompt.

        Returns (body, meta): the expanded version when the retry produced a
        longer chapter, otherwise the original draft. Records a
        ResearchWarning when the chapter is still below min_chars afterwards,
        so the shortfall is visible in the final report instead of silent.
        """
        target = self._chapter_target_chars or (
            self.DEFAULT_CHAPTER_CHARS_JA if self.language == "ja"
            else self.DEFAULT_CHAPTER_CHARS_EN
        )
        print(f"[ReportGeneratorV2] Section {section_number} draft is short "
              f"({len(draft_body)} chars < {min_chars}); retrying with an "
              f"expansion prompt")

        style_instructions = context.get_style_instructions()

        if self.language == "ja":
            prompt = f"""以下はレポートの章「{section_number}. {section_title}」の草稿ですが、短すぎます（現在約{len(draft_body)}字、目標約{target}字）。

【草稿】
{draft_body}

【利用可能な情報源】
{content_summary}

{style_instructions}

【指示】
1. 情報源の具体的な数値・事実・引用を活用し、目標分量まで内容を拡充して書き直す
2. 引用番号 [SOURCE N] は維持・追加する（この表記は一字一句変えない）
3. 情報源にない情報を捏造しない。同じ内容の水増しではなく、根拠・背景・比較・分析・示唆を加えて深める
4. 章見出し「## {section_number}. {section_title}」から始め、本文をマークダウンでそのまま書く（JSONやコードブロックで囲まない）

本文の後、最後に次の区切り記号とメタ情報を必ず出力する:

{CHAPTER_META_DELIMITER}
{{"key_points": ["要点1", "要点2"], "terms_used": ["使用した専門用語1"], "facts_stated": ["記述した事実1"]}}"""
        else:
            prompt = f"""The following draft of report chapter "{section_number}. {section_title}" is too short (about {len(draft_body)} chars now, target about {target} chars).

[DRAFT]
{draft_body}

[AVAILABLE SOURCES]
{content_summary}

{style_instructions}

[INSTRUCTIONS]
1. Rewrite and expand the chapter toward the target length using concrete figures, facts, and quotes from the sources
2. Keep and add [SOURCE N] citation markers exactly as-is
3. Do not fabricate information absent from the sources; deepen with evidence, background, comparison, analysis, and implications rather than padding
4. Start with the heading "## {section_number}. {section_title}" and write the body directly as markdown (no JSON, no code block)

After the body, output this delimiter followed by metadata:

{CHAPTER_META_DELIMITER}
{{"key_points": ["point1", "point2"], "terms_used": ["technical term1"], "facts_stated": ["fact1"]}}"""

        try:
            response = self.llm.generate(prompt)
            if response and response.content:
                new_body, new_meta = split_prose_and_meta(
                    response.content, CHAPTER_META_DELIMITER
                )
                if new_body and len(new_body.strip()) > len(draft_body.strip()):
                    draft_body = new_body
                    draft_meta = new_meta or draft_meta
        except Exception as e:
            print(f"[ReportGeneratorV2] Expansion retry failed for section "
                  f"{section_number}: {e}")

        if len(draft_body.strip()) < min_chars:
            ResearchWarnings.get_instance().add(
                ResearchWarnings.MEDIUM,
                "ReportGeneratorV2",
                f"Section {section_number} '{section_title}' remains short even "
                f"after an expansion retry ({len(draft_body)} chars, minimum "
                f"{min_chars}, target {target}). The gathered evidence for this "
                f"section is likely insufficient — consider more research "
                f"iterations, more pages per query, or gap-fill rounds.",
            )

        return draft_body, draft_meta

    def _extract_chapter_summary(self, chapter: ChapterContent) -> str:
        """Extract a summary from chapter content."""
        content = chapter.content

        # Try to get first paragraph or first 300 chars
        paragraphs = content.split("\n\n")
        for p in paragraphs:
            p = p.strip()
            if len(p) > 50 and not p.startswith("#"):
                return p[:300] + ("..." if len(p) > 300 else "")

        return content[:300] + ("..." if len(content) > 300 else "")

    def _refine_chapters(
        self,
        chapters: Dict[str, ChapterContent],
        consistency_report: ConsistencyReport,
        context: ReportContext,
    ) -> Dict[str, ChapterContent]:
        """Refine chapters based on consistency issues."""

        # Get sections that need refinement
        sections_to_refine = set()
        for issue in consistency_report.issues:
            if issue.section and issue.section != "全体":
                sections_to_refine.add(issue.section)
            for related in issue.related_sections:
                sections_to_refine.add(related)

        # Refine each affected section
        for section_num in sections_to_refine:
            if section_num not in chapters:
                continue

            chapter = chapters[section_num]
            issues_for_section = consistency_report.get_issues_by_section(section_num)

            if not issues_for_section:
                continue

            # Build refinement prompt
            issues_text = "\n".join([
                f"- [{i.issue_type.value}] {i.description}"
                for i in issues_for_section
            ])

            if self.language == "ja":
                prompt = f"""以下の章を修正してください。

【現在の内容】
{chapter.content}

【検出された問題】
{issues_text}

【修正指示】
上記の問題を解決するように内容を修正してください。
用語の統一、矛盾の解消、文体の調整を行ってください。

修正後の本文のみを出力してください（JSON不要）:"""
            else:
                prompt = f"""Please revise the following chapter.

[CURRENT CONTENT]
{chapter.content}

[DETECTED ISSUES]
{issues_text}

[REVISION INSTRUCTIONS]
Fix the above issues. Unify terminology, resolve contradictions, adjust style.

Output only the revised content (no JSON):"""

            try:
                response = self.llm.generate(prompt)
                refined_content = response.content.strip()

                # Update chapter
                chapter.content = strip_label_prefixes(refined_content)
                chapter.word_count = len(refined_content)
                chapter.is_draft = False

            except Exception as e:
                print(f"[ReportGeneratorV2] Refinement failed for {section_num}: {e}")
                ResearchWarnings.get_instance().add(
                    ResearchWarnings.MEDIUM,
                    "ReportGeneratorV2",
                    f"Consistency refinement failed for section {section_num}. "
                    f"Chapter remains as draft with potential consistency issues. "
                    f"Error: {e}",
                )

        return chapters

    def _polish_chapters(
        self,
        chapters: Dict[str, ChapterContent],
        context: ReportContext,
    ) -> Dict[str, ChapterContent]:
        """
        Polish all chapters for naturalness (one LLM call per chapter).

        Facts, numbers, citations and heading structure are kept intact; only
        the prose quality is improved. A rewrite whose length falls outside
        60%-150% of the original is rejected to guard against summarizing or
        inflating rewrites.
        """
        style_instructions = context.get_style_instructions()
        prev_tail = ""

        for section_num in sorted(chapters.keys()):
            chapter = chapters[section_num]
            if not chapter.content or not chapter.content.strip():
                continue

            if self.language == "ja":
                prompt = f"""あなたは日本語の報告書の推敲者です。以下の章の内容・事実・構成は変えずに、日本語としての自然さだけを磨いてください。

{style_instructions}

【直前の章の末尾（つながりの参考。書き換え対象ではない）】
{prev_tail or "（これが最初の章）"}

【推敲対象の章】
{chapter.content}

【推敲ルール】
1. 事実・数値・固有名詞・出典表記・見出し構成は一切変更しない。情報の追加・削除もしない。
2. 翻訳調・単調な文末・不自然な語順・冗長表現を直し、段落内と段落間の流れを滑らかにする。
3. 文体（です・ます調／である調）は上記スタイル指示に統一する。
4. 「結論：」「要点：」のようなラベル書きや細切れの短文の羅列があれば、通常の流れる文章に書き直す。段落同士は接続表現でつなぎ、章全体を調子の通った文章にする。
5. 修正が不要な文はそのまま残してよい。

推敲後の本文のみを出力（前置き・JSON・コードブロック不要）:"""
            else:
                prompt = f"""You are a prose editor for a research report. Polish the following chapter for naturalness only, without changing its content, facts, or structure.

{style_instructions}

[END OF PREVIOUS CHAPTER (for transition context; not to be rewritten)]
{prev_tail or "(this is the first chapter)"}

[CHAPTER TO POLISH]
{chapter.content}

[EDITING RULES]
1. Do not change facts, numbers, proper nouns, citations, or heading structure. Do not add or remove information.
2. Fix awkward phrasing, monotonous sentence endings, and redundancy; smooth the flow within and between paragraphs.
3. Keep the style consistent with the instructions above.
4. Sentences that need no edits may be kept as-is.

Output only the polished body (no preamble, no JSON, no code block):"""

            try:
                response = self.llm.generate(prompt)
                polished = response.content.strip()

                original_len = len(chapter.content)
                if polished and 0.6 * original_len <= len(polished) <= 1.5 * original_len:
                    polished = strip_label_prefixes(polished)
                    chapter.content = polished
                    chapter.word_count = len(polished)
                    chapter.is_draft = False
                else:
                    print(f"[ReportGeneratorV2] Polish rejected for {section_num} "
                          f"(length {len(polished)} vs original {original_len})")
                    ResearchWarnings.get_instance().add(
                        ResearchWarnings.LOW,
                        "ReportGeneratorV2",
                        f"第{section_num}章の推敲結果が長さ検査で棄却され、"
                        f"下書きのまま出力されています。",
                    )
            except Exception as e:
                print(f"[ReportGeneratorV2] Polish failed for {section_num}: {e}")
                ResearchWarnings.get_instance().add(
                    ResearchWarnings.LOW,
                    "ReportGeneratorV2",
                    f"第{section_num}章の推敲パスが失敗し、下書きのまま"
                    f"出力されています。Error: {e}",
                )

            prev_tail = chapter.content[-300:]

        return chapters

    def _get_sections_from_plan(self, research_plan: Any) -> List[Dict[str, str]]:
        """Extract sections from research plan."""
        sections = []

        if hasattr(research_plan, 'table_of_contents'):
            toc = research_plan.table_of_contents
            if hasattr(toc, 'items'):
                for item in toc.items:
                    sections.append({
                        "section": item.section,
                        "title": item.title,
                        "description": getattr(item, 'description', ''),
                    })
                    # Add subsections
                    if hasattr(item, 'subsections'):
                        for sub in item.subsections:
                            sections.append({
                                "section": sub.section,
                                "title": sub.title,
                                "description": getattr(sub, 'description', ''),
                            })

        return sections

    def _report_progress(self, message: str, current: int, total: int) -> None:
        """Report progress if callback is set."""
        if self.progress_callback:
            self.progress_callback(message, current, total)
        else:
            print(f"[ReportGeneratorV2] {message} ({current}/{total})")

    def generate_final_document(
        self,
        result: GenerationResult,
        include_glossary: bool = True,
        include_consistency_summary: bool = False,
        evidence_locker=None,
    ) -> str:
        """
        Generate the final document from generation result.

        Args:
            result: GenerationResult from generate_report
            include_glossary: Include glossary section
            include_consistency_summary: Include consistency check summary
            evidence_locker: EvidenceLocker for generating references section

        Returns:
            Complete document as markdown string
        """
        # Build URL-to-reference mapping for citation renumbering
        evidence_list = []
        url_to_ref = {}
        if evidence_locker is not None:
            evidence_list = evidence_locker.get_all_evidence()
            for i, evidence in enumerate(evidence_list, 1):
                url_to_ref[evidence.url] = i

        lines = []

        # Title
        lines.append(f"# {result.context.report_title}")
        lines.append("")
        lines.append(f"*Generated: {datetime.now().strftime('%Y-%m-%d')}*")
        lines.append("")

        # Table of Contents (use natural sort for proper section ordering)
        lines.append("## 目次" if self.language == "ja" else "## Table of Contents")
        lines.append("")
        for section_num, chapter in sorted(result.chapters.items(), key=_section_sort_key):
            lines.append(f"- {section_num}. {chapter.section_title}")
        lines.append("")

        # Chapters (use natural sort: 1, 1.1, 2, 2.1, ... 10, 10.1)
        for section_num, chapter in sorted(result.chapters.items(), key=_section_sort_key):
            content = chapter.content
            # Renumber [SOURCE N] citations to final reference numbers
            if url_to_ref and chapter.citations:
                content = self._renumber_citations(content, chapter.citations, url_to_ref)
            lines.append(content)
            lines.append("")

        # Glossary
        if include_glossary and result.context.glossary:
            glossary_section = self.glossary_manager.generate_glossary_section(
                {k: {
                    "term": v.term,
                    "definition": v.definition,
                    "aliases": v.aliases,
                } for k, v in result.context.glossary.items()},
                title="用語集" if self.language == "ja" else "Glossary"
            )
            lines.append(glossary_section)

        # Consistency summary
        if include_consistency_summary and result.consistency_report:
            summary = self.consistency_checker.generate_summary(result.consistency_report)
            lines.append(summary)

        # References section
        if evidence_list:
            lines.append("")
            lines.append("---")
            lines.append("")
            lines.append("## 参考文献" if self.language == "ja" else "## References")
            lines.append("")
            for i, evidence in enumerate(evidence_list, 1):
                quality_badge = self._get_quality_badge(evidence.quality_category)
                lines.append(f"{i}. {quality_badge} {evidence.citation_text}")
            lines.append("")

        return "\n".join(lines)

    @staticmethod
    def _get_quality_badge(quality) -> str:
        """Get quality badge string for a citation."""
        # Import here to avoid circular imports
        try:
            from ...evidence.locker import QualityCategory
            badges = {
                QualityCategory.AUTHORITATIVE: "[A]",
                QualityCategory.HIGH: "[H]",
                QualityCategory.MEDIUM: "[M]",
                QualityCategory.LOW: "[L]",
                QualityCategory.UNVERIFIED: "[?]",
            }
            return badges.get(quality, "")
        except (ImportError, AttributeError):
            return ""

    @staticmethod
    def _renumber_citations(
        content: str,
        section_sources: List[str],
        url_to_ref: Dict[str, int],
    ) -> str:
        """Renumber [SOURCE N] citations to final reference numbers [N].

        Args:
            content: Chapter content with [SOURCE N] markers
            section_sources: List of source URLs used for this chapter
            url_to_ref: Mapping from URL to final reference number

        Returns:
            Content with renumbered citation markers
        """
        if not content or not section_sources:
            return content

        def replace_citation(match):
            original_num = int(match.group(1))
            source_index = original_num - 1
            if 0 <= source_index < len(section_sources):
                source_url = section_sources[source_index]
                if source_url in url_to_ref:
                    return f"[{url_to_ref[source_url]}]"
            return f"[?{original_num}]"

        return re.sub(r'\[SOURCE:?\s*(\d+)\]', replace_citation, content)

    def save_report(
        self,
        markdown_content: str,
        output_dir: Path,
        filename: str,
        format: str = "markdown",
        strict_format: bool = False,
    ) -> Path:
        """
        Save the generated report in the specified format.

        Args:
            markdown_content: The report content as markdown string
            output_dir: Directory to save the report
            filename: Base filename (without extension)
            format: Output format ('markdown', 'docx', 'pdf', 'html')
            strict_format: If True, raise ReportFormatError instead of falling
                          back to markdown when DOCX generation fails.

        Returns:
            Path to the saved report file

        Raises:
            ReportFormatError: When strict_format=True and DOCX generation fails.
        """
        output_dir = Path(output_dir)
        output_dir.mkdir(parents=True, exist_ok=True)

        format_lower = format.lower() if isinstance(format, str) else format.value.lower()

        if format_lower == "docx":
            return self._save_as_docx(markdown_content, output_dir, filename, strict_format=strict_format)
        elif format_lower == "html":
            return self._save_as_html(markdown_content, output_dir, filename)
        elif format_lower == "pdf":
            # PDF: save as markdown with note (PDF generation requires additional setup)
            md_path = self._save_as_markdown(markdown_content, output_dir, filename)
            ResearchWarnings.get_instance().add(
                ResearchWarnings.CRITICAL,
                "ReportGeneratorV2",
                "PDF format requested but not supported in V2 generator. "
                "Report saved as markdown (.md) instead. "
                "Use report_generator_version='v1' for PDF support.",
            )
            return md_path
        else:
            return self._save_as_markdown(markdown_content, output_dir, filename)

    def _save_as_markdown(self, content: str, output_dir: Path, filename: str) -> Path:
        """Save report as markdown file."""
        filepath = output_dir / f"{filename}.md"
        filepath.write_text(content, encoding="utf-8")
        return filepath

    def _save_as_docx(
        self,
        markdown_content: str,
        output_dir: Path,
        filename: str,
        strict_format: bool = False,
    ) -> Path:
        """Convert markdown content to DOCX and save.

        Uses 3-layer defense against XML illegal characters:
        1. Pre-sanitize the entire markdown content
        2. Sanitize every text insertion into the Document
        3. BytesIO pre-validation before writing to disk

        Args:
            markdown_content: Markdown text to convert
            output_dir: Output directory
            filename: Base filename (without extension)
            strict_format: If True, raise ReportFormatError on failure instead
                          of falling back to markdown.

        Returns:
            Path to the saved DOCX file.

        Raises:
            ReportFormatError: When strict_format=True and DOCX generation fails.
        """
        try:
            from docx import Document
            from docx.shared import Pt, Inches, Cm
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            if strict_format:
                raise ReportFormatError(
                    "python-docx is not installed. Cannot generate DOCX in strict_format mode. "
                    "Install with: pip install python-docx"
                )
            logger.error("[ReportGeneratorV2] python-docx not installed. Falling back to markdown.")
            ResearchWarnings.get_instance().add(
                ResearchWarnings.CRITICAL,
                "ReportGeneratorV2",
                "DOCX format requested but python-docx is not installed. "
                "Report saved as markdown (.md) instead. "
                "Install with: pip install python-docx",
            )
            return self._save_as_markdown(markdown_content, output_dir, filename)

        # Layer 1: Pre-sanitize the entire markdown content
        markdown_content = _sanitize_for_xml(markdown_content)

        doc = Document()
        lines = markdown_content.split("\n")
        i = 0

        while i < len(lines):
            line = lines[i]
            stripped = line.strip()

            # Empty line
            if not stripped:
                i += 1
                continue

            # Layer 2: Sanitize every text fragment before insertion
            stripped = _sanitize_for_xml(stripped)

            # Headings
            if stripped.startswith("#"):
                level, text = self._parse_heading(stripped)
                text = _sanitize_for_xml(text)
                try:
                    doc.add_heading(text, level=level)
                except Exception as e:
                    logger.warning(f"[DOCX] Heading add failed: {e}, adding as paragraph")
                    doc.add_paragraph(text)
                i += 1
                continue

            # Markdown table detection (lines starting with |)
            if stripped.startswith("|") and "|" in stripped[1:]:
                table_lines = []
                while i < len(lines) and lines[i].strip().startswith("|"):
                    table_lines.append(_sanitize_for_xml(lines[i].strip()))
                    i += 1
                if table_lines:
                    self._add_markdown_table_to_docx(doc, table_lines)
                continue

            # Unordered list items
            if re.match(r'^[-*+]\s', stripped):
                text = _sanitize_for_xml(re.sub(r'^[-*+]\s+', '', stripped))
                try:
                    para = doc.add_paragraph(style='List Bullet')
                except KeyError:
                    para = doc.add_paragraph()
                    text = "- " + text
                self._add_formatted_runs(para, text)
                i += 1
                continue

            # Ordered list items
            if re.match(r'^\d+\.\s', stripped):
                text = _sanitize_for_xml(re.sub(r'^\d+\.\s+', '', stripped))
                try:
                    para = doc.add_paragraph(style='List Number')
                except KeyError:
                    para = doc.add_paragraph()
                    text = stripped  # Keep the number prefix
                self._add_formatted_runs(para, text)
                i += 1
                continue

            # Italic metadata line (e.g., *Generated: 2024-01-01*)
            if stripped.startswith("*") and stripped.endswith("*") and not stripped.startswith("**"):
                para = doc.add_paragraph()
                run = para.add_run(_sanitize_for_xml(stripped.strip("*")))
                run.italic = True
                run.font.size = Pt(9)
                para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                i += 1
                continue

            # Image reference: ![alt](path)
            img_match = re.match(r'!\[([^\]]*)\]\(([^)]+)\)', stripped)
            if img_match:
                alt_text = _sanitize_for_xml(img_match.group(1))
                img_path = img_match.group(2)
                self._add_image_to_docx(doc, img_path, alt_text)
                i += 1
                continue

            # Regular paragraph: collect consecutive non-empty, non-special lines
            para_lines = []
            start_i = i  # Track starting position to prevent infinite loop
            while i < len(lines):
                l = lines[i].strip()
                if not l or l.startswith("#") or re.match(r'^[-*+]\s', l) or re.match(r'^\d+\.\s', l) or l.startswith("|"):
                    break
                # Check for image reference
                if re.match(r'!\[([^\]]*)\]\(([^)]+)\)', l):
                    break
                para_lines.append(_sanitize_for_xml(l))
                i += 1

            # Safety: if no lines were consumed, force advance to avoid infinite loop
            if i == start_i:
                i += 1

            if para_lines:
                para = doc.add_paragraph()
                self._add_formatted_runs(para, " ".join(para_lines))
            continue

        # Layer 3: BytesIO pre-validation - test XML serialization before writing to disk
        filepath = output_dir / f"{filename}.docx"
        from io import BytesIO

        try:
            buf = BytesIO()
            doc.save(buf)
            # Serialization succeeded - write to disk
            buf.seek(0)
            with open(filepath, 'wb') as f:
                f.write(buf.read())
        except Exception as save_error:
            logger.error(f"[ReportGeneratorV2] BytesIO pre-validation failed: {save_error}")
            # Retry: aggressive per-paragraph sanitization
            try:
                logger.info("[ReportGeneratorV2] Retrying with per-paragraph sanitization...")
                self._sanitize_docx_paragraphs(doc)
                self._sanitize_docx_tables(doc)
                buf2 = BytesIO()
                doc.save(buf2)
                buf2.seek(0)
                with open(filepath, 'wb') as f:
                    f.write(buf2.read())
                logger.info("[ReportGeneratorV2] Retry succeeded after per-element sanitization.")
            except Exception as retry_error:
                if strict_format:
                    # Save debug markdown so content is not lost
                    debug_md_path = self._save_as_markdown(markdown_content, output_dir, f"{filename}_debug")
                    raise ReportFormatError(
                        f"DOCX generation failed even after sanitization. "
                        f"Original error: {save_error} / Retry error: {retry_error}. "
                        f"Debug markdown saved to: {debug_md_path}",
                        debug_md_path=str(debug_md_path),
                        original_error=retry_error,
                    )
                logger.error(
                    f"[ReportGeneratorV2] Retry also failed: {retry_error}. "
                    f"Falling back to markdown."
                )
                return self._save_as_markdown(markdown_content, output_dir, filename)
        return filepath

    @staticmethod
    def _sanitize_docx_paragraphs(doc):
        """Sanitize all paragraph text in a Document to remove illegal XML chars."""
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                if run.text:
                    run.text = _sanitize_for_xml(run.text)

    @staticmethod
    def _sanitize_docx_tables(doc):
        """Sanitize all table cell text in a Document to remove illegal XML chars."""
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            if run.text:
                                run.text = _sanitize_for_xml(run.text)

    def _add_markdown_table_to_docx(self, doc, table_lines: list):
        """Convert markdown table lines to a DOCX table."""
        try:
            from docx.shared import Pt

            # Parse header row
            rows_data = []
            for line in table_lines:
                cells = [c.strip() for c in line.strip("|").split("|")]
                # Skip separator row (contains only -, :, spaces)
                if all(re.match(r'^[-:]+$', c.strip()) for c in cells if c.strip()):
                    continue
                rows_data.append(cells)

            if not rows_data:
                return

            num_cols = max(len(row) for row in rows_data)
            table = doc.add_table(rows=len(rows_data), cols=num_cols)
            table.style = "Table Grid"

            for row_idx, row_data in enumerate(rows_data):
                for col_idx, cell_text in enumerate(row_data):
                    if col_idx < num_cols:
                        cell = table.rows[row_idx].cells[col_idx]
                        cell.text = _sanitize_for_xml(cell_text.strip())
                        # Bold header row
                        if row_idx == 0:
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    run.bold = True
                                    run.font.size = Pt(9)

            doc.add_paragraph()  # Spacing after table
        except Exception as e:
            print(f"[ReportGeneratorV2] Table insertion failed: {e}")

    def _add_image_to_docx(self, doc, img_path: str, alt_text: str = ""):
        """Add an image to the DOCX document."""
        try:
            from docx.shared import Inches
            import os

            # Check if file exists (handle both absolute and relative paths)
            if os.path.exists(img_path):
                doc.add_picture(img_path, width=Inches(5.5))
            elif os.path.exists(os.path.join(str(self.output_dir if hasattr(self, 'output_dir') else '.'), img_path)):
                full_path = os.path.join(str(self.output_dir if hasattr(self, 'output_dir') else '.'), img_path)
                doc.add_picture(full_path, width=Inches(5.5))

            if alt_text:
                para = doc.add_paragraph()
                run = para.add_run(alt_text)
                run.italic = True
                from docx.shared import Pt
                run.font.size = Pt(9)
        except Exception as e:
            # If image can't be added, add alt text as placeholder
            if alt_text:
                doc.add_paragraph(f"[Image: {alt_text}]")

    def _save_as_html(self, markdown_content: str, output_dir: Path, filename: str) -> Path:
        """Convert markdown content to HTML and save."""
        # Simple markdown to HTML conversion
        html_body = self._markdown_to_html(markdown_content)
        html_doc = f"""<!DOCTYPE html>
<html lang="{self.language}">
<head>
<meta charset="utf-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<style>
body {{ font-family: 'Helvetica Neue', Arial, 'Hiragino Kaku Gothic ProN', sans-serif; max-width: 900px; margin: 0 auto; padding: 2em; line-height: 1.8; }}
h1 {{ border-bottom: 2px solid #333; padding-bottom: 0.3em; }}
h2 {{ border-bottom: 1px solid #ccc; padding-bottom: 0.2em; }}
ul, ol {{ padding-left: 1.5em; }}
</style>
</head>
<body>
{html_body}
</body>
</html>"""
        filepath = output_dir / f"{filename}.html"
        filepath.write_text(html_doc, encoding="utf-8")
        return filepath

    @staticmethod
    def _parse_heading(line: str) -> tuple:
        """Parse a markdown heading line into (level, text)."""
        match = re.match(r'^(#{1,6})\s+(.*)', line)
        if match:
            level = min(len(match.group(1)), 4)  # docx supports levels 0-4
            return level - 1, match.group(2).strip()
        return 0, line.strip("#").strip()

    @staticmethod
    def _add_formatted_runs(paragraph, text: str):
        """Add text with bold/italic formatting as runs to a paragraph."""
        text = _sanitize_for_xml(text)
        # Split by bold (**text**) and italic (*text*) markers
        parts = re.split(r'(\*\*[^*]+\*\*|\*[^*]+\*)', text)
        for part in parts:
            if not part:
                continue
            if part.startswith("**") and part.endswith("**"):
                run = paragraph.add_run(_sanitize_for_xml(part[2:-2]))
                run.bold = True
            elif part.startswith("*") and part.endswith("*"):
                run = paragraph.add_run(_sanitize_for_xml(part[1:-1]))
                run.italic = True
            else:
                paragraph.add_run(_sanitize_for_xml(part))

    @staticmethod
    def _markdown_to_html(markdown_text: str) -> str:
        """Simple markdown to HTML conversion."""
        lines = markdown_text.split("\n")
        html_lines = []
        in_list = False
        list_type = None

        for line in lines:
            stripped = line.strip()

            if not stripped:
                if in_list:
                    html_lines.append(f"</{list_type}>")
                    in_list = False
                    list_type = None
                html_lines.append("")
                continue

            # Headings
            heading_match = re.match(r'^(#{1,6})\s+(.*)', stripped)
            if heading_match:
                if in_list:
                    html_lines.append(f"</{list_type}>")
                    in_list = False
                level = len(heading_match.group(1))
                text = heading_match.group(2)
                html_lines.append(f"<h{level}>{text}</h{level}>")
                continue

            # Unordered list
            if re.match(r'^[-*+]\s', stripped):
                text = re.sub(r'^[-*+]\s+', '', stripped)
                if not in_list or list_type != "ul":
                    if in_list:
                        html_lines.append(f"</{list_type}>")
                    html_lines.append("<ul>")
                    in_list = True
                    list_type = "ul"
                html_lines.append(f"  <li>{text}</li>")
                continue

            # Ordered list
            ol_match = re.match(r'^\d+\.\s+(.*)', stripped)
            if ol_match:
                text = ol_match.group(1)
                if not in_list or list_type != "ol":
                    if in_list:
                        html_lines.append(f"</{list_type}>")
                    html_lines.append("<ol>")
                    in_list = True
                    list_type = "ol"
                html_lines.append(f"  <li>{text}</li>")
                continue

            # Regular paragraph
            if in_list:
                html_lines.append(f"</{list_type}>")
                in_list = False
                list_type = None
            # Apply inline formatting
            formatted = re.sub(r'\*\*([^*]+)\*\*', r'<strong>\1</strong>', stripped)
            formatted = re.sub(r'\*([^*]+)\*', r'<em>\1</em>', formatted)
            html_lines.append(f"<p>{formatted}</p>")

        if in_list:
            html_lines.append(f"</{list_type}>")

        return "\n".join(html_lines)
