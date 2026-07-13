"""
Figure and Table Generator - Add figures and tables to research reports.

This module provides functionality to:
- Extract relevant images from referenced web pages
- Identify and extract numerical data for tables (LLM + pattern + HTML parsing)
- Generate matplotlib charts with LLM-recommended chart types
- Support PIE, LINE, BAR, AREA, SCATTER, STACKED_BAR, HORIZONTAL_BAR charts
- Add figures and tables to existing reports (Markdown, DOCX, PDF)
"""

import io
import os
import re
import json
import base64
import hashlib
import tempfile
from dataclasses import dataclass, field
from datetime import datetime
from enum import Enum
from pathlib import Path
from typing import Optional, Dict, Any, List, Tuple
from urllib.parse import urlparse, urljoin

import requests


class FigureType(str, Enum):
    """Types of figures."""
    IMAGE = "image"          # Image from web source
    CHART = "chart"          # Generated matplotlib chart
    DIAGRAM = "diagram"      # Diagram or flowchart


class ChartType(str, Enum):
    """Types of charts that can be generated."""
    LINE = "line"
    BAR = "bar"
    PIE = "pie"
    AREA = "area"
    SCATTER = "scatter"
    STACKED_BAR = "stacked_bar"
    HORIZONTAL_BAR = "horizontal_bar"


@dataclass
class Figure:
    """A figure (image or chart) for the report."""
    figure_id: str
    figure_type: FigureType
    title: str
    caption: str
    source_url: Optional[str] = None
    source_title: Optional[str] = None
    section_id: str = ""
    image_path: Optional[Path] = None
    image_data: Optional[bytes] = None
    alt_text: str = ""
    width: Optional[int] = None
    height: Optional[int] = None

    def to_dict(self) -> Dict[str, Any]:
        """Convert to dictionary."""
        return {
            "figure_id": self.figure_id,
            "figure_type": self.figure_type.value,
            "title": self.title,
            "caption": self.caption,
            "source_url": self.source_url,
            "source_title": self.source_title,
            "section_id": self.section_id,
            "image_path": str(self.image_path) if self.image_path else None,
            "alt_text": self.alt_text,
            "width": self.width,
            "height": self.height,
        }


@dataclass
class TableData:
    """Data for a table."""
    table_id: str
    title: str
    caption: str
    headers: List[str]
    rows: List[List[Any]]
    source_url: Optional[str] = None
    source_title: Optional[str] = None
    section_id: str = ""
    data_type: str = "general"  # general, time_series, comparison, distribution, correlation
    unit: Optional[str] = None

    def to_dict(self) -> Dict[str, Any]:
        """Convert to dictionary."""
        return {
            "table_id": self.table_id,
            "title": self.title,
            "caption": self.caption,
            "headers": self.headers,
            "rows": self.rows,
            "source_url": self.source_url,
            "source_title": self.source_title,
            "section_id": self.section_id,
            "data_type": self.data_type,
            "unit": self.unit,
        }


@dataclass
class FigureTableCollection:
    """Collection of figures and tables for a report."""
    figures: List[Figure] = field(default_factory=list)
    tables: List[TableData] = field(default_factory=list)
    charts: List[Figure] = field(default_factory=list)  # Generated charts

    def get_figures_for_section(self, section_id: str) -> List[Figure]:
        """Get all figures for a section."""
        return [f for f in self.figures if f.section_id == section_id]

    def get_tables_for_section(self, section_id: str) -> List[TableData]:
        """Get all tables for a section."""
        return [t for t in self.tables if t.section_id == section_id]

    def get_charts_for_section(self, section_id: str) -> List[Figure]:
        """Get all charts for a section."""
        return [c for c in self.charts if c.section_id == section_id]

    def to_dict(self) -> Dict[str, Any]:
        """Convert to dictionary."""
        return {
            "figures": [f.to_dict() for f in self.figures],
            "tables": [t.to_dict() for t in self.tables],
            "charts": [c.to_dict() for c in self.charts],
        }


# ============================================================================
# Chart Style Configuration
# ============================================================================

# Professional color palette (colorblind-friendly)
CHART_COLORS = [
    "#4C78A8",  # Steel Blue
    "#F58518",  # Orange
    "#E45756",  # Red
    "#72B7B2",  # Teal
    "#54A24B",  # Green
    "#EECA3B",  # Yellow
    "#B279A2",  # Purple
    "#FF9DA6",  # Pink
    "#9D755D",  # Brown
    "#BAB0AC",  # Gray
]

PIE_COLORS = [
    "#4C78A8",
    "#F58518",
    "#E45756",
    "#72B7B2",
    "#54A24B",
    "#EECA3B",
    "#B279A2",
    "#FF9DA6",
    "#9D755D",
    "#BAB0AC",
]


class FigureTableGenerator:
    """
    Generate figures and tables for research reports.

    Extracts images from referenced sources and creates tables/charts
    from numerical data found in the content.
    """

    # Common image extensions
    IMAGE_EXTENSIONS = {'.jpg', '.jpeg', '.png', '.gif', '.webp', '.svg'}

    # Patterns for numerical data extraction
    NUMBER_PATTERN = re.compile(r'[\d,]+\.?\d*')
    YEAR_PATTERN = re.compile(r'(19|20)\d{2}')
    PERCENTAGE_PATTERN = re.compile(r'(\d+\.?\d*)\s*[%％]')
    CURRENCY_PATTERN = re.compile(r'[\$€¥£]\s*[\d,]+\.?\d*|[\d,]+\.?\d*\s*(円|ドル|ユーロ|億|万)')

    def __init__(
        self,
        llm_client=None,
        output_dir: Path = None,
        language: str = "ja",
        max_images_per_section: int = 2,
        image_max_width: int = 800,
        download_timeout: int = 10,
        proxies: Dict[str, str] = None,
        verify_ssl: bool = True,
        chart_library: str = "matplotlib",
    ):
        """
        Initialize FigureTableGenerator.

        Args:
            llm_client: LLM client for content analysis
            output_dir: Directory to save generated files
            language: Output language
            max_images_per_section: Maximum images per section
            image_max_width: Maximum image width in pixels
            download_timeout: Timeout for downloading images
            proxies: Proxy settings dict
            verify_ssl: Whether to verify SSL certificates
            chart_library: Chart rendering library ("matplotlib" or "seaborn";
                falls back to matplotlib when seaborn is not installed)
        """
        self.llm = llm_client
        self.output_dir = output_dir or Path("./output/figures")
        self.output_dir.mkdir(parents=True, exist_ok=True)
        self.language = language
        self.max_images_per_section = max_images_per_section
        self.image_max_width = image_max_width
        self.download_timeout = download_timeout
        self.proxies = proxies
        self.verify_ssl = verify_ssl
        self.chart_library = (chart_library or "matplotlib").lower()
        self._seaborn_warned = False

    # ========================================================================
    # Main Entry Point
    # ========================================================================

    def generate_figures_and_tables(
        self,
        session,
        evidence_locker,
        include_images: bool = True,
        include_tables: bool = True,
        include_charts: bool = True,
    ) -> FigureTableCollection:
        """
        Generate figures and tables for all sections.

        Args:
            session: Research session with content
            evidence_locker: Evidence locker with sources
            include_images: Include images from sources
            include_tables: Include extracted tables
            include_charts: Include generated charts

        Returns:
            FigureTableCollection with all figures and tables
        """
        collection = FigureTableCollection()

        # Get section information
        sections = self._get_sections(session)

        for section_id, section_data in sections.items():
            if section_id.startswith("_"):
                continue

            # Get evidence for this section
            section_evidence = evidence_locker.get_section_evidence(section_id)
            content = section_data.get("content", "")

            # Extract images from sources
            if include_images:
                images = self._extract_images_for_section(
                    section_id=section_id,
                    section_data=section_data,
                    evidence_list=section_evidence,
                )
                collection.figures.extend(images)

            # Extract numerical data and create tables
            if include_tables:
                tables = self._extract_tables_for_section(
                    section_id=section_id,
                    section_data=section_data,
                    content=content,
                    evidence_list=section_evidence,
                )
                collection.tables.extend(tables)

            # Generate charts from tables
            if include_charts and collection.tables:
                section_tables = [t for t in collection.tables if t.section_id == section_id]
                for table in section_tables:
                    chart = self._generate_chart_for_table(table)
                    if chart:
                        collection.charts.append(chart)

        return collection

    def generate_from_recommendations(
        self,
        session,
        evidence_locker,
        recommendations: List,
        include_images: bool = True,
        include_tables: bool = True,
    ) -> FigureTableCollection:
        """
        Generate figures and tables from ChartAnalyzer recommendations.

        This method uses intelligent chart recommendations that include
        insights and meaningful messages rather than mechanical table extraction.

        Args:
            session: Research session with content
            evidence_locker: Evidence locker with sources
            recommendations: List of ChartRecommendation from ChartAnalyzer
            include_images: Include images from sources
            include_tables: Include extracted tables (in addition to recommended charts)

        Returns:
            FigureTableCollection with figures, tables, and recommended charts
        """
        collection = FigureTableCollection()

        # Get section information
        sections = self._get_sections(session)

        # Process each section for images and tables
        for section_id, section_data in sections.items():
            if section_id.startswith("_"):
                continue

            section_evidence = evidence_locker.get_section_evidence(section_id)
            content = section_data.get("content", "")

            # Extract images if enabled
            if include_images:
                images = self._extract_images_for_section(
                    section_id=section_id,
                    section_data=section_data,
                    evidence_list=section_evidence,
                )
                collection.figures.extend(images)

            # Extract tables if enabled (but don't generate charts from them)
            if include_tables:
                tables = self._extract_tables_for_section(
                    section_id=section_id,
                    section_data=section_data,
                    content=content,
                    evidence_list=section_evidence,
                )
                collection.tables.extend(tables)

        # Generate charts from recommendations
        for rec in recommendations:
            chart = self._generate_chart_from_recommendation(rec)
            if chart:
                collection.charts.append(chart)

        return collection

    def _generate_chart_from_recommendation(self, recommendation) -> Optional[Figure]:
        """
        Generate a chart from a ChartRecommendation.

        Args:
            recommendation: ChartRecommendation from ChartAnalyzer

        Returns:
            Figure object with chart image, or None if failed
        """
        try:
            import matplotlib
            matplotlib.use('Agg')
            import matplotlib.pyplot as plt
            import matplotlib.font_manager as fm

            # Setup fonts and style (matplotlib or seaborn theme)
            sns = self._apply_chart_style(plt, fm)

            # Map recommendation chart type to our ChartType
            type_map = {
                "line": ChartType.LINE,
                "bar": ChartType.BAR,
                "horizontal_bar": ChartType.HORIZONTAL_BAR,
                "stacked_bar": ChartType.STACKED_BAR,
                "pie": ChartType.PIE,
                "area": ChartType.AREA,
                "scatter": ChartType.SCATTER,
                "combo": ChartType.LINE,  # Fallback to line for combo
            }

            rec_type = recommendation.chart_type.value if hasattr(recommendation.chart_type, 'value') else str(recommendation.chart_type)
            chart_type = type_map.get(rec_type, ChartType.BAR)

            # Create figure
            fig, ax = plt.subplots(figsize=(10, 6))

            # Prepare data from recommendation
            data_points = recommendation.data_points
            if not data_points:
                plt.close(fig)
                return None

            # Sort by year for time series, or by value for comparisons
            if recommendation.purpose.value in ("show_trend", "show_growth"):
                data_points = sorted(data_points, key=lambda x: x.year or 0)
                x_labels = [str(dp.year) for dp in data_points]
            else:
                data_points = sorted(data_points, key=lambda x: x.normalized_value, reverse=True)
                x_labels = [dp.subject for dp in data_points]

            values = [dp.normalized_value for dp in data_points]

            # Render based on chart type
            if chart_type == ChartType.PIE:
                self._render_pie_from_recommendation(ax, data_points)
            elif chart_type == ChartType.HORIZONTAL_BAR:
                self._render_horizontal_bar_from_recommendation(ax, data_points)
            elif chart_type == ChartType.LINE:
                if sns is not None:
                    sns.lineplot(x=list(range(len(values))), y=values,
                                 marker='o', linewidth=2, ax=ax)
                else:
                    ax.plot(range(len(values)), values, marker='o', linewidth=2, color=CHART_COLORS[0])
                ax.set_xticks(range(len(x_labels)))
                ax.set_xticklabels(x_labels, rotation=45, ha='right')
            elif chart_type == ChartType.AREA:
                ax.fill_between(range(len(values)), values, alpha=0.7, color=CHART_COLORS[0])
                ax.plot(range(len(values)), values, linewidth=1, color=CHART_COLORS[0])
                ax.set_xticks(range(len(x_labels)))
                ax.set_xticklabels(x_labels, rotation=45, ha='right')
            else:  # BAR
                if sns is not None:
                    sns.barplot(x=list(range(len(values))), y=values, ax=ax)
                    bars = ax.patches
                else:
                    bars = ax.bar(range(len(values)), values, color=CHART_COLORS[0])
                ax.set_xticks(range(len(x_labels)))
                ax.set_xticklabels(x_labels, rotation=45, ha='right')
                # Add value labels for small datasets
                if len(values) <= 8:
                    for bar, val in zip(bars, values):
                        ax.text(
                            bar.get_x() + bar.get_width() / 2,
                            bar.get_height(),
                            f'{val:,.1f}',
                            ha='center', va='bottom', fontsize=9
                        )

            # Set title and labels
            ax.set_title(recommendation.title, fontsize=14, fontweight='bold', pad=15)
            if recommendation.y_axis_label:
                ax.set_ylabel(recommendation.y_axis_label, fontsize=11)
            if recommendation.x_axis_label:
                ax.set_xlabel(recommendation.x_axis_label, fontsize=11)

            # Add insight as subtitle if available
            if recommendation.main_message:
                ax.text(
                    0.5, -0.15,
                    recommendation.main_message[:100] + "..." if len(recommendation.main_message) > 100 else recommendation.main_message,
                    transform=ax.transAxes,
                    fontsize=9,
                    ha='center',
                    style='italic',
                    wrap=True,
                )

            plt.tight_layout()

            # Save chart
            chart_filename = f"chart_{recommendation.chart_id}.png"
            chart_path = self.output_dir / chart_filename

            plt.savefig(
                chart_path,
                dpi=150,
                bbox_inches='tight',
                facecolor='white',
                edgecolor='none',
            )

            # Read image data
            with open(chart_path, 'rb') as f:
                image_data = f.read()

            plt.close(fig)

            # Create Figure object
            caption = recommendation.main_message or f"Chart: {recommendation.title}"

            return Figure(
                figure_id=f"chart_{recommendation.chart_id}",
                figure_type=FigureType.CHART,
                title=recommendation.title,
                caption=caption,
                source_url=recommendation.source_urls[0] if recommendation.source_urls else "",
                section_id=recommendation.section_id,
                image_path=chart_path,
                image_data=image_data,
                alt_text=f"Chart showing {recommendation.title}",
            )

        except Exception as e:
            import logging
            logging.getLogger(__name__).warning(
                f"Failed to generate chart from recommendation: {e}"
            )
            return None

    def _render_pie_from_recommendation(self, ax, data_points) -> None:
        """Render pie chart from recommendation data points."""
        labels = [dp.subject for dp in data_points]
        values = [dp.value for dp in data_points]

        # Filter out zero/negative values
        filtered = [(l, v) for l, v in zip(labels, values) if v > 0]
        if not filtered:
            return

        labels, values = zip(*filtered)

        colors = PIE_COLORS[:len(values)]
        wedges, texts, autotexts = ax.pie(
            values,
            labels=labels,
            colors=colors,
            autopct='%1.1f%%',
            startangle=90,
        )

        for w in wedges:
            w.set_edgecolor('white')
            w.set_linewidth(1.5)

        ax.axis('equal')

    def _render_horizontal_bar_from_recommendation(self, ax, data_points) -> None:
        """Render horizontal bar chart from recommendation data points."""
        import numpy as np

        labels = [dp.subject for dp in data_points]
        values = [dp.normalized_value for dp in data_points]

        y = np.arange(len(labels))
        bars = ax.barh(y, values, color=CHART_COLORS[0])

        ax.set_yticks(y)
        ax.set_yticklabels(labels)
        ax.invert_yaxis()

        # Add value labels
        for bar, val in zip(bars, values):
            ax.text(
                bar.get_width() + max(values) * 0.01,
                bar.get_y() + bar.get_height() / 2,
                f'{val:,.1f}',
                va='center', fontsize=9
            )

    def _get_sections(self, session) -> Dict[str, Dict[str, Any]]:
        """Get section information from session."""
        return session.section_contents

    # ========================================================================
    # Image Extraction
    # ========================================================================

    def _extract_images_for_section(
        self,
        section_id: str,
        section_data: Dict[str, Any],
        evidence_list: List,
    ) -> List[Figure]:
        """Extract relevant images for a section."""
        figures = []

        # First, check if images are already stored in section_data
        existing_images = section_data.get("images", [])
        print(f"[FigureTableGenerator] Section {section_id}: found {len(existing_images)} stored images")
        for idx, img_data in enumerate(existing_images[:self.max_images_per_section]):
            src = img_data.get("src", "")
            if not src:
                continue

            page_url = img_data.get("page_url", "")
            image_path, image_data = self._download_image(src, referer=page_url or None)

            if image_path or image_data:
                figure = Figure(
                    figure_id=f"fig_{section_id}_{idx+1}",
                    figure_type=FigureType.IMAGE,
                    title=img_data.get("suggested_caption", "") or img_data.get("alt", "") or f"Figure {idx+1}",
                    caption=img_data.get("suggested_caption", "") or img_data.get("alt", ""),
                    source_url=src,
                    source_title=img_data.get("page_title", ""),
                    section_id=section_id,
                    image_path=image_path,
                    image_data=image_data,
                    alt_text=img_data.get("alt", ""),
                )
                figures.append(figure)
                print(f"[FigureTableGenerator] Downloaded image from stored data: {src[:60]}")

        # If not enough images from stored data, try evidence URLs
        if len(figures) < self.max_images_per_section and evidence_list:
            print(f"[FigureTableGenerator] Section {section_id}: trying {len(evidence_list)} evidence URLs for images")
            for idx, evidence in enumerate(evidence_list[:5]):  # Check more evidence URLs
                if len(figures) >= self.max_images_per_section:
                    break
                # Skip non-web evidence (PDFs etc.)
                url = getattr(evidence, 'url', '')
                if not url or not url.startswith('http'):
                    continue
                # Skip PDF/document URLs - they don't have extractable images
                if any(url.lower().endswith(ext) for ext in ['.pdf', '.docx', '.xlsx', '.csv']):
                    continue

                try:
                    images_from_url = self._extract_images_from_url(
                        url,
                        getattr(evidence, 'title', ''),
                        limit=2,
                    )
                    for img in images_from_url:
                        img.section_id = section_id
                        img.figure_id = f"fig_{section_id}_{len(figures)+1}"
                        figures.append(img)
                        print(f"[FigureTableGenerator] Extracted image from evidence URL: {url[:60]}")
                        if len(figures) >= self.max_images_per_section:
                            break
                except Exception as e:
                    print(f"[FigureTableGenerator] Failed to extract images from {url[:60]}: {e}")
                    continue

        print(f"[FigureTableGenerator] Section {section_id}: total {len(figures)} images extracted")
        return figures

    def _extract_images_from_url(
        self,
        url: str,
        page_title: str,
        limit: int = 3,
    ) -> List[Figure]:
        """Extract images from a web page URL."""
        figures = []

        try:
            from bs4 import BeautifulSoup

            response = requests.get(
                url,
                timeout=(5, 20),
                headers={
                    "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36",
                    "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8",
                    "Accept-Language": "en-US,en;q=0.9,ja;q=0.8",
                },
                proxies=self.proxies,
                verify=self.verify_ssl,
            )
            response.raise_for_status()

            soup = BeautifulSoup(response.content, 'html.parser')

            img_tags = soup.find_all('img')

            for img in img_tags:
                if len(figures) >= limit:
                    break

                src = img.get('src', '') or img.get('data-src', '')
                if not src:
                    continue

                # Don't urljoin data URIs (base64-encoded inline images)
                if not src.startswith('data:'):
                    src = urljoin(url, src)

                # Skip small icons and tracking pixels
                width = img.get('width')
                height = img.get('height')
                try:
                    if width and int(str(width).rstrip('px%')) < 100:
                        continue
                except (ValueError, TypeError):
                    pass
                try:
                    if height and int(str(height).rstrip('px%')) < 100:
                        continue
                except (ValueError, TypeError):
                    pass

                # Skip common non-content images
                skip_keywords = ['icon', 'logo', 'avatar', 'button', 'banner',
                                 'ad-', 'ads/', 'pixel', 'tracking', 'spacer']
                if any(skip in src.lower() for skip in skip_keywords):
                    continue

                image_path, image_data = self._download_image(src, referer=url)

                if image_path or image_data:
                    alt_text = img.get('alt', '')
                    title = img.get('title', '') or alt_text

                    figure = Figure(
                        figure_id="",
                        figure_type=FigureType.IMAGE,
                        title=title or f"Image from {page_title}",
                        caption=f"Source: {page_title}",
                        source_url=src,
                        source_title=page_title,
                        section_id="",
                        image_path=image_path,
                        image_data=image_data,
                        alt_text=alt_text,
                    )
                    figures.append(figure)

        except Exception as e:
            print(f"[FigureTableGenerator] Error extracting images from {url}: {e}")

        return figures

    def _download_image(
        self,
        url: str,
        referer: str = None,
    ) -> Tuple[Optional[Path], Optional[bytes]]:
        """Download an image from URL or decode a base64 data URI."""
        # Handle base64 data URIs (data:image/png;base64,...)
        if url.startswith('data:'):
            return self._decode_data_uri(url)

        import time
        from urllib.parse import urlparse

        headers = {
            "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36",
            "Accept": "image/avif,image/webp,image/apng,image/svg+xml,image/*,*/*;q=0.8",
            "Accept-Language": "en-US,en;q=0.9,ja;q=0.8",
        }
        if referer:
            headers["Referer"] = referer
        else:
            parsed = urlparse(url)
            headers["Referer"] = f"{parsed.scheme}://{parsed.netloc}/"

        # 2 attempts with a short (5, 10) timeout: figure images are
        # nice-to-have, and the previous 3 x (5,20)s + exponential backoff
        # could burn ~78s per dead URL, sequentially, across many images
        max_retries = 2
        last_exception = None
        for attempt in range(max_retries):
            try:
                response = requests.get(
                    url,
                    timeout=(5, 10),
                    headers=headers,
                    proxies=self.proxies,
                    verify=self.verify_ssl,
                )
                if response.status_code in (429, 503) and attempt < max_retries - 1:
                    wait = 2 ** attempt
                    print(f"[FigureTableGenerator] HTTP {response.status_code} for {url}, retrying in {wait}s...")
                    time.sleep(wait)
                    continue
                response.raise_for_status()
                break
            except requests.exceptions.ConnectionError as e:
                last_exception = e
                if attempt < max_retries - 1:
                    wait = 2 ** attempt
                    print(f"[FigureTableGenerator] Connection error for {url}, retrying in {wait}s...")
                    time.sleep(wait)
                    continue
                print(f"[FigureTableGenerator] Error downloading image {url}: {e}")
                return None, None
            except requests.exceptions.Timeout as e:
                last_exception = e
                if attempt < max_retries - 1:
                    wait = 2 ** attempt
                    print(f"[FigureTableGenerator] Timeout for {url}, retrying in {wait}s...")
                    time.sleep(wait)
                    continue
                print(f"[FigureTableGenerator] Error downloading image {url}: {e}")
                return None, None
            except requests.exceptions.HTTPError as e:
                print(f"[FigureTableGenerator] HTTP error downloading image {url}: {e}")
                return None, None
            except Exception as e:
                print(f"[FigureTableGenerator] Error downloading image {url}: {e}")
                return None, None
        else:
            print(f"[FigureTableGenerator] Failed to download image after {max_retries} retries: {url}: {last_exception}")
            return None, None

        try:

            content_type = response.headers.get('content-type', '').lower()

            # Check content-type: accept image types, octet-stream, and URLs with image extensions
            is_image_content = any(t in content_type for t in ['image', 'jpeg', 'png', 'gif', 'webp', 'svg'])
            is_octet_stream = 'octet-stream' in content_type
            has_image_extension = self._get_extension_from_url(url) is not None

            if not is_image_content and not (is_octet_stream and has_image_extension) and not has_image_extension:
                return None, None

            # Validate that the content is actually an image by checking magic bytes
            content = response.content
            if len(content) < 8:
                return None, None

            # Check common image file signatures (magic bytes)
            image_signatures = [
                b'\xff\xd8\xff',      # JPEG
                b'\x89PNG',            # PNG
                b'GIF87a', b'GIF89a',  # GIF
                b'RIFF',              # WEBP (RIFF....WEBP)
                b'<svg',              # SVG
                b'<?xml',             # SVG (XML format)
            ]
            is_valid_image = any(content[:len(sig)] == sig for sig in image_signatures)
            # WEBP specific check
            if content[:4] == b'RIFF' and len(content) > 11 and content[8:12] == b'WEBP':
                is_valid_image = True

            if not is_valid_image and not is_image_content:
                return None, None

            url_hash = hashlib.md5(url.encode()).hexdigest()[:12]
            ext = self._get_extension_from_url(url) or '.png'
            filename = f"img_{url_hash}{ext}"
            filepath = self.output_dir / filename

            with open(filepath, 'wb') as f:
                f.write(content)

            return filepath, content

        except Exception as e:
            print(f"[FigureTableGenerator] Error downloading image {url}: {e}")
            return None, None

    def _decode_data_uri(
        self,
        data_uri: str,
    ) -> Tuple[Optional[Path], Optional[bytes]]:
        """Decode a base64 data URI and save as image file.

        Handles URIs like: data:image/png;base64,iVBORw0KGgo...
        """
        try:
            # Parse: data:[<mediatype>][;base64],<data>
            if ',' not in data_uri:
                return None, None

            header, encoded = data_uri.split(',', 1)

            # Determine extension from MIME type
            ext = '.png'  # default
            mime_to_ext = {
                'image/png': '.png',
                'image/jpeg': '.jpg',
                'image/gif': '.gif',
                'image/webp': '.webp',
                'image/svg+xml': '.svg',
            }
            for mime, extension in mime_to_ext.items():
                if mime in header:
                    ext = extension
                    break

            if ';base64' not in header:
                return None, None

            content = base64.b64decode(encoded)
            if len(content) < 8:
                return None, None

            url_hash = hashlib.md5(encoded[:64].encode()).hexdigest()[:12]
            filename = f"img_{url_hash}{ext}"
            filepath = self.output_dir / filename

            with open(filepath, 'wb') as f:
                f.write(content)

            return filepath, content

        except Exception as e:
            print(f"[FigureTableGenerator] Error decoding data URI: {e}")
            return None, None

    def _get_extension_from_url(self, url: str) -> Optional[str]:
        """Get file extension from URL."""
        parsed = urlparse(url)
        path = parsed.path.lower()
        for ext in self.IMAGE_EXTENSIONS:
            if path.endswith(ext):
                return ext
        return None

    # ========================================================================
    # Table Extraction
    # ========================================================================

    def _extract_tables_for_section(
        self,
        section_id: str,
        section_data: Dict[str, Any],
        content: str,
        evidence_list: List,
    ) -> List[TableData]:
        """Extract numerical data and create tables for a section.

        Extraction priority (evidence-first to reduce hallucination):
          1. HTML table parsing from evidence URLs (most reliable)
          2. LLM-based extraction from evidence content (evidence-grounded)
          3. Pattern-based extraction from evidence/content (fallback)
        """
        tables = []
        section_title = section_data.get("title", "")

        # Build evidence text once for reuse
        evidence_text = self._build_evidence_text(evidence_list)

        # Method 1 (Priority): HTML table parsing from evidence sources
        if evidence_list:
            html_tables = self._extract_html_tables(
                section_id=section_id,
                evidence_list=evidence_list,
                section_title=section_title,
            )
            tables.extend(html_tables)
            if html_tables:
                print(f"[FigureTableGenerator] Section {section_id}: "
                      f"extracted {len(html_tables)} table(s) from HTML sources")

        # Method 2: LLM-based extraction from evidence content (not LLM prose)
        if not tables and self.llm:
            # Prefer evidence text; fall back to section content only if no evidence
            source_text = evidence_text or content
            if source_text:
                extracted_tables = self._extract_tables_with_llm(
                    section_id=section_id,
                    content=source_text,
                    section_title=section_title,
                )
                tables.extend(extracted_tables)
                if extracted_tables:
                    print(f"[FigureTableGenerator] Section {section_id}: "
                          f"extracted {len(extracted_tables)} table(s) via LLM from evidence")

        # Method 3: Pattern-based extraction (fallback)
        if not tables:
            # Use evidence text if available, otherwise fall back to section content
            pattern_source = evidence_text or content
            extracted_tables = self._extract_tables_by_pattern(
                section_id=section_id,
                content=pattern_source,
                section_title=section_title,
            )
            tables.extend(extracted_tables)
            if extracted_tables:
                print(f"[FigureTableGenerator] Section {section_id}: "
                      f"extracted {len(extracted_tables)} table(s) via pattern matching")

        if not tables:
            print(f"[FigureTableGenerator] Section {section_id}: "
                  f"no tables extracted (evidence_text={len(evidence_text)} chars, "
                  f"content={len(content)} chars)")

        # Add source attribution
        if tables and evidence_list:
            primary_source = evidence_list[0] if evidence_list else None
            for table in tables:
                if primary_source and not table.source_url:
                    table.source_url = primary_source.url
                    table.source_title = primary_source.title

        return tables

    def _build_evidence_text(self, evidence_list: List) -> str:
        """Build combined text from evidence content for table extraction.

        Uses the original source content (content_excerpt / original_content)
        rather than LLM-generated prose to reduce hallucination risk.
        """
        parts = []
        for evidence in evidence_list[:5]:  # Limit to avoid token overflow
            text = (
                getattr(evidence, 'content_excerpt', '')
                or getattr(evidence, 'original_content', '')
                or ''
            )
            if text.strip():
                parts.append(text.strip())
        return "\n\n".join(parts)

    def _extract_tables_with_llm(
        self,
        section_id: str,
        content: str,
        section_title: str,
    ) -> List[TableData]:
        """Use LLM to extract tabular data from content."""
        tables = []

        # Use more content (up to 8000 chars) for better extraction
        truncated = content[:8000]

        if self.language == "ja":
            prompt = f"""以下のテキストから、表形式で表現すべき数値データを抽出してください。

対象データ:
- 時系列データ（年次、月次推移）
- 項目間の比較データ
- 構成比・シェアデータ（合計100%になるもの）
- 統計データ・ランキング
- 相関関係のあるデータ（2つの数値指標の関係）

セクションタイトル: {section_title}

テキスト:
{truncated}

以下のJSON形式で返答してください。データがない場合は空配列[]を返してください。
```json
[
  {{
    "title": "表のタイトル",
    "headers": ["列1", "列2", "列3"],
    "rows": [["データ1", "データ2", "データ3"]],
    "data_type": "time_series|comparison|distribution|correlation|statistics",
    "unit": "単位（億円、%、人など）"
  }}
]
```

data_typeの分類基準:
- time_series: 時間軸に沿ったデータ（年、月、四半期）
- comparison: 複数項目の並列比較
- distribution: 構成比・割合（合計100%前後）
- correlation: 2つの数値指標の関係性
- statistics: 集計値、平均値、中央値など

JSONのみを返してください。"""
        else:
            prompt = f"""Analyze the following text and extract numerical data suitable for table format.

Look for:
- Time series data (yearly, monthly trends)
- Comparisons between items
- Distribution/share data (totaling ~100%)
- Statistics, rankings, percentages
- Correlated data (relationship between two metrics)

Section Title: {section_title}

Content:
{truncated}

Return a JSON array of tables:
```json
[
  {{
    "title": "Table title",
    "headers": ["Col1", "Col2", "Col3"],
    "rows": [["data1", "data2", "data3"]],
    "data_type": "time_series|comparison|distribution|correlation|statistics",
    "unit": "unit of measurement"
  }}
]
```

data_type classification:
- time_series: Data along time axis (year, month, quarter)
- comparison: Side-by-side comparison of items
- distribution: Composition/share (totals ~100%)
- correlation: Relationship between two numeric metrics
- statistics: Aggregated values, averages, medians

Return ONLY valid JSON, no other text."""

        try:
            response = self.llm.generate(prompt)
            response_text = response.content

            extracted = self._parse_json_array(response_text)

            for idx, table_data in enumerate(extracted):
                if not table_data.get("headers") or not table_data.get("rows"):
                    continue
                # Skip tables with only 1 row of data (not useful)
                if len(table_data["rows"]) < 2:
                    continue

                table = TableData(
                    table_id=f"table_{section_id}_{idx+1}",
                    title=table_data.get("title", f"Table {idx+1}"),
                    caption=table_data.get("title", ""),
                    headers=table_data["headers"],
                    rows=table_data["rows"],
                    section_id=section_id,
                    data_type=table_data.get("data_type", "general"),
                    unit=table_data.get("unit"),
                )
                tables.append(table)

        except Exception as e:
            print(f"[FigureTableGenerator] Error extracting tables with LLM: {e}")

        return tables

    def _extract_html_tables(
        self,
        section_id: str,
        evidence_list: List,
        section_title: str,
        max_tables: int = 2,
    ) -> List[TableData]:
        """Extract tables from HTML pages referenced in evidence."""
        tables = []

        try:
            from bs4 import BeautifulSoup
        except ImportError:
            return tables

        for evidence in evidence_list[:3]:
            if len(tables) >= max_tables:
                break

            try:
                response = requests.get(
                    evidence.url,
                    timeout=self.download_timeout,
                    headers={"User-Agent": "Mozilla/5.0 (compatible; ResearchBot/1.0)"},
                    proxies=self.proxies,
                    verify=self.verify_ssl,
                )
                response.raise_for_status()
                soup = BeautifulSoup(response.content, 'html.parser')

                html_tables = soup.find_all('table')

                for t_idx, html_table in enumerate(html_tables):
                    if len(tables) >= max_tables:
                        break

                    headers, rows = self._parse_html_table(html_table)
                    if not headers or len(rows) < 2:
                        continue
                    # Skip navigation or layout tables (too few columns)
                    if len(headers) < 2:
                        continue

                    table = TableData(
                        table_id=f"table_{section_id}_html_{len(tables)+1}",
                        title=self._guess_table_title(html_table, section_title),
                        caption=f"Source: {evidence.title}",
                        headers=headers,
                        rows=rows[:20],  # Limit rows
                        source_url=evidence.url,
                        source_title=evidence.title,
                        section_id=section_id,
                        data_type=self._guess_data_type(headers, rows),
                    )
                    tables.append(table)

            except Exception as e:
                print(f"[FigureTableGenerator] Error parsing HTML tables from {evidence.url}: {e}")

        return tables

    def _parse_html_table(self, table_tag) -> Tuple[List[str], List[List[str]]]:
        """Parse an HTML table tag into headers and rows."""
        headers = []
        rows = []

        # Extract headers
        thead = table_tag.find('thead')
        if thead:
            th_tags = thead.find_all('th')
            headers = [th.get_text(strip=True) for th in th_tags]

        # If no thead, try first row
        if not headers:
            first_row = table_tag.find('tr')
            if first_row:
                th_tags = first_row.find_all('th')
                if th_tags:
                    headers = [th.get_text(strip=True) for th in th_tags]

        # Extract rows
        tbody = table_tag.find('tbody') or table_tag
        for tr in tbody.find_all('tr'):
            cells = tr.find_all(['td', 'th'])
            if cells:
                row = [cell.get_text(strip=True) for cell in cells]
                # Skip if this is the header row
                if row == headers:
                    continue
                if any(cell.strip() for cell in row):
                    rows.append(row)

        # If no headers found, use first data row as headers
        if not headers and rows:
            headers = rows.pop(0)

        return headers, rows

    def _guess_table_title(self, table_tag, fallback_title: str) -> str:
        """Guess table title from surrounding HTML elements."""
        # Check for caption tag
        caption = table_tag.find('caption')
        if caption:
            return caption.get_text(strip=True)

        # Check preceding sibling for heading
        prev = table_tag.find_previous_sibling(['h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'p', 'strong'])
        if prev:
            text = prev.get_text(strip=True)
            if len(text) < 100:
                return text

        return f"{fallback_title} - Data"

    def _guess_data_type(self, headers: List[str], rows: List[List[str]]) -> str:
        """Guess the data type from headers and content."""
        headers_lower = [h.lower() for h in headers]

        # Time series: has year/date column
        time_keywords = ['year', 'date', '年', '月', '四半期', 'quarter', 'period', '年度']
        if any(any(kw in h for kw in time_keywords) for h in headers_lower):
            return "time_series"

        # Distribution: has percentage/share column
        dist_keywords = ['%', 'share', 'シェア', '割合', '構成比', 'ratio', 'percentage']
        if any(any(kw in h for kw in dist_keywords) for h in headers_lower):
            return "distribution"

        # Check if data contains mostly percentages
        pct_count = 0
        total_cells = 0
        for row in rows[:5]:
            for cell in row[1:]:
                total_cells += 1
                if isinstance(cell, str) and ('%' in cell or '％' in cell):
                    pct_count += 1
        if total_cells > 0 and pct_count / total_cells > 0.5:
            return "distribution"

        return "comparison"

    def _extract_tables_by_pattern(
        self,
        section_id: str,
        content: str,
        section_title: str,
    ) -> List[TableData]:
        """Extract tabular data using pattern matching."""
        tables = []

        # Look for year-based data patterns
        years = self.YEAR_PATTERN.findall(content)
        if years and len(set(years)) >= 2:
            time_series = self._extract_time_series(content, years)
            if time_series:
                table = TableData(
                    table_id=f"table_{section_id}_ts",
                    title=f"{section_title} - Time Series Data",
                    caption="Extracted time series data",
                    headers=["Year"] + list(time_series.keys()),
                    rows=self._format_time_series_rows(time_series),
                    section_id=section_id,
                    data_type="time_series",
                )
                tables.append(table)

        # Look for percentage comparisons
        percentages = self.PERCENTAGE_PATTERN.findall(content)
        if not tables and percentages and len(percentages) >= 3:
            pct_table = self._extract_percentage_data(content, section_id, section_title)
            if pct_table:
                tables.append(pct_table)

        return tables

    def _extract_time_series(
        self,
        content: str,
        years: List[str],
    ) -> Dict[str, Dict[str, Any]]:
        """Extract time series data from content."""
        time_series = {}

        for year in set(years):
            pattern = rf'{year}[年\s:：]?\s*([\d,]+\.?\d*)\s*(億|万|%|円|ドル)?'
            matches = re.findall(pattern, content)
            if matches:
                for value, unit in matches:
                    metric = unit or "Value"
                    if metric not in time_series:
                        time_series[metric] = {}
                    try:
                        time_series[metric][year] = float(value.replace(',', ''))
                    except ValueError:
                        pass

        return time_series

    def _extract_percentage_data(
        self,
        content: str,
        section_id: str,
        section_title: str,
    ) -> Optional[TableData]:
        """Extract percentage/distribution data from content."""
        # Find patterns like "ItemA: 30%" or "ItemA（30%）" or "ItemA 30%"
        patterns = [
            re.compile(r'([^\d\n]{2,20})[：:]\s*(\d+\.?\d*)\s*[%％]'),
            re.compile(r'([^\d\n]{2,20})[\(（]\s*(\d+\.?\d*)\s*[%％][\)）]'),
            re.compile(r'「([^」]+)」\s*(\d+\.?\d*)\s*[%％]'),
        ]

        items = []
        for pattern in patterns:
            matches = pattern.findall(content)
            for name, value in matches:
                name = name.strip().strip('・-– ')
                try:
                    items.append((name, float(value)))
                except ValueError:
                    pass

        if len(items) < 3:
            return None

        # Deduplicate
        seen = set()
        unique_items = []
        for name, value in items:
            if name not in seen:
                seen.add(name)
                unique_items.append((name, value))

        if len(unique_items) < 3:
            return None

        headers = ["Item", "Percentage (%)"]
        rows = [[name, f"{value:.1f}"] for name, value in unique_items[:10]]

        return TableData(
            table_id=f"table_{section_id}_pct",
            title=f"{section_title} - Distribution",
            caption="Extracted distribution data",
            headers=headers,
            rows=rows,
            section_id=section_id,
            data_type="distribution",
            unit="%",
        )

    def _format_time_series_rows(
        self,
        time_series: Dict[str, Dict[str, Any]],
    ) -> List[List[Any]]:
        """Format time series data into table rows."""
        if not time_series:
            return []

        all_years = set()
        for metric_data in time_series.values():
            all_years.update(metric_data.keys())

        years = sorted(all_years)
        rows = []

        for year in years:
            row = [year]
            for metric in time_series.keys():
                value = time_series[metric].get(year, "")
                row.append(value)
            rows.append(row)

        return rows

    def _parse_json_array(self, text: str) -> List[Dict]:
        """Robustly parse a JSON array from LLM response text."""
        # Try direct parse
        text = text.strip()
        try:
            result = json.loads(text)
            if isinstance(result, list):
                return result
        except json.JSONDecodeError:
            pass

        # Extract from code block
        code_block = re.search(r'```(?:json)?\s*\n?(.*?)\n?```', text, re.DOTALL)
        if code_block:
            try:
                result = json.loads(code_block.group(1).strip())
                if isinstance(result, list):
                    return result
            except json.JSONDecodeError:
                pass

        # Find outermost brackets
        start = text.find('[')
        end = text.rfind(']')
        if start != -1 and end > start:
            try:
                result = json.loads(text[start:end + 1])
                if isinstance(result, list):
                    return result
            except json.JSONDecodeError:
                pass

        return []

    # ========================================================================
    # LLM Visualization Recommendation
    # ========================================================================

    def _recommend_visualization(self, table: TableData) -> ChartType:
        """
        Use LLM to recommend the best chart type for a table.

        Falls back to rule-based selection if LLM is unavailable.
        """
        if not self.llm:
            return self._rule_based_chart_selection(table)

        # Prepare compact data summary
        data_summary = f"Title: {table.title}\n"
        data_summary += f"Headers: {table.headers}\n"
        data_summary += f"Row count: {len(table.rows)}\n"
        data_summary += f"Sample rows (first 5):\n"
        for row in table.rows[:5]:
            data_summary += f"  {row}\n"
        data_summary += f"Data type hint: {table.data_type}\n"
        data_summary += f"Unit: {table.unit or 'N/A'}"

        prompt = f"""Analyze the following data and recommend the SINGLE best chart type for visualization.

{data_summary}

Available chart types:
- line: Time series, trends over time, continuous data
- bar: Categorical comparison, discrete items side by side
- pie: Composition/share of a whole (only when ≤8 categories, data sums to ~100%)
- area: Cumulative trends, stacked time series
- scatter: Correlation between two numeric variables
- stacked_bar: Part-to-whole comparison across categories
- horizontal_bar: Comparison with long category labels or many items (>6)

Return ONLY the chart type name (one word), nothing else.
Example: line"""

        try:
            response = self.llm.generate(prompt)
            recommended = response.content.strip().lower().replace('"', '').replace("'", "")

            # Map to ChartType
            type_map = {
                "line": ChartType.LINE,
                "bar": ChartType.BAR,
                "pie": ChartType.PIE,
                "area": ChartType.AREA,
                "scatter": ChartType.SCATTER,
                "stacked_bar": ChartType.STACKED_BAR,
                "horizontal_bar": ChartType.HORIZONTAL_BAR,
            }
            if recommended in type_map:
                return type_map[recommended]

        except Exception as e:
            print(f"[FigureTableGenerator] LLM chart recommendation failed: {e}")

        return self._rule_based_chart_selection(table)

    def _rule_based_chart_selection(self, table: TableData) -> ChartType:
        """Rule-based chart type selection based on data characteristics."""
        num_rows = len(table.rows)
        num_cols = len(table.headers)

        if table.data_type == "time_series":
            if num_cols > 3:
                return ChartType.AREA
            return ChartType.LINE

        if table.data_type == "distribution":
            if num_rows <= 8:
                return ChartType.PIE
            return ChartType.HORIZONTAL_BAR

        if table.data_type == "correlation":
            return ChartType.SCATTER

        if table.data_type == "comparison":
            if num_cols > 3 and num_rows > 3:
                return ChartType.STACKED_BAR
            if num_rows > 6:
                return ChartType.HORIZONTAL_BAR
            return ChartType.BAR

        # Default: BAR for small data, HORIZONTAL_BAR for many items
        if num_rows > 6:
            return ChartType.HORIZONTAL_BAR
        return ChartType.BAR

    # ========================================================================
    # Chart Generation
    # ========================================================================

    def _generate_chart_for_table(
        self,
        table: TableData,
    ) -> Optional[Figure]:
        """Generate a chart for a table using the recommended chart type."""
        try:
            import matplotlib
            matplotlib.use('Agg')
            import matplotlib.pyplot as plt
            import matplotlib.font_manager as fm

            # Set up fonts and style (matplotlib or seaborn theme)
            self._apply_chart_style(plt, fm)

            # Get LLM-recommended or rule-based chart type
            chart_type = self._recommend_visualization(table)

            # Prepare data
            if len(table.headers) < 2 or len(table.rows) < 2:
                return None

            x_labels = [str(row[0]) for row in table.rows]
            data_columns = self._extract_numeric_columns(table)

            if not data_columns:
                return None

            # Create figure
            fig, ax = plt.subplots(figsize=(10, 6))

            # Render chart
            success = self._render_chart(ax, chart_type, x_labels, data_columns, table)

            if not success:
                plt.close(fig)
                return None

            # Set title and labels
            ax.set_title(table.title, fontsize=14, fontweight='bold', pad=15)
            if table.unit:
                ax.set_ylabel(table.unit, fontsize=11)

            plt.tight_layout()

            # Save chart
            chart_filename = f"chart_{table.table_id}.png"
            chart_path = self.output_dir / chart_filename
            plt.savefig(chart_path, dpi=150, bbox_inches='tight',
                        facecolor='white', edgecolor='none')
            plt.close(fig)

            with open(chart_path, 'rb') as f:
                image_data = f.read()

            chart_type_label = chart_type.value.replace("_", " ").title()
            chart_figure = Figure(
                figure_id=f"chart_{table.table_id}",
                figure_type=FigureType.CHART,
                title=f"{table.title}",
                caption=f"{chart_type_label} chart: {table.caption}" if table.caption else table.title,
                source_url=table.source_url,
                source_title=table.source_title,
                section_id=table.section_id,
                image_path=chart_path,
                image_data=image_data,
                alt_text=f"Chart showing {table.title}",
            )

            return chart_figure

        except ImportError:
            print("[FigureTableGenerator] matplotlib not installed. Install with: pip install matplotlib")
            return None
        except Exception as e:
            print(f"[FigureTableGenerator] Error generating chart: {e}")
            return None

    def _get_seaborn(self):
        """Import seaborn if selected and available; warn once on failure."""
        if self.chart_library != "seaborn":
            return None
        try:
            import seaborn as sns
            return sns
        except ImportError:
            if not self._seaborn_warned:
                print("[FigureTableGenerator] seaborn not installed; "
                      "falling back to matplotlib. Install with: pip install seaborn")
                self._seaborn_warned = True
            return None

    def _apply_chart_style(self, plt, fm):
        """
        Apply the configured chart style (matplotlib or seaborn theme).

        Seaborn's set_theme resets rcParams, so Japanese fonts are
        re-applied afterwards in both branches.
        """
        sns = self._get_seaborn()
        if sns is not None:
            sns.set_theme(style="whitegrid", palette="deep")
            self._setup_matplotlib_fonts(plt, fm)
            plt.rcParams.update({
                "axes.spines.top": False,
                "axes.spines.right": False,
                "figure.facecolor": "white",
            })
        else:
            self._setup_matplotlib_fonts(plt, fm)
            self._setup_chart_style(plt)
        return sns

    def _setup_matplotlib_fonts(self, plt, fm):
        """Configure matplotlib for Japanese font support."""
        try:
            japanese_fonts = ['IPAGothic', 'IPAexGothic', 'Noto Sans CJK JP',
                              'Yu Gothic', 'Hiragino Sans', 'MS Gothic']
            available_fonts = [f.name for f in fm.fontManager.ttflist]
            for font in japanese_fonts:
                if font in available_fonts:
                    plt.rcParams['font.family'] = font
                    break
        except Exception:
            pass

    def _setup_chart_style(self, plt):
        """Set professional chart styling."""
        plt.rcParams.update({
            'axes.grid': True,
            'grid.alpha': 0.3,
            'grid.linestyle': '--',
            'axes.spines.top': False,
            'axes.spines.right': False,
            'axes.labelsize': 11,
            'xtick.labelsize': 10,
            'ytick.labelsize': 10,
            'legend.fontsize': 10,
            'figure.facecolor': 'white',
        })

    def _extract_numeric_columns(
        self, table: TableData,
    ) -> List[Tuple[str, List[float]]]:
        """Extract numeric columns from table data."""
        data_columns = []

        for col_idx in range(1, len(table.headers)):
            col_data = []
            for row in table.rows:
                try:
                    if col_idx < len(row):
                        val = row[col_idx]
                        if isinstance(val, (int, float)):
                            col_data.append(float(val))
                        elif isinstance(val, str):
                            clean_val = val.replace(',', '').replace('億', '').replace('万', '')
                            clean_val = clean_val.replace('%', '').replace('％', '').strip()
                            col_data.append(float(clean_val))
                        else:
                            col_data.append(0)
                    else:
                        col_data.append(0)
                except (ValueError, TypeError):
                    col_data.append(0)
            data_columns.append((table.headers[col_idx], col_data))

        # Filter out all-zero columns
        data_columns = [(label, data) for label, data in data_columns
                        if any(v != 0 for v in data)]

        return data_columns

    def _render_chart(
        self,
        ax,
        chart_type: ChartType,
        x_labels: List[str],
        data_columns: List[Tuple[str, List[float]]],
        table: TableData,
    ) -> bool:
        """Render a specific chart type. Returns True on success."""
        try:
            if chart_type == ChartType.LINE:
                return self._render_line_chart(ax, x_labels, data_columns)
            elif chart_type == ChartType.BAR:
                return self._render_bar_chart(ax, x_labels, data_columns)
            elif chart_type == ChartType.PIE:
                return self._render_pie_chart(ax, x_labels, data_columns)
            elif chart_type == ChartType.AREA:
                return self._render_area_chart(ax, x_labels, data_columns)
            elif chart_type == ChartType.SCATTER:
                return self._render_scatter_chart(ax, x_labels, data_columns)
            elif chart_type == ChartType.STACKED_BAR:
                return self._render_stacked_bar_chart(ax, x_labels, data_columns)
            elif chart_type == ChartType.HORIZONTAL_BAR:
                return self._render_horizontal_bar_chart(ax, x_labels, data_columns)
            else:
                return self._render_bar_chart(ax, x_labels, data_columns)
        except Exception as e:
            print(f"[FigureTableGenerator] Failed to render {chart_type.value}: {e}")
            # Fallback to bar chart
            try:
                return self._render_bar_chart(ax, x_labels, data_columns)
            except Exception:
                return False

    def _render_line_chart(self, ax, x_labels, data_columns) -> bool:
        """Render a line chart."""
        for i, (label, data) in enumerate(data_columns):
            color = CHART_COLORS[i % len(CHART_COLORS)]
            ax.plot(x_labels, data, marker='o', label=label, color=color,
                    linewidth=2, markersize=6)
        if len(data_columns) > 1:
            ax.legend(loc='best')
        ax.tick_params(axis='x', rotation=45)
        return True

    def _render_bar_chart(self, ax, x_labels, data_columns) -> bool:
        """Render a grouped bar chart."""
        import numpy as np

        x = np.arange(len(x_labels))
        width = 0.8 / max(len(data_columns), 1)

        for i, (label, data) in enumerate(data_columns):
            color = CHART_COLORS[i % len(CHART_COLORS)]
            offset = (i - len(data_columns) / 2 + 0.5) * width
            bars = ax.bar(x + offset, data, width, label=label, color=color,
                          edgecolor='white', linewidth=0.5)
            # Add value labels on bars for small datasets
            if len(x_labels) <= 6 and len(data_columns) <= 3:
                for bar, val in zip(bars, data):
                    if val != 0:
                        ax.text(bar.get_x() + bar.get_width() / 2, bar.get_height(),
                                f'{val:,.0f}', ha='center', va='bottom', fontsize=8)

        ax.set_xticks(x)
        ax.set_xticklabels(x_labels, rotation=45, ha='right')
        if len(data_columns) > 1:
            ax.legend(loc='best')
        return True

    def _render_pie_chart(self, ax, x_labels, data_columns) -> bool:
        """Render a pie chart."""
        if not data_columns:
            return False

        # Use first data column
        label_name, values = data_columns[0]
        colors = PIE_COLORS[:len(values)]

        # Filter out zero/negative values
        filtered = [(label, val) for label, val in zip(x_labels, values) if val > 0]
        if len(filtered) < 2:
            return False

        labels, vals = zip(*filtered)

        # Calculate percentages for autopct
        total = sum(vals)
        wedges, texts, autotexts = ax.pie(
            vals,
            labels=labels,
            colors=colors[:len(vals)],
            autopct=lambda pct: f'{pct:.1f}%' if pct >= 3 else '',
            startangle=90,
            pctdistance=0.75,
            wedgeprops={'edgecolor': 'white', 'linewidth': 1.5},
        )

        # Style text
        for text in texts:
            text.set_fontsize(9)
        for autotext in autotexts:
            autotext.set_fontsize(8)
            autotext.set_fontweight('bold')

        ax.set_aspect('equal')
        ax.grid(False)
        ax.spines['left'].set_visible(False)
        ax.spines['bottom'].set_visible(False)
        ax.tick_params(left=False, bottom=False, labelleft=False, labelbottom=False)
        return True

    def _render_area_chart(self, ax, x_labels, data_columns) -> bool:
        """Render a stacked area chart."""
        import numpy as np

        x = np.arange(len(x_labels))

        # Stack data
        bottom = np.zeros(len(x_labels))
        for i, (label, data) in enumerate(data_columns):
            color = CHART_COLORS[i % len(CHART_COLORS)]
            data_arr = np.array(data)
            ax.fill_between(x, bottom, bottom + data_arr, label=label,
                            color=color, alpha=0.7)
            ax.plot(x, bottom + data_arr, color=color, linewidth=0.8)
            bottom += data_arr

        ax.set_xticks(x)
        ax.set_xticklabels(x_labels, rotation=45, ha='right')
        if len(data_columns) > 1:
            ax.legend(loc='upper left')
        ax.set_xlim(0, len(x_labels) - 1)
        return True

    def _render_scatter_chart(self, ax, x_labels, data_columns) -> bool:
        """Render a scatter chart."""
        if len(data_columns) < 2:
            # Use x_labels as numeric if possible, otherwise use index
            try:
                x_vals = [float(str(l).replace(',', '')) for l in x_labels]
            except ValueError:
                x_vals = list(range(len(x_labels)))
            y_label, y_vals = data_columns[0]

            color = CHART_COLORS[0]
            ax.scatter(x_vals, y_vals, color=color, s=60, alpha=0.7, edgecolors='white')
            ax.set_xlabel(x_labels[0] if x_labels else "X")
            ax.set_ylabel(y_label)
        else:
            # Two data columns: use first as X, second as Y
            x_label, x_vals = data_columns[0]
            y_label, y_vals = data_columns[1]

            color = CHART_COLORS[0]
            ax.scatter(x_vals, y_vals, color=color, s=60, alpha=0.7, edgecolors='white')
            ax.set_xlabel(x_label)
            ax.set_ylabel(y_label)

            # Add labels for each point
            if len(x_labels) <= 15:
                for i, label in enumerate(x_labels):
                    if i < len(x_vals) and i < len(y_vals):
                        ax.annotate(label, (x_vals[i], y_vals[i]),
                                    textcoords="offset points", xytext=(5, 5),
                                    fontsize=7, alpha=0.8)

        return True

    def _render_stacked_bar_chart(self, ax, x_labels, data_columns) -> bool:
        """Render a stacked bar chart."""
        import numpy as np

        x = np.arange(len(x_labels))
        width = 0.6

        bottom = np.zeros(len(x_labels))
        for i, (label, data) in enumerate(data_columns):
            color = CHART_COLORS[i % len(CHART_COLORS)]
            data_arr = np.array(data)
            ax.bar(x, data_arr, width, label=label, color=color,
                   bottom=bottom, edgecolor='white', linewidth=0.5)
            bottom += data_arr

        ax.set_xticks(x)
        ax.set_xticklabels(x_labels, rotation=45, ha='right')
        ax.legend(loc='upper left', bbox_to_anchor=(1.02, 1))
        return True

    def _render_horizontal_bar_chart(self, ax, x_labels, data_columns) -> bool:
        """Render a horizontal bar chart."""
        import numpy as np

        if not data_columns:
            return False

        y = np.arange(len(x_labels))

        if len(data_columns) == 1:
            label, data = data_columns[0]
            color = CHART_COLORS[0]
            bars = ax.barh(y, data, color=color, edgecolor='white', linewidth=0.5)
            # Add value labels
            for bar, val in zip(bars, data):
                if val != 0:
                    ax.text(bar.get_width() + max(data) * 0.02, bar.get_y() + bar.get_height() / 2,
                            f'{val:,.0f}', ha='left', va='center', fontsize=9)
        else:
            height = 0.8 / max(len(data_columns), 1)
            for i, (label, data) in enumerate(data_columns):
                color = CHART_COLORS[i % len(CHART_COLORS)]
                offset = (i - len(data_columns) / 2 + 0.5) * height
                ax.barh(y + offset, data, height, label=label, color=color,
                        edgecolor='white', linewidth=0.5)
            ax.legend(loc='best')

        ax.set_yticks(y)
        ax.set_yticklabels(x_labels)
        ax.invert_yaxis()
        return True

    # ========================================================================
    # Markdown Integration
    # ========================================================================

    def _insert_section_figures(
        self,
        result_lines: list,
        section_id: str,
        collection: 'FigureTableCollection',
        inserted_sections: set,
    ) -> None:
        """Insert figures, tables, and charts for a section into result_lines."""
        if section_id in inserted_sections:
            return
        inserted_sections.add(section_id)

        # Add figures for this section
        section_figures = collection.get_figures_for_section(section_id)
        for fig in section_figures:
            result_lines.append('')
            if fig.image_path:
                image_path = Path(fig.image_path) if not isinstance(fig.image_path, Path) else fig.image_path
                if image_path.exists():
                    result_lines.append(f'![{fig.title}]({fig.image_path})')
                else:
                    result_lines.append(f'<!-- Image not found: {fig.image_path} -->')
            result_lines.append(f'*{fig.caption}*')
            if fig.source_url:
                result_lines.append(f'Source: [{fig.source_title or fig.source_url}]({fig.source_url})')
            result_lines.append('')

        # Add tables for this section
        section_tables = collection.get_tables_for_section(section_id)
        for table in section_tables:
            result_lines.append('')
            result_lines.append(f'**{table.title}**')
            result_lines.append('')

            header_line = '| ' + ' | '.join(str(h) for h in table.headers) + ' |'
            separator = '|' + '|'.join(['---'] * len(table.headers)) + '|'
            result_lines.append(header_line)
            result_lines.append(separator)

            for row in table.rows:
                row_line = '| ' + ' | '.join(str(cell) for cell in row) + ' |'
                result_lines.append(row_line)

            result_lines.append('')
            # Caption (distinct from title)
            if table.caption and table.caption != table.title:
                result_lines.append(f'*{table.caption}*')
            if table.source_url:
                result_lines.append(f'*Source: [{table.source_title or "Link"}]({table.source_url})*')
            result_lines.append('')

        # Add charts for this section
        section_charts = collection.get_charts_for_section(section_id)
        for chart in section_charts:
            result_lines.append('')
            if chart.image_path:
                image_path = Path(chart.image_path) if not isinstance(chart.image_path, Path) else chart.image_path
                if image_path.exists():
                    result_lines.append(f'![{chart.title}]({chart.image_path})')
                else:
                    result_lines.append(f'<!-- Chart not found: {chart.image_path} -->')
            result_lines.append(f'*{chart.caption}*')
            result_lines.append('')

    def add_figures_to_markdown(
        self,
        markdown_content: str,
        collection: FigureTableCollection,
    ) -> str:
        """Add figures and tables to markdown content."""
        lines = markdown_content.split('\n')
        result_lines = []
        current_section = None
        inserted_sections = set()

        for line in lines:
            # When a new section header is found, flush figures for previous section first
            if line.startswith('## ') or line.startswith('### '):
                match = re.match(r'^##+ (\d+(?:\.\d+)?)[.\s]', line)
                if match:
                    # Flush previous section before starting new one
                    if current_section:
                        self._insert_section_figures(result_lines, current_section, collection, inserted_sections)
                    current_section = match.group(1)

            result_lines.append(line)

            # After section content, add figures and tables
            if current_section and (line.strip() == '' or line.startswith('---')):
                self._insert_section_figures(result_lines, current_section, collection, inserted_sections)
                current_section = None

        # Flush figures for the last section (if document ends without blank line)
        if current_section:
            self._insert_section_figures(result_lines, current_section, collection, inserted_sections)

        # Append orphaned figures (section_id="" or unmatched) at the end
        orphan_figures = [f for f in collection.figures if f.section_id not in inserted_sections]
        orphan_charts = [c for c in collection.charts if c.section_id not in inserted_sections]
        orphan_tables = [t for t in collection.tables if t.section_id not in inserted_sections]

        if orphan_figures or orphan_charts or orphan_tables:
            result_lines.append('')
            result_lines.append('---')
            result_lines.append('')
            for fig in orphan_figures:
                if fig.image_path:
                    image_path = Path(fig.image_path) if not isinstance(fig.image_path, Path) else fig.image_path
                    if image_path.exists():
                        result_lines.append(f'![{fig.title}]({fig.image_path})')
                        result_lines.append(f'*{fig.caption}*')
                        result_lines.append('')
            for chart in orphan_charts:
                if chart.image_path:
                    image_path = Path(chart.image_path) if not isinstance(chart.image_path, Path) else chart.image_path
                    if image_path.exists():
                        result_lines.append(f'![{chart.title}]({chart.image_path})')
                        result_lines.append(f'*{chart.caption}*')
                        result_lines.append('')
            for table in orphan_tables:
                result_lines.append(f'**{table.title}**')
                result_lines.append('')
                header_line = '| ' + ' | '.join(str(h) for h in table.headers) + ' |'
                separator = '|' + '|'.join(['---'] * len(table.headers)) + '|'
                result_lines.append(header_line)
                result_lines.append(separator)
                for row in table.rows:
                    row_line = '| ' + ' | '.join(str(cell) for cell in row) + ' |'
                    result_lines.append(row_line)
                result_lines.append('')

        return '\n'.join(result_lines)

    # ========================================================================
    # DOCX Integration
    # ========================================================================

    def add_figures_to_docx(
        self,
        docx_path: Path,
        collection: FigureTableCollection,
        output_path: Path = None,
    ) -> Optional[Path]:
        """
        Add figures and tables to an existing DOCX file.

        Args:
            docx_path: Path to original DOCX
            collection: Collection of figures and tables
            output_path: Output path (defaults to *_with_figures.docx)

        Returns:
            Path to updated DOCX or None on failure
        """
        try:
            from docx import Document
            from docx.shared import Inches, Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            print("[FigureTableGenerator] python-docx not installed. Install with: pip install python-docx")
            return None

        # Formats python-docx can embed; SVG/WebP downloads must be skipped,
        # not attempted (add_picture raises UnrecognizedImageError)
        EMBEDDABLE_IMAGE_SUFFIXES = {'.png', '.jpg', '.jpeg', '.gif', '.bmp', '.tiff'}

        try:
            doc = Document(docx_path)
            output_path = output_path or docx_path.parent / f"{docx_path.stem}_with_figures.docx"

            # Track which section we're in.
            # IMPORTANT: insertion positions must be indices into the body's
            # block-level children (paragraphs AND tables), because
            # _move_paragraph_after/_move_table_after reposition within that
            # list. Indexing doc.paragraphs would drift by one for every
            # table already present in the document.
            body_children = list(doc.element.body)

            current_section = None
            insert_after = []
            matched_sections = set()

            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()

                # Detect section headers (match "1. " or "1.2. " format)
                match = re.match(r'^(\d+(?:\.\d+)?)[.\s]', text)
                if match and paragraph.style.name.startswith('Heading'):
                    try:
                        body_idx = body_children.index(paragraph._element)  # noqa: SLF001
                    except ValueError:
                        continue  # paragraph nested in a table etc.
                    if current_section:
                        insert_after.append((body_idx, current_section))
                    current_section = match.group(1)

            # Insert at end for last section
            if current_section:
                insert_after.append((len(body_children), current_section))

            def _image_exists(img_path) -> bool:
                """Check if image path exists on disk."""
                if not img_path:
                    return False
                p = Path(img_path) if not isinstance(img_path, Path) else img_path
                return p.exists()

            def _embeddable(img_path) -> bool:
                if not _image_exists(img_path):
                    return False
                suffix = Path(img_path).suffix.lower()
                if suffix and suffix not in EMBEDDABLE_IMAGE_SUFFIXES:
                    print(f"[FigureTableGenerator] skipping non-embeddable "
                          f"image format {suffix}: {img_path}")
                    return False
                return True

            skipped_elements = 0

            # Insert in reverse order to maintain indices
            for insert_idx, section_id in reversed(insert_after):
                matched_sections.add(section_id)
                section_tables = collection.get_tables_for_section(section_id)
                section_charts = collection.get_charts_for_section(section_id)
                section_figures = collection.get_figures_for_section(section_id)

                elements = []

                # Add charts
                for chart in reversed(section_charts):
                    if _embeddable(chart.image_path):
                        elements.append(('image', chart))

                # Add tables
                for table in reversed(section_tables):
                    elements.append(('table', table))

                # Add figures
                for fig in reversed(section_figures):
                    if _embeddable(fig.image_path):
                        elements.append(('image', fig))

                for elem_type, elem in elements:
                  # One corrupt image or malformed table must not abort the
                  # whole insertion (previously a single failure dropped
                  # every figure in the report)
                  try:
                    if elem_type == 'image':
                        # Add image
                        p = doc.add_paragraph()
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run = p.add_run()
                        run.add_picture(str(elem.image_path), width=Inches(5.5))

                        # Add caption
                        caption_p = doc.add_paragraph()
                        caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        caption_run = caption_p.add_run(elem.caption)
                        caption_run.font.size = Pt(9)
                        caption_run.italic = True

                        # Move to correct position
                        self._move_paragraph_after(doc, p, insert_idx)
                        self._move_paragraph_after(doc, caption_p, insert_idx + 1)

                    elif elem_type == 'table':
                        # Add table title
                        title_p = doc.add_paragraph()
                        title_run = title_p.add_run(elem.title)
                        title_run.bold = True
                        title_run.font.size = Pt(10)

                        # Add table
                        docx_table = doc.add_table(
                            rows=len(elem.rows) + 1,
                            cols=len(elem.headers),
                        )
                        docx_table.style = 'Light Grid'

                        # Headers
                        for j, header in enumerate(elem.headers):
                            docx_table.rows[0].cells[j].text = str(header)

                        # Data rows
                        for row_idx, row in enumerate(elem.rows):
                            for col_idx, cell in enumerate(row):
                                if col_idx < len(elem.headers):
                                    docx_table.rows[row_idx + 1].cells[col_idx].text = str(cell)

                        # Add caption below the table
                        caption_p = None
                        if elem.caption and elem.caption != elem.title:
                            caption_p = doc.add_paragraph()
                            caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            caption_run = caption_p.add_run(elem.caption)
                            caption_run.font.size = Pt(9)
                            caption_run.italic = True

                        # Move title, table, and caption to correct position
                        self._move_paragraph_after(doc, title_p, insert_idx)
                        self._move_table_after(doc, docx_table, insert_idx + 1)
                        if caption_p:
                            self._move_paragraph_after(doc, caption_p, insert_idx + 2)
                  except Exception as elem_err:
                    skipped_elements += 1
                    label = getattr(elem, 'title', '') or getattr(elem, 'caption', '')
                    print(f"[FigureTableGenerator] skipping one "
                          f"{elem_type} ('{str(label)[:40]}') in section "
                          f"{section_id}: {elem_err}")

            # Append orphaned figures (section_id="" or unmatched) at end of document
            orphan_figures = [f for f in collection.figures if f.section_id not in matched_sections]
            orphan_charts = [c for c in collection.charts if c.section_id not in matched_sections]
            for item in orphan_figures + orphan_charts:
                if not _embeddable(item.image_path):
                    continue
                try:
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = p.add_run()
                    img_path = Path(item.image_path) if not isinstance(item.image_path, Path) else item.image_path
                    run.add_picture(str(img_path), width=Inches(5.5))
                    caption_p = doc.add_paragraph()
                    caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    caption_run = caption_p.add_run(item.caption)
                    caption_run.font.size = Pt(9)
                    caption_run.italic = True
                except Exception as elem_err:
                    skipped_elements += 1
                    print(f"[FigureTableGenerator] skipping one orphan image: {elem_err}")

            if skipped_elements:
                print(f"[FigureTableGenerator] {skipped_elements} element(s) "
                      f"skipped during DOCX insertion (see messages above)")

            doc.save(output_path)
            return output_path

        except Exception as e:
            print(f"[FigureTableGenerator] Error adding figures to DOCX: {e}")
            return None

    @staticmethod
    def _move_paragraph_after(doc, paragraph, index):
        """Move a paragraph/table to after the given index in document body.

        Uses lxml XML manipulation to relocate the element, since
        python-docx does not natively support paragraph reordering.
        """
        try:
            body = doc.element.body
            # Collect all block-level children (paragraphs + tables)
            children = list(body)
            para_elem = paragraph._element  # noqa: SLF001

            # Clamp index
            if index < 0:
                index = 0
            if index >= len(children):
                # Already at the end, nothing to move
                return

            # Remove from current position
            body.remove(para_elem)
            # Re-read children after removal
            children = list(body)

            if index >= len(children):
                body.append(para_elem)
            else:
                children[index].addprevious(para_elem)
        except Exception as e:
            # Non-fatal: figure stays at end of document
            print(f"[FigureTableGenerator] Could not reposition element: {e}")

    @staticmethod
    def _move_table_after(doc, table, index):
        """Move a table element to after the given index in document body."""
        try:
            body = doc.element.body
            children = list(body)
            tbl_elem = table._tbl  # noqa: SLF001

            if index < 0:
                index = 0
            if index >= len(children):
                return

            body.remove(tbl_elem)
            children = list(body)

            if index >= len(children):
                body.append(tbl_elem)
            else:
                children[index].addprevious(tbl_elem)
        except Exception as e:
            print(f"[FigureTableGenerator] Could not reposition table: {e}")

    # ========================================================================
    # PDF Integration
    # ========================================================================

    def add_figures_to_pdf(
        self,
        markdown_content: str,
        collection: FigureTableCollection,
        output_path: Path,
    ) -> Optional[Path]:
        """
        Generate a PDF with figures and tables from markdown + collection.

        Uses the markdown integration and then converts to PDF.

        Args:
            markdown_content: Original markdown content
            collection: Collection of figures and tables
            output_path: Output PDF path

        Returns:
            Path to generated PDF or None on failure
        """
        # First, add figures to markdown
        updated_md = self.add_figures_to_markdown(markdown_content, collection)

        # Try to convert to PDF using available tools
        try:
            import markdown

            # Convert markdown to HTML
            html_content = markdown.markdown(
                updated_md,
                extensions=['tables', 'fenced_code'],
            )

            # Embed chart images as base64 in HTML
            html_content = self._embed_images_in_html(html_content, collection)

            # Wrap in HTML template
            html_doc = f"""<!DOCTYPE html>
<html>
<head>
<meta charset="utf-8">
<style>
body {{ font-family: 'Noto Sans CJK JP', 'Yu Gothic', sans-serif; margin: 40px; line-height: 1.6; }}
table {{ border-collapse: collapse; width: 100%; margin: 15px 0; }}
th, td {{ border: 1px solid #ddd; padding: 8px; text-align: left; }}
th {{ background-color: #4C78A8; color: white; }}
tr:nth-child(even) {{ background-color: #f9f9f9; }}
img {{ max-width: 100%; height: auto; display: block; margin: 15px auto; }}
h1 {{ color: #2c3e50; border-bottom: 2px solid #4C78A8; padding-bottom: 10px; }}
h2 {{ color: #34495e; border-bottom: 1px solid #bdc3c7; padding-bottom: 5px; }}
em {{ color: #7f8c8d; font-size: 0.9em; }}
</style>
</head>
<body>
{html_content}
</body>
</html>"""

            # Write HTML (can be converted to PDF via weasyprint or wkhtmltopdf)
            html_path = output_path.with_suffix('.html')
            with open(html_path, 'w', encoding='utf-8') as f:
                f.write(html_doc)

            # Try weasyprint for PDF conversion
            try:
                from weasyprint import HTML
                HTML(filename=str(html_path)).write_pdf(str(output_path))
                return output_path
            except ImportError:
                pass

            # Fallback: return HTML if PDF conversion unavailable
            print("[FigureTableGenerator] PDF conversion requires weasyprint. HTML generated instead.")
            return html_path

        except ImportError:
            print("[FigureTableGenerator] markdown library not installed for PDF generation.")
            return None
        except Exception as e:
            print(f"[FigureTableGenerator] Error generating PDF: {e}")
            return None

    def _embed_images_in_html(self, html_content: str, collection: FigureTableCollection) -> str:
        """Replace image file paths with base64 embedded images in HTML."""
        all_figures = collection.figures + collection.charts
        for fig in all_figures:
            if fig.image_path and fig.image_path.exists():
                try:
                    with open(fig.image_path, 'rb') as f:
                        img_data = base64.b64encode(f.read()).decode('utf-8')
                    ext = fig.image_path.suffix.lstrip('.')
                    if ext == 'jpg':
                        ext = 'jpeg'
                    data_uri = f"data:image/{ext};base64,{img_data}"
                    html_content = html_content.replace(str(fig.image_path), data_uri)
                except Exception:
                    pass
        return html_content

    # ========================================================================
    # Export / Utility
    # ========================================================================

    def export_collection(
        self,
        collection: FigureTableCollection,
        filepath: Path,
    ) -> None:
        """Export figure/table collection to JSON."""
        with open(filepath, 'w', encoding='utf-8') as f:
            json.dump(collection.to_dict(), f, ensure_ascii=False, indent=2)


def add_figures_to_report(
    report_path: Path,
    session,
    evidence_locker,
    llm_client=None,
    output_dir: Path = None,
    language: str = "ja",
    proxies: Dict[str, str] = None,
    verify_ssl: bool = True,
    chart_library: str = "matplotlib",
) -> Path:
    """
    Add figures and tables to an existing report.

    Args:
        report_path: Path to the report file
        session: Research session
        evidence_locker: Evidence locker
        llm_client: Optional LLM client for analysis
        output_dir: Output directory for images
        language: Language for captions
        proxies: Proxy settings dict
        verify_ssl: Whether to verify SSL certificates
        chart_library: Chart rendering library ("matplotlib" or "seaborn")

    Returns:
        Path to the updated report
    """
    if output_dir is None:
        output_dir = report_path.parent / "figures"
    output_dir.mkdir(parents=True, exist_ok=True)

    generator = FigureTableGenerator(
        llm_client=llm_client,
        output_dir=output_dir,
        language=language,
        proxies=proxies,
        verify_ssl=verify_ssl,
        chart_library=chart_library,
    )

    collection = generator.generate_figures_and_tables(
        session=session,
        evidence_locker=evidence_locker,
    )

    # Read existing report with encoding fallback
    content = None
    encodings = ['utf-8', 'utf-8-sig', 'cp932', 'shift_jis', 'latin-1']
    for encoding in encodings:
        try:
            with open(report_path, 'r', encoding=encoding) as f:
                content = f.read()
            break
        except UnicodeDecodeError:
            continue

    if content is None:
        with open(report_path, 'r', encoding='utf-8', errors='replace') as f:
            content = f.read()

    suffix = report_path.suffix.lower()

    if suffix == '.md':
        updated_content = generator.add_figures_to_markdown(content, collection)
        updated_path = report_path.parent / f"{report_path.stem}_with_figures.md"
        with open(updated_path, 'w', encoding='utf-8') as f:
            f.write(updated_content)

    elif suffix == '.docx':
        updated_path = generator.add_figures_to_docx(
            docx_path=report_path,
            collection=collection,
        )
        if not updated_path:
            updated_path = report_path

    elif suffix == '.pdf' or suffix == '.html':
        updated_path = generator.add_figures_to_pdf(
            markdown_content=content,
            collection=collection,
            output_path=report_path.parent / f"{report_path.stem}_with_figures.pdf",
        )
        if not updated_path:
            updated_path = report_path

    else:
        # Fallback: treat as markdown
        updated_content = generator.add_figures_to_markdown(content, collection)
        updated_path = report_path.parent / f"{report_path.stem}_with_figures.md"
        with open(updated_path, 'w', encoding='utf-8') as f:
            f.write(updated_content)

    # Export collection metadata
    collection_path = output_dir / "figures_tables.json"
    generator.export_collection(collection, collection_path)

    return updated_path
