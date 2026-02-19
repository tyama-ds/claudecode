"""
Report Generator - Generate research reports in multiple formats.
"""

import re
from datetime import datetime
from enum import Enum
from pathlib import Path
from typing import Optional, Dict, Any, List, Tuple

from ..evidence.locker import EvidenceLocker, QualityCategory, SourceType
from ..evidence.quality_evaluator import get_quality_summary
from ..research.researcher import ResearchSession
from ..verification.verifier import VerificationResult
from .length_controller import (
    ContentLengthController,
    LengthTarget,
    LengthInfo,
    get_length_summary,
)


class ReportFormat(str, Enum):
    """Supported report formats."""
    MARKDOWN = "markdown"
    DOCX = "docx"
    PDF = "pdf"
    HTML = "html"


class ReportGenerator:
    """
    Generate research reports in various formats.

    Supports:
    - Markdown (.md)
    - Word Document (.docx)
    - PDF (.pdf)
    - HTML (.html)
    """

    def __init__(
        self,
        output_dir: Path = None,
        include_toc: bool = True,
        include_citations: bool = True,
        include_images: bool = True,
        include_quality_summary: bool = True,
        language: str = "ja",
    ):
        """
        Initialize ReportGenerator.

        Args:
            output_dir: Directory for output files
            include_toc: Include table of contents
            include_citations: Include citations/references
            include_images: Include images in report
            include_quality_summary: Include quality statistics summary
            language: Report language
        """
        self.output_dir = output_dir or Path("./output/reports")
        self.output_dir.mkdir(parents=True, exist_ok=True)
        self.include_toc = include_toc
        self.include_citations = include_citations
        self.include_images = include_images
        self.include_quality_summary = include_quality_summary
        self.language = language

    def generate_report(
        self,
        session: ResearchSession,
        evidence_locker: EvidenceLocker,
        format: ReportFormat = ReportFormat.MARKDOWN,
        verification_result: VerificationResult = None,
        filename: str = None,
        min_quality: QualityCategory = None,
        quality_categories: List[QualityCategory] = None,
        source_types: List[SourceType] = None,
        target_pages: int = None,
        target_characters: int = None,
    ) -> Path:
        """
        Generate a complete research report.

        Args:
            session: Research session with content
            evidence_locker: Evidence locker with sources
            format: Output format
            verification_result: Optional verification results
            filename: Custom filename (without extension)
            min_quality: Minimum quality category for citations
            quality_categories: Specific quality categories to include
            source_types: Source types to include in citations
            target_pages: Target page count (approximate)
            target_characters: Target character count

        Returns:
            Path to generated report
        """
        if not filename:
            filename = f"research_report_{session.session_id}"

        # Create a filtered evidence list if quality filters are specified
        if min_quality or quality_categories or source_types:
            filtered_evidence = evidence_locker.filter_evidence(
                min_quality=min_quality,
                quality_categories=quality_categories,
                source_types=source_types,
            )
        else:
            filtered_evidence = None  # Use all evidence

        # Apply content length adjustment if targets are specified
        adjusted_session = self._apply_length_adjustment(
            session=session,
            format=format,
            target_pages=target_pages,
            target_characters=target_characters,
        )

        if format == ReportFormat.MARKDOWN:
            return self._generate_markdown(
                adjusted_session, evidence_locker, verification_result, filename, filtered_evidence
            )
        elif format == ReportFormat.DOCX:
            return self._generate_docx(
                adjusted_session, evidence_locker, verification_result, filename, filtered_evidence
            )
        elif format == ReportFormat.PDF:
            return self._generate_pdf(
                adjusted_session, evidence_locker, verification_result, filename, filtered_evidence
            )
        elif format == ReportFormat.HTML:
            return self._generate_html(
                adjusted_session, evidence_locker, verification_result, filename, filtered_evidence
            )
        else:
            raise ValueError(f"Unsupported format: {format}")

    def _apply_length_adjustment(
        self,
        session: ResearchSession,
        format: ReportFormat,
        target_pages: int = None,
        target_characters: int = None,
    ) -> ResearchSession:
        """
        Apply content length adjustment to meet page/character targets.

        Args:
            session: Original research session
            format: Output format
            target_pages: Target page count
            target_characters: Target character count

        Returns:
            Session with adjusted content (or original if no targets)
        """
        # No adjustment needed if no targets specified
        if target_pages is None and target_characters is None:
            return session

        # Create length target
        length_target = LengthTarget(
            target_pages=target_pages,
            target_characters=target_characters,
        )

        # Create controller
        controller = ContentLengthController(
            target=length_target,
            format_type=format.value,
            language=self.language,
        )

        # Check if adjustment is needed
        needs_adj, ratio = controller.needs_adjustment(session.section_contents)

        if not needs_adj:
            return session

        # Create a copy of the session with adjusted content
        adjusted_contents = controller.adjust_content(
            session.section_contents,
            adjustment_ratio=ratio,
        )

        # Create a modified session (shallow copy with new section_contents)
        # We don't modify the original session
        from copy import copy
        adjusted_session = copy(session)
        adjusted_session.section_contents = adjusted_contents

        return adjusted_session

    def get_length_info(
        self,
        session: ResearchSession,
        format: ReportFormat = ReportFormat.PDF,
    ) -> LengthInfo:
        """
        Get length information for a research session.

        Args:
            session: Research session
            format: Output format for estimation

        Returns:
            LengthInfo with character count and page estimate
        """
        controller = ContentLengthController(
            format_type=format.value,
            language=self.language,
        )
        return controller.calculate_length(session.section_contents)

    def get_length_summary(
        self,
        session: ResearchSession,
        format: ReportFormat = ReportFormat.PDF,
    ) -> str:
        """
        Get human-readable length summary for a session.

        Args:
            session: Research session
            format: Output format for estimation

        Returns:
            Summary string
        """
        return get_length_summary(
            section_contents=session.section_contents,
            format_type=format.value,
            language=self.language,
        )

    def _generate_markdown(
        self,
        session: ResearchSession,
        evidence_locker: EvidenceLocker,
        verification_result: VerificationResult,
        filename: str,
        filtered_evidence: List = None,
    ) -> Path:
        """Generate Markdown report."""
        plan = session.research_plan
        content_parts = []

        # Use filtered evidence or all evidence
        evidence_list = filtered_evidence if filtered_evidence is not None else evidence_locker.get_all_evidence()

        # Renumber citations to match final reference order
        renumbered_contents = self._renumber_all_citations(session, evidence_list)

        # Title and metadata
        content_parts.append(f"# {plan.title if plan else session.query}\n")
        content_parts.append(f"**Research Date:** {session.started_at[:10]}\n")
        content_parts.append(f"**Session ID:** {session.session_id}\n")
        content_parts.append("\n---\n")

        # Executive Summary
        exec_summary = renumbered_contents.get("_executive_summary", {})
        if exec_summary:
            content_parts.append("## Executive Summary\n")
            content_parts.append(exec_summary.get("executive_summary", "") + "\n")

            if exec_summary.get("key_findings"):
                content_parts.append("\n### Key Findings\n")
                for finding in exec_summary["key_findings"]:
                    content_parts.append(f"- {finding}\n")

            content_parts.append("\n---\n")

        # Table of Contents
        if self.include_toc and plan:
            content_parts.append("## Table of Contents\n")
            for item in plan.table_of_contents.items:
                content_parts.append(f"- [{item.section}. {item.title}](#{self._slugify(item.title)})\n")
                for sub in item.subsections:
                    content_parts.append(f"  - [{sub.section}. {sub.title}](#{self._slugify(sub.title)})\n")
            content_parts.append("\n---\n")

        # Main Content Sections
        if plan:
            for item in plan.table_of_contents.items:
                section_content = renumbered_contents.get(item.section, {})

                content_parts.append(f"\n## {item.section}. {item.title}\n")

                if section_content:
                    # Section content with citation markers converted to footnotes
                    main_content = section_content.get("content", "")
                    content_parts.append(self._process_citations(main_content) + "\n")

                    # Images
                    if self.include_images and section_content.get("images"):
                        content_parts.append("\n### Related Figures\n")
                        for i, img in enumerate(section_content["images"][:3]):
                            caption = img.get("suggested_caption", f"Figure {i+1}")
                            content_parts.append(f"![{caption}]({img.get('src', '')})\n")
                            content_parts.append(f"*{caption}*\n\n")

                # Subsections
                for sub in item.subsections:
                    sub_content = renumbered_contents.get(sub.section, {})
                    content_parts.append(f"\n### {sub.section}. {sub.title}\n")
                    if sub_content:
                        content_parts.append(self._process_citations(
                            sub_content.get("content", "")
                        ) + "\n")

        # Verification Summary
        if verification_result:
            content_parts.append("\n---\n")
            content_parts.append("## Verification Summary\n")
            content_parts.append(f"**Overall Reliability:** {verification_result.overall_reliability_score:.1%}\n\n")
            content_parts.append("| Confidence Level | Count |\n")
            content_parts.append("|-----------------|-------|\n")
            content_parts.append(f"| High | {verification_result.high_confidence_count} |\n")
            content_parts.append(f"| Medium | {verification_result.medium_confidence_count} |\n")
            content_parts.append(f"| Low | {verification_result.low_confidence_count} |\n")
            content_parts.append(f"| Unsupported | {verification_result.unsupported_count} |\n")

            if verification_result.hallucination_risk_count > 0:
                content_parts.append(f"\n**Hallucination Risks:** {verification_result.hallucination_risk_count} claims flagged\n")

        # Quality Summary
        if self.include_quality_summary:
            content_parts.append("\n---\n")
            content_parts.append("## Source Quality Summary\n")
            quality_stats = evidence_locker.get_quality_statistics()

            if self.language == "ja":
                content_parts.append(f"**総エビデンス数:** {quality_stats['total_evidence']}\n")
                content_parts.append(f"**高品質情報の割合:** {quality_stats['high_quality_percentage']}%\n")
                content_parts.append(f"**平均品質スコア:** {quality_stats['average_quality_score']}\n\n")

                content_parts.append("### 品質カテゴリ分布\n")
                content_parts.append("| カテゴリ | 件数 |\n")
                content_parts.append("|----------|------|\n")
                category_labels = {
                    "authoritative": "権威的",
                    "high": "高品質",
                    "medium": "中品質",
                    "low": "低品質",
                    "unverified": "未検証",
                }
                for cat, count in quality_stats.get("quality_distribution", {}).items():
                    label = category_labels.get(cat, cat)
                    content_parts.append(f"| {label} | {count} |\n")
            else:
                content_parts.append(f"**Total Evidence:** {quality_stats['total_evidence']}\n")
                content_parts.append(f"**High Quality Percentage:** {quality_stats['high_quality_percentage']}%\n")
                content_parts.append(f"**Average Quality Score:** {quality_stats['average_quality_score']}\n\n")

                content_parts.append("### Quality Distribution\n")
                content_parts.append("| Category | Count |\n")
                content_parts.append("|----------|-------|\n")
                for cat, count in quality_stats.get("quality_distribution", {}).items():
                    content_parts.append(f"| {cat.title()} | {count} |\n")

            # Note if filtering was applied
            if filtered_evidence is not None:
                total_all = len(evidence_locker.get_all_evidence())
                if self.language == "ja":
                    content_parts.append(f"\n*注: このレポートでは品質フィルタリングにより {len(evidence_list)}/{total_all} 件のエビデンスを使用しています。*\n")
                else:
                    content_parts.append(f"\n*Note: This report uses {len(evidence_list)}/{total_all} evidence items after quality filtering.*\n")

        # References
        if self.include_citations:
            content_parts.append("\n---\n")
            content_parts.append("## References\n")

            for i, evidence in enumerate(evidence_list, 1):
                quality_badge = self._get_quality_badge(evidence.quality_category)
                content_parts.append(f"{i}. {quality_badge} {evidence.citation_text}\n")

        # Write file
        filepath = self.output_dir / f"{filename}.md"
        with open(filepath, "w", encoding="utf-8") as f:
            f.write("".join(content_parts))

        return filepath

    def _get_quality_badge(self, quality: QualityCategory) -> str:
        """Get quality badge for citation."""
        badges = {
            QualityCategory.AUTHORITATIVE: "[A]",
            QualityCategory.HIGH: "[H]",
            QualityCategory.MEDIUM: "[M]",
            QualityCategory.LOW: "[L]",
            QualityCategory.UNVERIFIED: "[?]",
        }
        return badges.get(quality, "")

    def _generate_docx(
        self,
        session: ResearchSession,
        evidence_locker: EvidenceLocker,
        verification_result: VerificationResult,
        filename: str,
        filtered_evidence: List = None,
    ) -> Path:
        """Generate Word document report."""
        try:
            from docx import Document
            from docx.shared import Inches, Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            raise ImportError(
                "python-docx not installed. Install with: pip install python-docx"
            )

        doc = Document()
        plan = session.research_plan

        # Use filtered evidence or all evidence
        evidence_list = filtered_evidence if filtered_evidence is not None else evidence_locker.get_all_evidence()

        # Renumber citations to match final reference order
        renumbered_contents = self._renumber_all_citations(session, evidence_list)

        # Title
        title = doc.add_heading(plan.title if plan else session.query, 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Metadata
        meta = doc.add_paragraph()
        meta.add_run(f"Research Date: {session.started_at[:10]}\n").bold = True
        meta.add_run(f"Session ID: {session.session_id}")

        doc.add_paragraph()  # Spacing

        # Executive Summary
        exec_summary = renumbered_contents.get("_executive_summary", {})
        if exec_summary:
            doc.add_heading("Executive Summary", level=1)
            doc.add_paragraph(exec_summary.get("executive_summary", ""))

            if exec_summary.get("key_findings"):
                doc.add_heading("Key Findings", level=2)
                for finding in exec_summary["key_findings"]:
                    doc.add_paragraph(finding, style="List Bullet")

        # Table of Contents placeholder
        if self.include_toc and plan:
            doc.add_heading("Table of Contents", level=1)
            for item in plan.table_of_contents.items:
                try:
                    doc.add_paragraph(
                        f"{item.section}. {item.title}",
                        style="TOC Heading"
                    )
                except KeyError:
                    # "TOC Heading" style may not exist in the default template
                    para = doc.add_paragraph(f"{item.section}. {item.title}")
                    if para.runs:
                        para.runs[0].bold = True

        # Main Content
        if plan:
            for item in plan.table_of_contents.items:
                section_content = renumbered_contents.get(item.section, {})

                doc.add_heading(f"{item.section}. {item.title}", level=1)

                if section_content:
                    content = section_content.get("content", "")
                    # Split into paragraphs
                    paragraphs = content.split("\n\n")
                    for para in paragraphs:
                        if para.strip():
                            doc.add_paragraph(self._strip_citations(para))

                # Subsections
                for sub in item.subsections:
                    sub_content = renumbered_contents.get(sub.section, {})
                    doc.add_heading(f"{sub.section}. {sub.title}", level=2)
                    if sub_content:
                        doc.add_paragraph(self._strip_citations(
                            sub_content.get("content", "")
                        ))

        # Verification Summary
        if verification_result:
            doc.add_heading("Verification Summary", level=1)
            doc.add_paragraph(
                f"Overall Reliability: {verification_result.overall_reliability_score:.1%}"
            )

            table = doc.add_table(rows=5, cols=2)
            table.style = "Table Grid"
            table.rows[0].cells[0].text = "Confidence Level"
            table.rows[0].cells[1].text = "Count"
            table.rows[1].cells[0].text = "High"
            table.rows[1].cells[1].text = str(verification_result.high_confidence_count)
            table.rows[2].cells[0].text = "Medium"
            table.rows[2].cells[1].text = str(verification_result.medium_confidence_count)
            table.rows[3].cells[0].text = "Low"
            table.rows[3].cells[1].text = str(verification_result.low_confidence_count)
            table.rows[4].cells[0].text = "Unsupported"
            table.rows[4].cells[1].text = str(verification_result.unsupported_count)

        # Quality Summary
        if self.include_quality_summary:
            doc.add_heading("Source Quality Summary", level=1)
            quality_stats = evidence_locker.get_quality_statistics()
            doc.add_paragraph(
                f"Total Evidence: {quality_stats['total_evidence']}\n"
                f"High Quality Percentage: {quality_stats['high_quality_percentage']}%\n"
                f"Average Quality Score: {quality_stats['average_quality_score']}"
            )

        # References
        if self.include_citations:
            doc.add_heading("References", level=1)
            for i, evidence in enumerate(evidence_list, 1):
                quality_badge = self._get_quality_badge(evidence.quality_category)
                doc.add_paragraph(f"{i}. {quality_badge} {evidence.citation_text}")

        # Save
        filepath = self.output_dir / f"{filename}.docx"
        doc.save(filepath)

        return filepath

    def _register_japanese_fonts(self):
        """Register Japanese fonts for PDF generation."""
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
        import os

        # Try to register a Japanese font
        font_registered = False
        registered_font_name = None

        # Common Japanese font paths on different systems
        japanese_fonts = [
            # Linux - Noto fonts (most common)
            ("/usr/share/fonts/truetype/noto/NotoSansCJK-Regular.ttc", "NotoSansCJK"),
            ("/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc", "NotoSansCJK"),
            ("/usr/share/fonts/noto-cjk/NotoSansCJK-Regular.ttc", "NotoSansCJK"),
            # Linux - other fonts
            ("/usr/share/fonts/truetype/takao-gothic/TakaoPGothic.ttf", "TakaoPGothic"),
            ("/usr/share/fonts/opentype/ipafont-gothic/ipag.ttf", "IPAGothic"),
            ("/usr/share/fonts/truetype/vlgothic/VL-Gothic-Regular.ttf", "VLGothic"),
            ("/usr/share/fonts/truetype/fonts-japanese-gothic.ttf", "JapaneseGothic"),
            # Google fonts location
            ("/usr/share/fonts/truetype/noto/NotoSansJP-Regular.ttf", "NotoSansJP"),
            # macOS
            ("/System/Library/Fonts/ヒラギノ角ゴシック W3.ttc", "Hiragino"),
            ("/Library/Fonts/Arial Unicode.ttf", "ArialUnicode"),
            # Windows
            ("C:/Windows/Fonts/msgothic.ttc", "MSGothic"),
            ("C:/Windows/Fonts/meiryo.ttc", "Meiryo"),
            ("C:/Windows/Fonts/YuGothic.ttc", "YuGothic"),
        ]

        for font_path, font_name in japanese_fonts:
            if os.path.exists(font_path):
                try:
                    pdfmetrics.registerFont(TTFont(font_name, font_path))
                    font_registered = True
                    registered_font_name = font_name
                    break
                except Exception as e:
                    # Font might be in use or incompatible
                    continue

        # Try CID fonts as fallback (built into reportlab for Asian languages)
        if not font_registered:
            try:
                from reportlab.pdfbase.cidfonts import UnicodeCIDFont
                pdfmetrics.registerFont(UnicodeCIDFont('HeiseiKakuGo-W5'))
                registered_font_name = 'HeiseiKakuGo-W5'
                font_registered = True
            except Exception:
                pass

            if not font_registered:
                try:
                    from reportlab.pdfbase.cidfonts import UnicodeCIDFont
                    pdfmetrics.registerFont(UnicodeCIDFont('HeiseiMin-W3'))
                    registered_font_name = 'HeiseiMin-W3'
                    font_registered = True
                except Exception:
                    pass

        return registered_font_name

    def _generate_pdf(
        self,
        session: ResearchSession,
        evidence_locker: EvidenceLocker,
        verification_result: VerificationResult,
        filename: str,
        filtered_evidence: List = None,
    ) -> Path:
        """Generate PDF report using reportlab with Japanese font support."""
        try:
            from reportlab.lib import colors
            from reportlab.lib.pagesizes import A4
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import inch
            from reportlab.platypus import (
                SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
                PageBreak
            )
            from reportlab.pdfbase import pdfmetrics
            from reportlab.pdfbase.ttfonts import TTFont
        except ImportError:
            raise ImportError(
                "reportlab not installed. Install with: pip install reportlab"
            )

        filepath = self.output_dir / f"{filename}.pdf"
        doc = SimpleDocTemplate(str(filepath), pagesize=A4)
        styles = getSampleStyleSheet()
        story = []

        plan = session.research_plan

        # Use filtered evidence or all evidence
        evidence_list = filtered_evidence if filtered_evidence is not None else evidence_locker.get_all_evidence()

        # Renumber citations to match final reference order
        renumbered_contents = self._renumber_all_citations(session, evidence_list)

        # Register Japanese fonts if language is Japanese
        japanese_font = None
        if self.language == "ja":
            japanese_font = self._register_japanese_fonts()

        # Determine fonts to use
        base_font = japanese_font if japanese_font else 'Helvetica'
        base_font_bold = japanese_font if japanese_font else 'Helvetica-Bold'

        # Custom styles with Japanese font support
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontName=base_font,
            fontSize=24,
            spaceAfter=30,
            alignment=1,  # Center
        )

        heading1_style = ParagraphStyle(
            'CustomH1',
            parent=styles['Heading1'],
            fontName=base_font,
            fontSize=16,
            spaceBefore=20,
            spaceAfter=10,
        )

        heading2_style = ParagraphStyle(
            'CustomH2',
            parent=styles['Heading2'],
            fontName=base_font,
            fontSize=14,
            spaceBefore=15,
            spaceAfter=8,
        )

        body_style = ParagraphStyle(
            'CustomBody',
            parent=styles['Normal'],
            fontName=base_font,
            fontSize=11,
            spaceBefore=6,
            spaceAfter=6,
            wordWrap='CJK',  # Enable CJK word wrapping
        )

        # Title
        story.append(Paragraph(
            plan.title if plan else session.query,
            title_style
        ))
        story.append(Spacer(1, 12))
        story.append(Paragraph(
            f"<b>Research Date:</b> {session.started_at[:10]}",
            body_style
        ))
        story.append(Paragraph(
            f"<b>Session ID:</b> {session.session_id}",
            body_style
        ))
        story.append(Spacer(1, 30))

        # Executive Summary
        exec_summary = renumbered_contents.get("_executive_summary", {})
        if exec_summary:
            story.append(Paragraph("Executive Summary", heading1_style))
            story.append(Paragraph(
                exec_summary.get("executive_summary", ""),
                body_style
            ))
            story.append(Spacer(1, 20))

        # Main Content
        if plan:
            for item in plan.table_of_contents.items:
                section_content = renumbered_contents.get(item.section, {})

                story.append(Paragraph(
                    f"{item.section}. {item.title}",
                    heading1_style
                ))

                if section_content:
                    content = self._strip_citations(
                        section_content.get("content", "")
                    )
                    # Split long content into paragraphs
                    for para in content.split("\n\n"):
                        if para.strip():
                            # Escape special characters for reportlab
                            safe_para = para.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
                            try:
                                story.append(Paragraph(safe_para, body_style))
                            except Exception:
                                story.append(Paragraph(para[:500] + "...", body_style))

                # Subsections
                for sub in item.subsections:
                    sub_content = renumbered_contents.get(sub.section, {})
                    story.append(Paragraph(
                        f"{sub.section}. {sub.title}",
                        heading2_style
                    ))
                    if sub_content:
                        safe_content = self._strip_citations(
                            sub_content.get("content", "")
                        ).replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
                        try:
                            story.append(Paragraph(safe_content[:2000], body_style))
                        except Exception:
                            story.append(Paragraph("Content unavailable", body_style))

        # Verification Summary
        if verification_result:
            story.append(PageBreak())
            story.append(Paragraph("Verification Summary", heading1_style))
            story.append(Paragraph(
                f"<b>Overall Reliability:</b> {verification_result.overall_reliability_score:.1%}",
                body_style
            ))
            story.append(Spacer(1, 10))

            # Verification table
            table_data = [
                ["Confidence Level", "Count"],
                ["High", str(verification_result.high_confidence_count)],
                ["Medium", str(verification_result.medium_confidence_count)],
                ["Low", str(verification_result.low_confidence_count)],
                ["Unsupported", str(verification_result.unsupported_count)],
            ]

            table = Table(table_data, colWidths=[3*inch, 1.5*inch])
            table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 12),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                ('GRID', (0, 0), (-1, -1), 1, colors.black),
            ]))
            story.append(table)

        # References
        if self.include_citations:
            story.append(PageBreak())
            story.append(Paragraph("References", heading1_style))

            for i, evidence in enumerate(evidence_list, 1):
                quality_badge = self._get_quality_badge(evidence.quality_category)
                safe_citation = evidence.citation_text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
                try:
                    story.append(Paragraph(f"{i}. {quality_badge} {safe_citation[:200]}", body_style))
                except Exception:
                    story.append(Paragraph(f"{i}. [Citation formatting error]", body_style))

        # Build PDF
        doc.build(story)

        return filepath

    def _generate_html(
        self,
        session: ResearchSession,
        evidence_locker: EvidenceLocker,
        verification_result: VerificationResult,
        filename: str,
        filtered_evidence: List = None,
    ) -> Path:
        """Generate HTML report."""
        plan = session.research_plan

        # Use filtered evidence or all evidence
        evidence_list = filtered_evidence if filtered_evidence is not None else evidence_locker.get_all_evidence()

        # Renumber citations to match final reference order
        renumbered_contents = self._renumber_all_citations(session, evidence_list)

        # Build content sections
        sections_html = []
        if plan:
            for item in plan.table_of_contents.items:
                section_content = renumbered_contents.get(item.section, {})

                section_html = f"""
                <section id="{self._slugify(item.title)}">
                    <h2>{item.section}. {item.title}</h2>
                    <div class="section-content">
                        {self._markdown_to_html(section_content.get('content', ''))}
                    </div>
                """

                # Images
                if self.include_images and section_content.get("images"):
                    section_html += '<div class="figures">'
                    for img in section_content["images"][:3]:
                        section_html += f"""
                        <figure>
                            <img src="{img.get('src', '')}" alt="{img.get('suggested_caption', '')}" loading="lazy">
                            <figcaption>{img.get('suggested_caption', '')}</figcaption>
                        </figure>
                        """
                    section_html += '</div>'

                # Subsections
                for sub in item.subsections:
                    sub_content = renumbered_contents.get(sub.section, {})
                    section_html += f"""
                    <section id="{self._slugify(sub.title)}">
                        <h3>{sub.section}. {sub.title}</h3>
                        <div class="section-content">
                            {self._markdown_to_html(sub_content.get('content', ''))}
                        </div>
                    </section>
                    """

                section_html += "</section>"
                sections_html.append(section_html)

        # Executive summary
        exec_summary = renumbered_contents.get("_executive_summary", {})
        exec_html = ""
        if exec_summary:
            findings_html = "".join(
                f"<li>{f}</li>" for f in exec_summary.get("key_findings", [])
            )
            exec_html = f"""
            <section class="executive-summary">
                <h2>Executive Summary</h2>
                <p>{exec_summary.get('executive_summary', '')}</p>
                <h3>Key Findings</h3>
                <ul>{findings_html}</ul>
            </section>
            """

        # Verification summary
        verification_html = ""
        if verification_result:
            verification_html = f"""
            <section class="verification-summary">
                <h2>Verification Summary</h2>
                <div class="reliability-score">
                    <span class="score">{verification_result.overall_reliability_score:.1%}</span>
                    <span class="label">Overall Reliability</span>
                </div>
                <table class="verification-table">
                    <tr><th>Confidence Level</th><th>Count</th></tr>
                    <tr><td>High</td><td>{verification_result.high_confidence_count}</td></tr>
                    <tr><td>Medium</td><td>{verification_result.medium_confidence_count}</td></tr>
                    <tr><td>Low</td><td>{verification_result.low_confidence_count}</td></tr>
                    <tr><td>Unsupported</td><td>{verification_result.unsupported_count}</td></tr>
                </table>
                {f'<p class="warning">Hallucination Risks: {verification_result.hallucination_risk_count} claims flagged</p>' if verification_result.hallucination_risk_count > 0 else ''}
            </section>
            """

        # Quality Summary
        quality_html = ""
        if self.include_quality_summary:
            quality_stats = evidence_locker.get_quality_statistics()
            quality_rows = "".join(
                f"<tr><td>{cat.title()}</td><td>{count}</td></tr>"
                for cat, count in quality_stats.get("quality_distribution", {}).items()
            )
            filter_note = ""
            if filtered_evidence is not None:
                total_all = len(evidence_locker.get_all_evidence())
                filter_note = f'<p class="filter-note">This report uses {len(evidence_list)}/{total_all} evidence items after quality filtering.</p>'

            quality_html = f"""
            <section class="quality-summary">
                <h2>Source Quality Summary</h2>
                <div class="quality-stats">
                    <div class="stat-item">
                        <span class="stat-value">{quality_stats['total_evidence']}</span>
                        <span class="stat-label">Total Evidence</span>
                    </div>
                    <div class="stat-item">
                        <span class="stat-value">{quality_stats['high_quality_percentage']}%</span>
                        <span class="stat-label">High Quality</span>
                    </div>
                    <div class="stat-item">
                        <span class="stat-value">{quality_stats['average_quality_score']}</span>
                        <span class="stat-label">Avg. Score</span>
                    </div>
                </div>
                <table class="quality-table">
                    <tr><th>Category</th><th>Count</th></tr>
                    {quality_rows}
                </table>
                {filter_note}
            </section>
            """

        # References
        references_html = ""
        if self.include_citations:
            refs = "".join(
                f'<li><span class="quality-badge quality-{e.quality_category.value}">{self._get_quality_badge(e.quality_category)}</span> {e.citation_text}</li>'
                for e in evidence_list
            )
            references_html = f"""
            <section class="references">
                <h2>References</h2>
                <ol>{refs}</ol>
                <div class="quality-legend">
                    <span class="legend-item"><span class="quality-badge quality-authoritative">[A]</span> Authoritative</span>
                    <span class="legend-item"><span class="quality-badge quality-high">[H]</span> High</span>
                    <span class="legend-item"><span class="quality-badge quality-medium">[M]</span> Medium</span>
                    <span class="legend-item"><span class="quality-badge quality-low">[L]</span> Low</span>
                    <span class="legend-item"><span class="quality-badge quality-unverified">[?]</span> Unverified</span>
                </div>
            </section>
            """

        # TOC
        toc_html = ""
        if self.include_toc and plan:
            toc_items = "".join(
                f'<li><a href="#{self._slugify(item.title)}">{item.section}. {item.title}</a></li>'
                for item in plan.table_of_contents.items
            )
            toc_html = f'<nav class="toc"><h2>Table of Contents</h2><ol>{toc_items}</ol></nav>'

        html_content = f"""
<!DOCTYPE html>
<html lang="{self.language}">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{plan.title if plan else session.query}</title>
    <style>
        :root {{
            --primary-color: #2c3e50;
            --accent-color: #3498db;
            --text-color: #333;
            --bg-color: #fff;
            --border-color: #e1e1e1;
        }}
        * {{ box-sizing: border-box; margin: 0; padding: 0; }}
        body {{
            font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
            line-height: 1.8;
            color: var(--text-color);
            background: var(--bg-color);
            max-width: 900px;
            margin: 0 auto;
            padding: 40px 20px;
        }}
        header {{
            text-align: center;
            margin-bottom: 40px;
            padding-bottom: 20px;
            border-bottom: 2px solid var(--accent-color);
        }}
        h1 {{ color: var(--primary-color); font-size: 2.5em; margin-bottom: 10px; }}
        h2 {{ color: var(--primary-color); margin: 30px 0 15px; padding-bottom: 10px; border-bottom: 1px solid var(--border-color); }}
        h3 {{ color: var(--primary-color); margin: 20px 0 10px; }}
        p {{ margin-bottom: 15px; }}
        .meta {{ color: #666; font-size: 0.9em; }}
        .toc {{ background: #f8f9fa; padding: 20px; border-radius: 8px; margin: 30px 0; }}
        .toc ol {{ padding-left: 25px; }}
        .toc li {{ margin: 8px 0; }}
        .toc a {{ color: var(--accent-color); text-decoration: none; }}
        .toc a:hover {{ text-decoration: underline; }}
        .executive-summary {{ background: #e8f4fd; padding: 25px; border-radius: 8px; margin: 30px 0; }}
        .section-content {{ padding: 10px 0; }}
        .figures {{ display: flex; gap: 20px; flex-wrap: wrap; margin: 20px 0; }}
        figure {{ flex: 1; min-width: 200px; text-align: center; }}
        figure img {{ max-width: 100%; border-radius: 8px; box-shadow: 0 2px 10px rgba(0,0,0,0.1); }}
        figcaption {{ font-size: 0.9em; color: #666; margin-top: 8px; }}
        .verification-summary {{ background: #fff3cd; padding: 25px; border-radius: 8px; margin: 30px 0; }}
        .reliability-score {{ text-align: center; margin: 20px 0; }}
        .reliability-score .score {{ font-size: 3em; font-weight: bold; color: var(--primary-color); display: block; }}
        .reliability-score .label {{ font-size: 0.9em; color: #666; }}
        .verification-table {{ width: 100%; border-collapse: collapse; margin: 20px 0; }}
        .verification-table th, .verification-table td {{ padding: 12px; border: 1px solid var(--border-color); text-align: left; }}
        .verification-table th {{ background: var(--primary-color); color: white; }}
        .warning {{ color: #dc3545; font-weight: bold; }}
        .references {{ margin-top: 40px; padding-top: 20px; border-top: 2px solid var(--border-color); }}
        .references ol {{ padding-left: 25px; }}
        .references li {{ margin: 10px 0; font-size: 0.9em; }}
        .quality-summary {{ background: #e8f4fd; padding: 25px; border-radius: 8px; margin: 30px 0; }}
        .quality-stats {{ display: flex; justify-content: space-around; margin: 20px 0; flex-wrap: wrap; }}
        .stat-item {{ text-align: center; padding: 15px; }}
        .stat-value {{ font-size: 2em; font-weight: bold; color: var(--primary-color); display: block; }}
        .stat-label {{ font-size: 0.9em; color: #666; }}
        .quality-table {{ width: 100%; border-collapse: collapse; margin: 20px 0; }}
        .quality-table th, .quality-table td {{ padding: 10px; border: 1px solid var(--border-color); text-align: left; }}
        .quality-table th {{ background: var(--primary-color); color: white; }}
        .filter-note {{ font-style: italic; color: #666; margin-top: 15px; }}
        .quality-badge {{ font-weight: bold; margin-right: 5px; padding: 2px 6px; border-radius: 3px; font-size: 0.8em; }}
        .quality-authoritative {{ background: #28a745; color: white; }}
        .quality-high {{ background: #17a2b8; color: white; }}
        .quality-medium {{ background: #ffc107; color: #333; }}
        .quality-low {{ background: #fd7e14; color: white; }}
        .quality-unverified {{ background: #6c757d; color: white; }}
        .quality-legend {{ margin-top: 20px; padding: 15px; background: #f8f9fa; border-radius: 5px; }}
        .legend-item {{ margin-right: 15px; display: inline-block; }}
        @media print {{
            body {{ max-width: none; }}
            .toc {{ page-break-after: always; }}
            section {{ page-break-inside: avoid; }}
        }}
    </style>
</head>
<body>
    <header>
        <h1>{plan.title if plan else session.query}</h1>
        <p class="meta">Research Date: {session.started_at[:10]} | Session: {session.session_id}</p>
    </header>

    {toc_html}
    {exec_html}
    {''.join(sections_html)}
    {verification_html}
    {quality_html}
    {references_html}

    <footer style="margin-top: 40px; text-align: center; color: #666; font-size: 0.8em;">
        <p>Generated by Deep Research Tool</p>
    </footer>
</body>
</html>
"""

        filepath = self.output_dir / f"{filename}.html"
        with open(filepath, "w", encoding="utf-8") as f:
            f.write(html_content)

        return filepath

    def _slugify(self, text: str) -> str:
        """Convert text to URL-friendly slug."""
        text = text.lower()
        text = re.sub(r'[^\w\s-]', '', text)
        text = re.sub(r'[\s_-]+', '-', text)
        return text.strip('-')

    def _build_url_to_reference_map(
        self,
        evidence_list: List,
    ) -> Dict[str, int]:
        """
        Build a mapping from source URL to final reference number.

        Args:
            evidence_list: List of evidence items in final order

        Returns:
            Dictionary mapping URL to reference number (1-indexed)
        """
        url_to_ref = {}
        for i, evidence in enumerate(evidence_list, 1):
            url_to_ref[evidence.url] = i
        return url_to_ref

    def _renumber_citations_in_section(
        self,
        content: str,
        section_sources: List[str],
        url_to_ref: Dict[str, int],
    ) -> str:
        """
        Renumber citations in section content to match final reference numbers.

        During synthesis, sources are numbered [SOURCE 1], [SOURCE 2], etc.
        based on the order they were passed to the synthesizer.
        This method converts those to match the final reference list order.

        Args:
            content: Section content with [SOURCE N] markers
            section_sources: List of URLs used for this section (in synthesis order)
            url_to_ref: Mapping from URL to final reference number

        Returns:
            Content with corrected reference numbers
        """
        if not content or not section_sources:
            return content

        # Find all [SOURCE N] patterns and replace with correct numbers
        def replace_citation(match):
            original_num = int(match.group(1))
            # Source numbers are 1-indexed
            source_index = original_num - 1

            if 0 <= source_index < len(section_sources):
                source_url = section_sources[source_index]
                if source_url in url_to_ref:
                    new_ref_num = url_to_ref[source_url]
                    return f"[{new_ref_num}]"

            # If we can't map, keep original but mark it
            return f"[?{original_num}]"

        # Replace [SOURCE N] and [SOURCE: N] patterns
        result = re.sub(r'\[SOURCE:?\s*(\d+)\]', replace_citation, content)
        return result

    def _renumber_all_citations(
        self,
        session: 'ResearchSession',
        evidence_list: List,
    ) -> Dict[str, Dict[str, Any]]:
        """
        Renumber all citations in session contents to match final reference order.

        Args:
            session: Research session with section contents
            evidence_list: Final ordered list of evidence

        Returns:
            Updated section_contents with correct citation numbers
        """
        from copy import deepcopy

        url_to_ref = self._build_url_to_reference_map(evidence_list)
        updated_contents = deepcopy(session.section_contents)

        for section_id, section_data in updated_contents.items():
            if section_id.startswith("_"):
                continue

            content = section_data.get("content", "")
            section_sources = section_data.get("sources", [])

            if content and section_sources:
                updated_content = self._renumber_citations_in_section(
                    content, section_sources, url_to_ref
                )
                updated_contents[section_id]["content"] = updated_content

        return updated_contents

    def _process_citations(self, text: str) -> str:
        """Convert citation markers to markdown footnotes."""
        # Convert [N] to [^N] for markdown footnotes
        text = re.sub(r'\[(\d+)\]', r'[^\1]', text)
        # Handle legacy format [SOURCE: N]
        text = re.sub(r'\[SOURCE:\s*(\d+)\]', r'[^\1]', text)
        # Convert [ANALYSIS] to italic marker
        text = text.replace('[ANALYSIS]', '*[Analysis]*')
        return text

    def _strip_citations(self, text: str) -> str:
        """Normalize citation markers for DOCX output.

        Keeps reference numbers as [N] in the text so that they remain
        visible and correspond to the References section at the end.
        Only removes legacy/internal markers that are not user-facing.
        """
        # Keep [N] references — these match the References section numbering.
        # Normalize [SOURCE: N] to [N]
        text = re.sub(r'\[SOURCE:\s*(\d+)\]', r'[\1]', text)
        # Remove unmapped citations [?N]
        text = re.sub(r'\[\?\d+\]', '', text)
        # Convert [ANALYSIS] to visible marker
        text = text.replace('[ANALYSIS]', '(Analysis)')
        return text.strip()

    def _markdown_to_html(self, text: str) -> str:
        """Simple markdown to HTML conversion."""
        if not text:
            return ""

        # Process citations first
        text = self._process_citations(text)

        # Basic markdown conversion
        # Headers (already handled in structure)
        # Bold
        text = re.sub(r'\*\*(.+?)\*\*', r'<strong>\1</strong>', text)
        # Italic
        text = re.sub(r'\*(.+?)\*', r'<em>\1</em>', text)
        # Links
        text = re.sub(r'\[([^\]]+)\]\(([^)]+)\)', r'<a href="\2">\1</a>', text)
        # Line breaks
        text = text.replace('\n\n', '</p><p>')
        text = text.replace('\n', '<br>')

        return f"<p>{text}</p>"
