"""
Semantic manifest - the canonical snapshot of everything the report SAYS.

At freeze time (right after the finalization loop verified the final
body) a manifest of all SEMANTIC content is built and hashed:

- chapter bodies (including limitations, Fermi, glossary, warnings —
  they are chapters by the time the loop runs)
- figure/chart titles, captions and alt texts
- table titles, captions, headers and every cell
- chart insights text

Everything in the manifest was part of what the verifier saw. After the
freeze, renderers may only perform NON-semantic work (citation number
substitution — applied before hashing —, layout, fonts, DOCX/PDF
conversion, drawing pixels from already-frozen chart data), so the
manifest recomputed just before the final output must hash IDENTICALLY.
A mismatch means something generated or altered meaning after
verification — a pipeline defect, reported loudly.
"""

import hashlib
import json
from typing import Any, Dict, List, Optional


def figure_semantics(collection) -> List[Dict[str, Any]]:
    """All semantic content of a FigureTableCollection, canonically ordered."""
    if collection is None:
        return []
    items: List[Dict[str, Any]] = []
    for kind, figs in (("figure", getattr(collection, "figures", []) or []),
                       ("chart", getattr(collection, "charts", []) or [])):
        for f in figs:
            entry = {
                "kind": kind,
                "id": str(getattr(f, "figure_id", "") or ""),
                "section": str(getattr(f, "section_id", "") or ""),
                "title": str(getattr(f, "title", "") or ""),
                "caption": str(getattr(f, "caption", "") or ""),
                "alt": str(getattr(f, "alt_text", "") or ""),
            }
            # chart STRUCTURED data (series values, axes, unit,
            # annotation) is semantic content: a changed plotted number
            # changes the manifest hash even though the pixels are
            # rendered later
            chart_data = getattr(f, "chart_data", None)
            if chart_data:
                def _c(v):
                    if isinstance(v, dict):
                        return {str(k): _c(x) for k, x in sorted(v.items())}
                    if isinstance(v, (list, tuple)):
                        return [_c(x) for x in v]
                    return str(v)
                entry["chart_data"] = _c(chart_data)
            items.append(entry)
    for t in getattr(collection, "tables", []) or []:
        items.append({
            "kind": "table",
            "id": str(getattr(t, "table_id", "") or ""),
            "section": str(getattr(t, "section_id", "") or ""),
            "title": str(getattr(t, "title", "") or ""),
            "caption": str(getattr(t, "caption", "") or ""),
            "unit": str(getattr(t, "unit", "") or ""),
            "headers": [str(h) for h in (getattr(t, "headers", []) or [])],
            "rows": [[str(c) for c in row]
                     for row in (getattr(t, "rows", []) or [])],
        })
    items.sort(key=lambda i: (i["kind"], i["section"], i["id"], i["title"]))
    return items


def build_semantic_manifest(chapters: Dict[str, str],
                            figure_collection=None,
                            extras: Optional[Dict[str, Any]] = None,
                            ) -> Dict[str, Any]:
    """Canonical semantic snapshot: chapter texts + figure semantics.

    ``extras`` carries semantic content that lives OUTSIDE the chapter
    map — the executive summary, key findings, recommendations, glossary
    entries, footnotes. Everything the reader will read is in the
    manifest; nothing semantic escapes the freeze because it happens to
    be stored under another key.
    """
    manifest = {
        "chapters": {str(k): str(v or "") for k, v in
                     sorted((chapters or {}).items())},
        "figures": figure_semantics(figure_collection),
    }
    if extras:
        def _canon(v):
            if isinstance(v, dict):
                return {str(k): _canon(x) for k, x in sorted(v.items())}
            if isinstance(v, (list, tuple)):
                return [_canon(x) for x in v]
            return str(v)
        manifest["extras"] = {str(k): _canon(v)
                              for k, v in sorted(extras.items())}
    return manifest


def manifest_hash(manifest: Dict[str, Any]) -> str:
    """Deterministic sha256 over the canonical JSON form."""
    canonical = json.dumps(manifest, ensure_ascii=False, sort_keys=True,
                           separators=(",", ":"))
    return hashlib.sha256(canonical.encode("utf-8")).hexdigest()


# ---------------------------------------------------------------------------
# Artifact-level freeze verification (audit item D)
#
# Hashing the in-memory chapters twice only proves the OBJECT did not
# change; a renderer bug could still alter or drop content on the way to
# disk. The functions below REBUILD the semantic text from the ACTUAL
# saved artifact and check that every frozen chapter paragraph survived
# into it (normalized: markup, citation numbers and whitespace removed —
# rendering may restyle, it must never change meaning-bearing text).
# ---------------------------------------------------------------------------

import re as _re

# [SOURCE 1] / [SOURCE: 1] editing tags, [1] display numbers and [^1]
# footnote markers are all citation plumbing, not semantic text
_CITATION_RE = _re.compile(r"\[(?:SOURCE:?\s*)?\^?\d+\]")
_MD_IMAGE_RE = _re.compile(r"!\[[^\]]*\]\([^)]*\)")
_MD_LINK_RE = _re.compile(r"\[([^\]]*)\]\([^)]*\)")
_TAG_RE = _re.compile(r"<[^>]+>")


def normalize_semantic_text(text: str) -> str:
    """Reduce text to its meaning-bearing characters.

    Removes markdown/HTML markup, citation tags/numbers and ALL
    whitespace, so a markdown source and its rendered artifact compare
    equal exactly when the words are the same.
    """
    text = text or ""
    text = _MD_IMAGE_RE.sub("", text)
    text = _MD_LINK_RE.sub(r"\1", text)
    text = _TAG_RE.sub("", text)
    text = _CITATION_RE.sub("", text)
    text = _re.sub(r"[*_`>#|\\-]+", "", text)
    text = _re.sub(r"[：:]\s*$", "", text)
    return _re.sub(r"\s+", "", text)


def extract_artifact_text(path) -> Optional[str]:
    """Text of a SAVED report artifact, read back LOCALLY
    (md/txt/html/docx/pdf — PDF via the already-bundled PyMuPDF/pypdf
    extractors, never a network service).

    Returns None when the format cannot be read back (the caller reports
    the check as SKIPPED — never silently passed).
    """
    from pathlib import Path
    p = Path(path)
    if not p.exists():
        return None
    suffix = p.suffix.lower()
    if suffix in (".md", ".txt", ".markdown"):
        return p.read_text(encoding="utf-8", errors="replace")
    if suffix in (".html", ".htm"):
        import html as _htmlmod
        return _htmlmod.unescape(
            p.read_text(encoding="utf-8", errors="replace"))
    if suffix == ".docx":
        try:
            import docx
            doc = docx.Document(str(p))
            parts = [para.text for para in doc.paragraphs]
            for table in doc.tables:
                for row in table.rows:
                    parts.extend(cell.text for cell in row.cells)
            return "\n".join(parts)
        except Exception:
            return None
    if suffix == ".pdf":
        try:
            import fitz                     # PyMuPDF (existing dep)
            with fitz.open(str(p)) as doc:
                return "\n".join(page.get_text() for page in doc)
        except Exception:
            pass
        try:
            import pypdf
            reader = pypdf.PdfReader(str(p))
            return "\n".join((page.extract_text() or "")
                             for page in reader.pages)
        except Exception:
            return None
    return None


def verify_frozen_in_artifact(manifest: Dict[str, Any],
                              artifact_text: str,
                              min_fragment_chars: int = 24,
                              allowed_extras: Optional[List[str]] = None,
                              check_additions: bool = True,
                              ) -> Dict[str, Any]:
    """Two-way, FAIL-CLOSED check between the frozen manifest and the
    ACTUAL saved artifact.

    Direction 1 (missing): every paragraph line of every frozen chapter,
    every figure/table title, caption and CELL (cells are checked
    regardless of length — a short "52%" cell counts), and every extras
    value must appear — normalized — in the artifact.

    Direction 2 (additions): substantial artifact paragraphs that are
    neither frozen content nor a declared deterministic extra
    (references, headings, TOC lines, dates...) are UNVERIFIED
    post-freeze additions and fail the check.

    checked == 0 is itself a FAILURE — an empty comparison can never
    pass. Returns {"ok", "checked", "missing", "additions"}.
    """
    haystack = normalize_semantic_text(artifact_text)
    checked = 0
    missing: List[str] = []
    needles: List[str] = []      # normalized frozen fragments

    def _check(fragment: str, min_chars: int = None) -> None:
        nonlocal checked
        needle = normalize_semantic_text(fragment)
        limit = min_fragment_chars if min_chars is None else min_chars
        if len(needle) < max(1, limit):
            return
        checked += 1
        needles.append(needle)
        if needle not in haystack:
            missing.append(fragment.strip()[:120])

    chapter_titles: List[str] = []
    for text in (manifest.get("chapters") or {}).values():
        for para in (text or "").split("\n\n"):
            for line in para.split("\n"):
                if line.strip().startswith("#"):
                    chapter_titles.append(line)
                _check(line)
    for item in manifest.get("figures") or []:
        _check(item.get("title", ""), min_chars=2)
        _check(item.get("caption", ""), min_chars=2)
        for header in item.get("headers", []) or []:
            _check(header, min_chars=1)          # short cells count
        for row in item.get("rows", []) or []:
            for cell in row:
                _check(cell, min_chars=1)        # short cells count
    extras = manifest.get("extras") or {}

    def _walk(value):
        if isinstance(value, dict):
            for v in value.values():
                _walk(v)
        elif isinstance(value, list):
            for v in value:
                _walk(v)
        else:
            _check(str(value), min_chars=2)
    _walk(extras)

    if checked == 0:
        return {"ok": False, "checked": 0,
                "missing": ["nothing to check (empty manifest)"],
                "additions": []}

    # ---- direction 2: unverified POST-FREEZE additions ----------------
    additions: List[str] = []
    if check_additions:
        frozen_blob = "\x1f".join(needles)
        allowed_norm = [normalize_semantic_text(e)
                        for e in [*(allowed_extras or []), *chapter_titles]]
        allowed_norm = [a for a in allowed_norm if len(a) >= 6]
        for para in _re.split(r"\n\s*\n", artifact_text or ""):
            for line in para.split("\n"):
                # footnote definitions / numbered reference lines are
                # citation plumbing, not semantic additions
                if _re.match(r"\s*\[\^?\d+\]\s*[:：]", line) or \
                        _re.match(r"\s*\d+\.\s+\S*https?://", line):
                    continue
                norm = normalize_semantic_text(line)
                if len(norm) < 40:      # short boilerplate is structural
                    continue
                if norm in frozen_blob or \
                        any(norm in n or n in norm for n in needles):
                    continue
                if any(a in norm or norm in a for a in allowed_norm):
                    continue
                additions.append(line.strip()[:120])

    return {"ok": not missing and not additions, "checked": checked,
            "missing": missing, "additions": additions}


def default_allowed_extras(evidence_locker=None) -> List[str]:
    """Deterministic artifact content that is NOT semantic body:
    references (locker citation texts/titles/urls) and standard
    structural headings."""
    extras: List[str] = []
    if evidence_locker is not None:
        try:
            for ev in evidence_locker.get_all_evidence():
                extras.append(str(getattr(ev, "citation_text", "") or ""))
                extras.append(str(getattr(ev, "title", "") or ""))
                extras.append(str(getattr(ev, "url", "") or ""))
        except Exception:
            pass
    extras.extend([
        "参考文献", "引用文献", "目次", "References", "Sources",
        "Table of Contents", "図表一覧", "Figures and Tables",
        "用語集", "Glossary",
    ])
    return [e for e in extras if e]


def check_artifact(manifest: Dict[str, Any], report_path,
                   evidence_locker=None) -> Dict[str, Any]:
    """One-call artifact verification shared by run()/manual/CLI paths.

    Returns {"status": "pass"|"fail"|"skipped:<reason>",
             "checked", "missing", "additions"}.
    """
    if not manifest:
        return {"status": "skipped:no_manifest", "checked": 0,
                "missing": [], "additions": []}
    artifact_text = extract_artifact_text(report_path)
    if artifact_text is None:
        return {"status": "skipped:unreadable_format", "checked": 0,
                "missing": [], "additions": []}
    result = verify_frozen_in_artifact(
        manifest, artifact_text,
        allowed_extras=default_allowed_extras(evidence_locker))
    result["status"] = "pass" if result["ok"] else "fail"
    return result


def figure_semantics_markdown(collection, language: str = "ja",
                              section_id: str = "figures") -> str:
    """The figure semantics as a verifiable markdown section.

    This text joins the finalization chapters, so every LLM-generated
    figure title, caption and table cell passes through the SAME
    verification as the body — BEFORE the freeze. It is rendered
    verbatim afterwards (render-only section, never LLM-edited).
    """
    items = figure_semantics(collection)
    if not items:
        return ""
    title = "図表一覧" if language == "ja" else "Figures and Tables"
    lines = [f"## {section_id}. {title}", ""]
    for item in items:
        label = {"figure": "図", "chart": "図", "table": "表"}.get(
            item["kind"], "図") if language == "ja" else item["kind"].title()
        head = item["title"] or item["id"]
        lines.append(f"### {label}: {head}")
        if item["caption"]:
            lines.append(item["caption"])
        if item.get("alt"):
            lines.append(item["alt"])
        if item["kind"] == "table" and item.get("headers"):
            lines.append("")
            lines.append("| " + " | ".join(item["headers"]) + " |")
            lines.append("|" + "---|" * len(item["headers"]))
            # EVERY row is part of the verified text — no row cap
            for row in item.get("rows", []):
                lines.append("| " + " | ".join(row) + " |")
        chart_data = item.get("chart_data")
        if chart_data:
            # chart series values/axes/unit join the verified text too
            labels = chart_data.get("labels") or []
            values = chart_data.get("values") or []
            pairs = [f"{l}={v}" for l, v in zip(labels, values)]
            for name, series in (chart_data.get("series") or {}).items():
                pairs.extend(f"{name}:{l}={v}"
                             for l, v in zip(labels, series))
            if pairs:
                head = "データ" if language == "ja" else "Data"
                lines.append(f"{head}: " + "; ".join(pairs))
            axes = " / ".join(x for x in (chart_data.get("x_axis"),
                                          chart_data.get("y_axis"),
                                          chart_data.get("unit")) if x)
            if axes:
                lines.append(("軸・単位: " if language == "ja"
                              else "Axes/unit: ") + axes)
        lines.append("")
    return "\n".join(lines).strip()
