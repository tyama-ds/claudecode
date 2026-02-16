"""
Integration tests for Deep Research Tool.
Tests component interactions without requiring actual API keys.
"""

import json
import tempfile
from pathlib import Path
from unittest.mock import Mock, MagicMock, patch
import sys

sys.path.insert(0, str(Path(__file__).parent.parent.parent))

import pytest

from deep_research_tool.config import Config, create_config, LLMProvider, SearchMethod
from deep_research_tool.evidence import EvidenceLocker, Evidence, EvidenceType
from deep_research_tool.research.query_generator import QueryGenerator, ResearchPlan, TableOfContents
from deep_research_tool.research.content_extractor import ContentExtractor, ExtractedContent
from deep_research_tool.research.researcher import Researcher, ResearchSession, ResearchState
from deep_research_tool.verification import Verifier, VerificationResult, ClaimVerification
from deep_research_tool.report import ReportGenerator, ReportFormat
from deep_research_tool.search import DuckDuckGoSearch, SearchResult
from deep_research_tool.utils import DocumentReader, truncate_text, chunk_text


class TestQueryGenerator:
    """Test QueryGenerator with mocked LLM."""

    def test_create_research_plan_with_mock(self):
        """Test creating research plan with mocked LLM response."""
        mock_llm = Mock()
        mock_llm.generate.return_value = Mock(
            content=json.dumps({
                "title": "Test Research Report",
                "summary": "Test summary",
                "table_of_contents": [
                    {"section": "1", "title": "Test Topic Market Size and Growth", "description": "Market analysis", "subsections": [
                        {"section": "1.1", "title": "Global Market Trends", "description": "Worldwide trends"},
                        {"section": "1.2", "title": "Regional Breakdown", "description": "By region"},
                    ]},
                    {"section": "2", "title": "Test Topic Technical Architecture", "description": "Technical details", "subsections": [
                        {"section": "2.1", "title": "Core Components", "description": "Main parts"},
                        {"section": "2.2", "title": "Integration Patterns", "description": "How parts connect"},
                    ]},
                    {"section": "3", "title": "Test Topic Competitive Landscape", "description": "Key players", "subsections": [
                        {"section": "3.1", "title": "Major Vendors", "description": "Top companies"},
                        {"section": "3.2", "title": "Market Share Distribution", "description": "Share data"},
                    ]},
                ],
                "search_queries": ["query1", "query2", "query3"],
                "key_terms": ["term1", "term2"],
                "suggested_sources": ["source1"],
                "methodology_notes": "Test notes",
                "estimated_complexity": "medium",
            })
        )

        generator = QueryGenerator(mock_llm, language="ja")
        plan = generator.create_research_plan("Test topic")

        assert plan.title == "Test Research Report"
        assert len(plan.table_of_contents.items) == 3
        assert len(plan.search_queries) == 3
        mock_llm.generate.assert_called_once()

    def test_generate_follow_up_queries_with_mock(self):
        """Test generating follow-up queries."""
        mock_llm = Mock()
        mock_llm.generate.return_value = Mock(
            content='["follow-up query 1", "follow-up query 2"]'
        )

        generator = QueryGenerator(mock_llm)
        from deep_research_tool.research.query_generator import TableOfContentsItem
        section = TableOfContentsItem(section="1", title="Test Section", description="Test")

        queries = generator.generate_follow_up_queries(section, "existing info", ["gap1"])

        assert len(queries) == 2
        assert "follow-up query 1" in queries

    def test_identify_gaps_with_mock(self):
        """Test identifying information gaps."""
        mock_llm = Mock()
        mock_llm.generate.return_value = Mock(
            content='["gap 1", "gap 2", "gap 3"]'
        )

        generator = QueryGenerator(mock_llm)
        from deep_research_tool.research.query_generator import TableOfContentsItem
        section = TableOfContentsItem(section="1", title="Test", description="Test")

        gaps = generator.identify_gaps(section, "current content")

        assert len(gaps) == 3


class TestContentExtractor:
    """Test ContentExtractor with mocked LLM."""

    def test_extract_relevant_content_with_mock(self):
        """Test content extraction."""
        mock_llm = Mock()
        mock_llm.generate.return_value = Mock(
            content=json.dumps({
                "processed_content": "Processed test content",
                "key_points": ["point1", "point2"],
                "quotes": [{"text": "quote", "context": "context"}],
                "relevance_score": 0.8,
                "extraction_notes": "Good source",
            })
        )

        extractor = ContentExtractor(mock_llm)
        result = extractor.extract_relevant_content(
            raw_content="Raw test content",
            source_url="https://example.com",
            source_title="Test Source",
            section_context="Test Section",
            research_query="test query",
        )

        assert result.processed_content == "Processed test content"
        assert result.relevance_score == 0.8
        assert len(result.key_points) == 2

    def test_synthesize_section_content_with_mock(self):
        """Test content synthesis."""
        mock_llm = Mock()
        mock_llm.generate.return_value = Mock(
            content=json.dumps({
                "content": "Synthesized content",
                "summary": "Brief summary",
                "source_references": [1, 2],
                "analysis_points": ["analysis"],
                "information_gaps": [],
                "confidence_level": "high",
            })
        )

        extractor = ContentExtractor(mock_llm)
        extracted = [
            ExtractedContent(
                source_url="https://example.com",
                source_title="Test",
                raw_content="Raw",
                processed_content="Processed",
            )
        ]

        result = extractor.synthesize_section_content(
            section_title="Test Section",
            section_description="Description",
            extracted_contents=extracted,
        )

        assert result["content"] == "Synthesized content"
        assert result["confidence_level"] == "high"


class TestVerifier:
    """Test Verifier with mocked LLM."""

    def test_verify_content_with_mock(self):
        """Test content verification."""
        mock_llm = Mock()
        # Mock for claim extraction
        mock_llm.generate.side_effect = [
            Mock(content='["claim 1", "claim 2", "claim 3"]'),  # extract_claims
            Mock(content=json.dumps({  # verify first claim
                "confidence": "high",
                "source_support": ["source1"],
                "reasoning": "Well supported",
                "is_hallucination_risk": False,
            })),
            Mock(content=json.dumps({  # verify second claim
                "confidence": "medium",
                "source_support": [],
                "reasoning": "Partially supported",
                "is_hallucination_risk": False,
            })),
            Mock(content=json.dumps({  # verify third claim
                "confidence": "low",
                "source_support": [],
                "reasoning": "Not well supported",
                "is_hallucination_risk": True,
            })),
        ]

        verifier = Verifier(mock_llm)
        locker = EvidenceLocker()
        locker.add_evidence(url="https://example.com", title="Source 1", content_excerpt="Content")

        result = verifier.verify_content(
            content="Test content with claims",
            evidence_locker=locker,
            document_title="Test Document",
        )

        assert result.total_claims == 3
        assert result.high_confidence_count == 1
        assert result.medium_confidence_count == 1
        assert result.low_confidence_count == 1
        assert result.hallucination_risk_count == 1

    def test_quick_verify_with_mock(self):
        """Test quick verification."""
        mock_llm = Mock()
        mock_llm.generate.return_value = Mock(
            content=json.dumps({
                "overall_reliability": "medium",
                "suspicious_claims": ["suspicious claim"],
                "verification_needed": ["item to verify"],
                "recommendations": ["recommendation"],
            })
        )

        verifier = Verifier(mock_llm)
        result = verifier.quick_verify("Content to verify")

        assert result["overall_reliability"] == "medium"
        assert len(result["suspicious_claims"]) == 1


class TestReportGenerator:
    """Test ReportGenerator."""

    def test_generate_markdown_report(self):
        """Test generating Markdown report."""
        with tempfile.TemporaryDirectory() as tmpdir:
            # Create mock session
            session = ResearchSession(
                query="Test research",
                requirements="Test requirements",
            )
            session.research_plan = ResearchPlan(
                title="Test Report",
                summary="Test summary",
                table_of_contents=TableOfContents(
                    title="Test Report",
                    items=[],
                ),
            )
            session.section_contents = {
                "_executive_summary": {
                    "executive_summary": "Executive summary text",
                    "key_findings": ["finding1", "finding2"],
                }
            }

            locker = EvidenceLocker(output_dir=Path(tmpdir) / "evidence")
            locker.add_evidence(
                url="https://example.com",
                title="Test Source",
                content_excerpt="Content",
            )

            generator = ReportGenerator(output_dir=Path(tmpdir) / "reports")
            report_path = generator.generate_report(
                session=session,
                evidence_locker=locker,
                format=ReportFormat.MARKDOWN,
            )

            assert report_path.exists()
            content = report_path.read_text()
            assert "Test Report" in content
            assert "Executive Summary" in content

    def test_generate_html_report(self):
        """Test generating HTML report."""
        with tempfile.TemporaryDirectory() as tmpdir:
            session = ResearchSession(query="Test")
            session.research_plan = ResearchPlan(
                title="Test HTML",
                summary="Summary",
                table_of_contents=TableOfContents(title="Test", items=[]),
            )
            session.section_contents = {}

            locker = EvidenceLocker(output_dir=Path(tmpdir) / "evidence")

            generator = ReportGenerator(output_dir=Path(tmpdir) / "reports")
            report_path = generator.generate_report(
                session=session,
                evidence_locker=locker,
                format=ReportFormat.HTML,
            )

            assert report_path.exists()
            assert report_path.suffix == ".html"

    def test_generate_pdf_report(self):
        """Test generating PDF report."""
        with tempfile.TemporaryDirectory() as tmpdir:
            from deep_research_tool.research.query_generator import TableOfContentsItem

            session = ResearchSession(
                query="Test PDF research",
                requirements="Test requirements",
            )
            session.research_plan = ResearchPlan(
                title="Test PDF Report",
                summary="Test summary for PDF",
                table_of_contents=TableOfContents(
                    title="Test PDF Report",
                    items=[
                        TableOfContentsItem(
                            section="1",
                            title="Introduction",
                            description="Test intro",
                            subsections=[],
                        ),
                    ],
                ),
            )
            session.section_contents = {
                "_executive_summary": {
                    "executive_summary": "This is the executive summary for PDF.",
                    "key_findings": ["Finding 1", "Finding 2", "Finding 3"],
                },
                "1": {
                    "title": "Introduction",
                    "content": "This is the introduction content for the PDF report.",
                    "summary": "Brief intro summary",
                },
            }

            locker = EvidenceLocker(output_dir=Path(tmpdir) / "evidence")
            locker.add_evidence(
                url="https://example.com/pdf-source",
                title="PDF Test Source",
                content_excerpt="Content for PDF citation",
            )

            generator = ReportGenerator(output_dir=Path(tmpdir) / "reports")
            report_path = generator.generate_report(
                session=session,
                evidence_locker=locker,
                format=ReportFormat.PDF,
            )

            assert report_path.exists()
            assert report_path.suffix == ".pdf"
            # Check file is not empty
            assert report_path.stat().st_size > 0

    def test_generate_docx_report(self):
        """Test generating Word document report."""
        with tempfile.TemporaryDirectory() as tmpdir:
            from deep_research_tool.research.query_generator import TableOfContentsItem

            session = ResearchSession(
                query="Test DOCX research",
                requirements="Test requirements",
            )
            session.research_plan = ResearchPlan(
                title="Test Word Document",
                summary="Test summary for DOCX",
                table_of_contents=TableOfContents(
                    title="Test Word Document",
                    items=[
                        TableOfContentsItem(
                            section="1",
                            title="Overview",
                            description="Test overview",
                            subsections=[
                                TableOfContentsItem(
                                    section="1.1",
                                    title="Background",
                                    description="Background info",
                                ),
                            ],
                        ),
                    ],
                ),
            )
            session.section_contents = {
                "_executive_summary": {
                    "executive_summary": "This is the executive summary for Word document.",
                    "key_findings": ["Key finding A", "Key finding B"],
                },
                "1": {
                    "title": "Overview",
                    "content": "This is the overview content for the Word document report.\n\nSecond paragraph here.",
                    "summary": "Brief overview",
                },
                "1.1": {
                    "title": "Background",
                    "content": "Background content with details.",
                    "summary": "Background summary",
                },
            }

            locker = EvidenceLocker(output_dir=Path(tmpdir) / "evidence")
            locker.add_evidence(
                url="https://example.com/docx-source-1",
                title="Word Test Source 1",
                content_excerpt="Content for citation 1",
            )
            locker.add_evidence(
                url="https://example.com/docx-source-2",
                title="Word Test Source 2",
                content_excerpt="Content for citation 2",
            )

            generator = ReportGenerator(output_dir=Path(tmpdir) / "reports")
            report_path = generator.generate_report(
                session=session,
                evidence_locker=locker,
                format=ReportFormat.DOCX,
            )

            assert report_path.exists()
            assert report_path.suffix == ".docx"
            # Check file is not empty
            assert report_path.stat().st_size > 0

            # Verify DOCX content using python-docx
            from docx import Document
            doc = Document(report_path)
            full_text = "\n".join([p.text for p in doc.paragraphs])
            assert "Test Word Document" in full_text or len(doc.paragraphs) > 0

    def test_generate_pdf_with_verification(self):
        """Test generating PDF with verification results."""
        with tempfile.TemporaryDirectory() as tmpdir:
            session = ResearchSession(query="Verified PDF")
            session.research_plan = ResearchPlan(
                title="Verified PDF Report",
                summary="Summary",
                table_of_contents=TableOfContents(title="Verified", items=[]),
            )
            session.section_contents = {}

            locker = EvidenceLocker(output_dir=Path(tmpdir) / "evidence")

            # Create verification result
            verification = VerificationResult(
                document_title="Verified PDF Report",
                total_claims=10,
                high_confidence_count=6,
                medium_confidence_count=2,
                low_confidence_count=1,
                unsupported_count=1,
                overall_reliability_score=0.8,
            )

            generator = ReportGenerator(output_dir=Path(tmpdir) / "reports")
            report_path = generator.generate_report(
                session=session,
                evidence_locker=locker,
                format=ReportFormat.PDF,
                verification_result=verification,
            )

            assert report_path.exists()
            assert report_path.suffix == ".pdf"
            assert report_path.stat().st_size > 0

    def test_generate_docx_with_quality_filtering(self):
        """Test generating DOCX with quality filtering."""
        with tempfile.TemporaryDirectory() as tmpdir:
            from deep_research_tool.evidence.locker import QualityCategory, SourceType

            session = ResearchSession(query="Quality Filtered DOCX")
            session.research_plan = ResearchPlan(
                title="Quality Filtered Report",
                summary="Summary",
                table_of_contents=TableOfContents(title="Filtered", items=[]),
            )
            session.section_contents = {}

            locker = EvidenceLocker(output_dir=Path(tmpdir) / "evidence")
            # Add evidence with different quality levels
            locker.add_evidence(
                url="https://gov.example.com/report",
                title="Government Report",
                content_excerpt="High quality content",
                quality_category=QualityCategory.AUTHORITATIVE,
                source_type=SourceType.OFFICIAL,
            )
            locker.add_evidence(
                url="https://blog.example.com/post",
                title="Blog Post",
                content_excerpt="Lower quality content",
                quality_category=QualityCategory.LOW,
                source_type=SourceType.BLOG,
            )

            generator = ReportGenerator(output_dir=Path(tmpdir) / "reports")
            report_path = generator.generate_report(
                session=session,
                evidence_locker=locker,
                format=ReportFormat.DOCX,
                min_quality=QualityCategory.HIGH,
            )

            assert report_path.exists()
            assert report_path.suffix == ".docx"

    def test_all_formats_generation(self):
        """Test generating reports in all supported formats."""
        with tempfile.TemporaryDirectory() as tmpdir:
            session = ResearchSession(query="All Formats Test")
            session.research_plan = ResearchPlan(
                title="Multi-Format Report",
                summary="Testing all formats",
                table_of_contents=TableOfContents(title="Multi", items=[]),
            )
            session.section_contents = {
                "_executive_summary": {
                    "executive_summary": "Test summary",
                    "key_findings": ["Finding"],
                }
            }

            locker = EvidenceLocker(output_dir=Path(tmpdir) / "evidence")
            locker.add_evidence(
                url="https://example.com/source",
                title="Test Source",
                content_excerpt="Test content",
            )

            generator = ReportGenerator(output_dir=Path(tmpdir) / "reports")

            formats_to_test = [
                (ReportFormat.MARKDOWN, ".md"),
                (ReportFormat.HTML, ".html"),
                (ReportFormat.PDF, ".pdf"),
                (ReportFormat.DOCX, ".docx"),
            ]

            for fmt, expected_suffix in formats_to_test:
                report_path = generator.generate_report(
                    session=session,
                    evidence_locker=locker,
                    format=fmt,
                    filename=f"test_{fmt.value}",
                )
                assert report_path.exists(), f"Failed to generate {fmt.value}"
                assert report_path.suffix == expected_suffix, f"Wrong suffix for {fmt.value}"
                assert report_path.stat().st_size > 0, f"Empty file for {fmt.value}"


class TestDuckDuckGoSearch:
    """Test DuckDuckGo search with mocked API."""

    def test_search_with_mock(self):
        """Test search functionality with mocked DDGS."""
        search = DuckDuckGoSearch()

        # Mock the internal _ddgs object
        mock_ddgs = Mock()
        mock_ddgs.text.return_value = [
            {"title": "Result 1", "href": "https://example1.com", "body": "Snippet 1"},
            {"title": "Result 2", "href": "https://example2.com", "body": "Snippet 2"},
        ]
        search._ddgs = mock_ddgs

        results = search.search("test query")

        assert len(results) == 2
        assert results[0].title == "Result 1"
        assert results[0].url == "https://example1.com"


class TestDocumentReader:
    """Test DocumentReader."""

    def test_read_text_file(self):
        """Test reading text file."""
        with tempfile.NamedTemporaryFile(mode="w", suffix=".txt", delete=False) as f:
            f.write("Test content\nLine 2")
            f.flush()
            filepath = Path(f.name)

        try:
            reader = DocumentReader()
            doc = reader.read_document(filepath)

            assert doc.content == "Test content\nLine 2"
            assert doc.file_type == "text"
            assert doc.error is None
        finally:
            filepath.unlink()

    def test_read_markdown_file(self):
        """Test reading markdown file."""
        with tempfile.NamedTemporaryFile(mode="w", suffix=".md", delete=False) as f:
            f.write("# Test Title\n\nContent here")
            f.flush()
            filepath = Path(f.name)

        try:
            reader = DocumentReader()
            doc = reader.read_document(filepath)

            assert doc.file_type == "markdown"
            assert doc.title == "Test Title"
        finally:
            filepath.unlink()

    def test_unsupported_file_type(self):
        """Test handling unsupported file type."""
        with tempfile.NamedTemporaryFile(suffix=".xyz", delete=False) as f:
            filepath = Path(f.name)

        try:
            reader = DocumentReader()
            doc = reader.read_document(filepath)

            assert doc.error is not None
            assert "Unsupported" in doc.error
        finally:
            filepath.unlink()


class TestUtilityFunctions:
    """Test utility functions."""

    def test_truncate_text(self):
        """Test text truncation."""
        text = "This is a long text that needs to be truncated"
        result = truncate_text(text, max_length=20)

        assert len(result) <= 20
        assert result.endswith("...")

    def test_truncate_text_short(self):
        """Test truncation of already short text."""
        text = "Short"
        result = truncate_text(text, max_length=20)

        assert result == "Short"

    def test_chunk_text(self):
        """Test text chunking."""
        text = "Paragraph one.\n\nParagraph two.\n\nParagraph three."
        chunks = chunk_text(text, chunk_size=30, overlap=5)

        assert len(chunks) >= 1
        # Verify overlap behavior
        for chunk in chunks:
            assert len(chunk) > 0


class TestResearchSession:
    """Test ResearchSession serialization."""

    def test_session_save_and_load(self):
        """Test saving and loading research session."""
        with tempfile.TemporaryDirectory() as tmpdir:
            session = ResearchSession(
                query="Test query",
                requirements="Test requirements",
            )
            session.section_contents = {"1": {"content": "Test content"}}

            filepath = Path(tmpdir) / "session.json"
            session.save(filepath)

            assert filepath.exists()

            loaded = ResearchSession.load(filepath)

            assert loaded.query == "Test query"
            assert loaded.requirements == "Test requirements"
            assert "1" in loaded.section_contents


class TestVerificationHTMLReport:
    """Test verification HTML report generation."""

    def test_generate_html_report(self):
        """Test HTML report generation."""
        with tempfile.TemporaryDirectory() as tmpdir:
            mock_llm = Mock()
            verifier = Verifier(mock_llm)

            result = VerificationResult(
                document_title="Test Document",
                total_claims=10,
                high_confidence_count=5,
                medium_confidence_count=3,
                low_confidence_count=1,
                unsupported_count=1,
                hallucination_risk_count=2,
                overall_reliability_score=0.75,
            )
            result.verified_claims = [
                ClaimVerification(
                    claim_text="Test claim",
                    confidence=Mock(value="high"),
                    source_support=["source1"],
                    reasoning="Well supported",
                    is_hallucination_risk=False,
                )
            ]

            output_path = Path(tmpdir) / "verification.html"
            html = verifier.generate_verification_report_html(result, output_path)

            assert output_path.exists()
            assert "Test Document" in html
            assert "75.0%" in html


if __name__ == "__main__":
    pytest.main([__file__, "-v"])
