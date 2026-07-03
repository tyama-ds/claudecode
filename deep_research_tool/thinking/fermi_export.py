"""
Word/PDF export for Fermi estimation results.

Requires optional dependencies:
    pip install python-docx   # for save_fermi_docx
    pip install reportlab     # for save_fermi_pdf
"""

import logging
from pathlib import Path
from typing import Union, TYPE_CHECKING

from .fermi_estimator import format_number

if TYPE_CHECKING:
    from .fermi_estimator import FermiEstimate


logger = logging.getLogger(__name__)


def save_fermi_docx(estimate: "FermiEstimate", filepath: Union[str, Path]) -> Path:
    """
    Save a Fermi estimate as a Word (docx) document.

    Args:
        estimate: FermiEstimate to export
        filepath: Output path (.docx appended if missing)

    Returns:
        Path to the saved document

    Raises:
        ImportError: If python-docx is not installed
    """
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        raise ImportError(
            "python-docx not installed. Install with: pip install python-docx"
        )

    filepath = Path(filepath)
    if filepath.suffix.lower() != ".docx":
        filepath = filepath.with_suffix(".docx")

    doc = Document()

    # Title
    title = doc.add_heading(f"フェルミ推定: {estimate.question}", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Estimate summary
    doc.add_heading("推定結果", level=1)
    summary = doc.add_paragraph()
    summary.add_run(
        f"推定値: {format_number(estimate.value)} {estimate.unit}\n"
    ).bold = True
    summary.add_run(
        f"範囲: {format_number(estimate.low)} 〜 {format_number(estimate.high)} "
        f"{estimate.unit}\n"
        f"確信度: {estimate.confidence:.0%}"
    )

    if estimate.formula:
        doc.add_heading("計算式", level=1)
        doc.add_paragraph(estimate.formula)

    # Factors table
    doc.add_heading("分解した要素", level=1)
    table = doc.add_table(rows=1, cols=6)
    table.style = "Table Grid"
    headers = ["要素", "中央値", "範囲", "単位", "演算", "根拠"]
    for cell, header in zip(table.rows[0].cells, headers):
        cell.text = header
    for f in estimate.factors:
        row = table.add_row().cells
        row[0].text = f.name
        row[1].text = format_number(f.mid)
        row[2].text = f"{format_number(f.low)}〜{format_number(f.high)}"
        row[3].text = f.unit
        row[4].text = "×" if f.operation == "multiply" else "÷"
        row[5].text = f.basis

    # Assumptions
    if estimate.assumptions:
        doc.add_heading("前提条件", level=1)
        for assumption in estimate.assumptions:
            doc.add_paragraph(assumption, style="List Bullet")

    # Reasoning
    if estimate.reasoning:
        doc.add_heading("推論過程", level=1)
        doc.add_paragraph(estimate.reasoning)

    doc.save(filepath)
    logger.info(f"Fermi estimate saved to {filepath}")
    return filepath


def save_fermi_pdf(estimate: "FermiEstimate", filepath: Union[str, Path]) -> Path:
    """
    Save a Fermi estimate as a PDF document (Japanese font supported).

    Args:
        estimate: FermiEstimate to export
        filepath: Output path (.pdf appended if missing)

    Returns:
        Path to the saved document

    Raises:
        ImportError: If reportlab is not installed
    """
    try:
        from reportlab.lib import colors
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.platypus import (
            SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
        )
    except ImportError:
        raise ImportError(
            "reportlab not installed. Install with: pip install reportlab"
        )

    from deep_research_tool.utils.fonts import register_japanese_font

    filepath = Path(filepath)
    if filepath.suffix.lower() != ".pdf":
        filepath = filepath.with_suffix(".pdf")

    japanese_font = register_japanese_font()
    base_font = japanese_font if japanese_font else "Helvetica"

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "FermiTitle", parent=styles["Heading1"],
        fontName=base_font, fontSize=18, spaceAfter=20, alignment=1,
    )
    heading_style = ParagraphStyle(
        "FermiHeading", parent=styles["Heading2"],
        fontName=base_font, fontSize=14, spaceBefore=16, spaceAfter=8,
    )
    body_style = ParagraphStyle(
        "FermiBody", parent=styles["Normal"],
        fontName=base_font, fontSize=10.5, leading=16,
    )

    doc = SimpleDocTemplate(str(filepath), pagesize=A4)
    story = [
        Paragraph(f"フェルミ推定: {estimate.question}", title_style),
        Paragraph("推定結果", heading_style),
        Paragraph(
            f"<b>推定値: {format_number(estimate.value)} {estimate.unit}</b><br/>"
            f"範囲: {format_number(estimate.low)} 〜 {format_number(estimate.high)} "
            f"{estimate.unit}<br/>"
            f"確信度: {estimate.confidence:.0%}",
            body_style,
        ),
    ]

    if estimate.formula:
        story.append(Paragraph("計算式", heading_style))
        story.append(Paragraph(estimate.formula, body_style))

    # Factors table
    story.append(Paragraph("分解した要素", heading_style))
    table_data = [["要素", "中央値", "範囲", "単位", "演算", "根拠"]]
    for f in estimate.factors:
        table_data.append([
            Paragraph(f.name, body_style),
            format_number(f.mid),
            f"{format_number(f.low)}〜{format_number(f.high)}",
            f.unit,
            "×" if f.operation == "multiply" else "÷",
            f.basis,
        ])
    table = Table(table_data)
    table.setStyle(TableStyle([
        ("FONTNAME", (0, 0), (-1, -1), base_font),
        ("FONTSIZE", (0, 0), (-1, -1), 9),
        ("BACKGROUND", (0, 0), (-1, 0), colors.lightgrey),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
    ]))
    story.append(table)

    if estimate.assumptions:
        story.append(Paragraph("前提条件", heading_style))
        for assumption in estimate.assumptions:
            story.append(Paragraph(f"・{assumption}", body_style))

    if estimate.reasoning:
        story.append(Paragraph("推論過程", heading_style))
        story.append(Paragraph(estimate.reasoning, body_style))
        story.append(Spacer(1, 12))

    doc.build(story)
    logger.info(f"Fermi estimate saved to {filepath}")
    return filepath
