"""
Document Reader - Read and extract content from various document formats.

Supports:
- PDF files
- Word documents (.docx)
- Text files (.txt)
- Markdown files (.md)
- HTML files (.html)
"""

from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path
from typing import List, Optional, Dict, Any


@dataclass
class DocumentContent:
    """Extracted content from a document."""
    filepath: str
    filename: str
    file_type: str
    title: str = ""
    content: str = ""
    metadata: Dict[str, Any] = field(default_factory=dict)
    pages: int = 0
    images: List[Dict[str, Any]] = field(default_factory=list)
    extracted_at: str = field(default_factory=lambda: datetime.now().isoformat())
    error: Optional[str] = None

    def to_dict(self) -> Dict[str, Any]:
        """Convert to dictionary."""
        return {
            "filepath": self.filepath,
            "filename": self.filename,
            "file_type": self.file_type,
            "title": self.title,
            "content": self.content,
            "metadata": self.metadata,
            "pages": self.pages,
            "images": self.images,
            "extracted_at": self.extracted_at,
            "error": self.error,
        }


class DocumentReader:
    """
    Read and extract content from various document formats.

    Automatically detects file type and uses appropriate extraction method.
    """

    SUPPORTED_TYPES = {
        ".pdf": "pdf",
        ".docx": "docx",
        ".doc": "doc",
        ".txt": "text",
        ".md": "markdown",
        ".html": "html",
        ".htm": "html",
    }

    def __init__(self, extract_images: bool = True, max_pages: int = None):
        """
        Initialize DocumentReader.

        Args:
            extract_images: Whether to extract images from documents
            max_pages: Maximum pages to process (None for all)
        """
        self.extract_images = extract_images
        self.max_pages = max_pages

    def read_document(self, filepath: Path) -> DocumentContent:
        """
        Read a document and extract its content.

        Args:
            filepath: Path to the document

        Returns:
            DocumentContent with extracted text and metadata
        """
        filepath = Path(filepath)

        if not filepath.exists():
            return DocumentContent(
                filepath=str(filepath),
                filename=filepath.name,
                file_type="unknown",
                error=f"File not found: {filepath}",
            )

        file_type = self.SUPPORTED_TYPES.get(filepath.suffix.lower(), "unknown")

        if file_type == "pdf":
            return self._read_pdf(filepath)
        elif file_type == "docx":
            return self._read_docx(filepath)
        elif file_type == "text":
            return self._read_text(filepath)
        elif file_type == "markdown":
            return self._read_markdown(filepath)
        elif file_type == "html":
            return self._read_html(filepath)
        else:
            return DocumentContent(
                filepath=str(filepath),
                filename=filepath.name,
                file_type=file_type,
                error=f"Unsupported file type: {filepath.suffix}",
            )

    def read_documents(self, filepaths: List[Path]) -> List[DocumentContent]:
        """
        Read multiple documents.

        Args:
            filepaths: List of document paths

        Returns:
            List of DocumentContent objects
        """
        return [self.read_document(Path(fp)) for fp in filepaths]

    def _read_pdf(self, filepath: Path) -> DocumentContent:
        """Read PDF file using PyPDF or PyMuPDF."""
        try:
            # Try PyMuPDF first (better quality extraction)
            return self._read_pdf_pymupdf(filepath)
        except ImportError:
            try:
                # Fall back to pypdf
                return self._read_pdf_pypdf(filepath)
            except ImportError:
                return DocumentContent(
                    filepath=str(filepath),
                    filename=filepath.name,
                    file_type="pdf",
                    error="No PDF library installed. Install with: pip install PyMuPDF or pip install pypdf",
                )

    def _read_pdf_pymupdf(self, filepath: Path) -> DocumentContent:
        """Read PDF using PyMuPDF (fitz)."""
        import fitz

        doc = fitz.open(filepath)
        content_parts = []
        images = []

        max_pages = self.max_pages or len(doc)

        for page_num in range(min(len(doc), max_pages)):
            page = doc[page_num]

            # Extract text
            text = page.get_text()
            content_parts.append(text)

            # Extract images if enabled
            if self.extract_images:
                image_list = page.get_images()
                for img_idx, img in enumerate(image_list[:5]):  # Limit images per page
                    xref = img[0]
                    try:
                        base_image = doc.extract_image(xref)
                        images.append({
                            "page": page_num + 1,
                            "index": img_idx,
                            "width": base_image.get("width", 0),
                            "height": base_image.get("height", 0),
                            "format": base_image.get("ext", ""),
                        })
                    except Exception:
                        pass

        # Get metadata
        metadata = doc.metadata or {}

        doc.close()

        return DocumentContent(
            filepath=str(filepath),
            filename=filepath.name,
            file_type="pdf",
            title=metadata.get("title", filepath.stem),
            content="\n\n".join(content_parts),
            metadata=metadata,
            pages=len(doc),
            images=images,
        )

    def _read_pdf_pypdf(self, filepath: Path) -> DocumentContent:
        """Read PDF using pypdf."""
        from pypdf import PdfReader

        reader = PdfReader(filepath)
        content_parts = []

        max_pages = self.max_pages or len(reader.pages)

        for page_num in range(min(len(reader.pages), max_pages)):
            page = reader.pages[page_num]
            text = page.extract_text()
            content_parts.append(text)

        # Get metadata
        metadata = {}
        if reader.metadata:
            metadata = {
                "title": reader.metadata.get("/Title", ""),
                "author": reader.metadata.get("/Author", ""),
                "creator": reader.metadata.get("/Creator", ""),
                "producer": reader.metadata.get("/Producer", ""),
            }

        return DocumentContent(
            filepath=str(filepath),
            filename=filepath.name,
            file_type="pdf",
            title=metadata.get("title", filepath.stem),
            content="\n\n".join(content_parts),
            metadata=metadata,
            pages=len(reader.pages),
        )

    def _read_docx(self, filepath: Path) -> DocumentContent:
        """Read Word document."""
        try:
            from docx import Document
        except ImportError:
            return DocumentContent(
                filepath=str(filepath),
                filename=filepath.name,
                file_type="docx",
                error="python-docx not installed. Install with: pip install python-docx",
            )

        doc = Document(filepath)
        content_parts = []

        # Extract paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                content_parts.append(para.text)

        # Extract tables
        for table in doc.tables:
            table_text = []
            for row in table.rows:
                row_text = [cell.text for cell in row.cells]
                table_text.append(" | ".join(row_text))
            if table_text:
                content_parts.append("\n".join(table_text))

        # Get core properties
        metadata = {}
        try:
            core_props = doc.core_properties
            metadata = {
                "title": core_props.title or "",
                "author": core_props.author or "",
                "created": str(core_props.created) if core_props.created else "",
                "modified": str(core_props.modified) if core_props.modified else "",
            }
        except Exception:
            pass

        return DocumentContent(
            filepath=str(filepath),
            filename=filepath.name,
            file_type="docx",
            title=metadata.get("title", filepath.stem),
            content="\n\n".join(content_parts),
            metadata=metadata,
        )

    def _read_text(self, filepath: Path) -> DocumentContent:
        """Read plain text file."""
        try:
            # Try multiple encodings
            encodings = ['utf-8', 'utf-8-sig', 'shift_jis', 'euc-jp', 'latin-1']

            content = None
            for encoding in encodings:
                try:
                    with open(filepath, 'r', encoding=encoding) as f:
                        content = f.read()
                    break
                except UnicodeDecodeError:
                    continue

            if content is None:
                return DocumentContent(
                    filepath=str(filepath),
                    filename=filepath.name,
                    file_type="text",
                    error="Unable to decode file with supported encodings",
                )

            return DocumentContent(
                filepath=str(filepath),
                filename=filepath.name,
                file_type="text",
                title=filepath.stem,
                content=content,
            )

        except Exception as e:
            return DocumentContent(
                filepath=str(filepath),
                filename=filepath.name,
                file_type="text",
                error=str(e),
            )

    def _read_markdown(self, filepath: Path) -> DocumentContent:
        """Read Markdown file."""
        result = self._read_text(filepath)
        result.file_type = "markdown"

        # Try to extract title from first heading
        if result.content and not result.error:
            lines = result.content.split('\n')
            for line in lines:
                if line.startswith('# '):
                    result.title = line[2:].strip()
                    break

        return result

    def _read_html(self, filepath: Path) -> DocumentContent:
        """Read HTML file and extract text content."""
        try:
            from bs4 import BeautifulSoup
        except ImportError:
            return DocumentContent(
                filepath=str(filepath),
                filename=filepath.name,
                file_type="html",
                error="beautifulsoup4 not installed. Install with: pip install beautifulsoup4",
            )

        result = self._read_text(filepath)
        if result.error:
            return result

        try:
            soup = BeautifulSoup(result.content, 'lxml')

            # Extract title
            title = ""
            if soup.title:
                title = soup.title.string or ""

            # Remove script and style elements
            for element in soup(["script", "style"]):
                element.decompose()

            # Get text
            text = soup.get_text(separator='\n', strip=True)

            return DocumentContent(
                filepath=str(filepath),
                filename=filepath.name,
                file_type="html",
                title=title or filepath.stem,
                content=text,
            )

        except Exception as e:
            return DocumentContent(
                filepath=str(filepath),
                filename=filepath.name,
                file_type="html",
                error=str(e),
            )

    def detect_document_type(self, filepath: Path) -> str:
        """
        Detect document type from file extension and content.

        Args:
            filepath: Path to the document

        Returns:
            Document type string
        """
        filepath = Path(filepath)
        return self.SUPPORTED_TYPES.get(filepath.suffix.lower(), "unknown")

    def is_supported(self, filepath: Path) -> bool:
        """Check if file type is supported."""
        filepath = Path(filepath)
        return filepath.suffix.lower() in self.SUPPORTED_TYPES

    @classmethod
    def get_supported_extensions(cls) -> List[str]:
        """Get list of supported file extensions."""
        return list(cls.SUPPORTED_TYPES.keys())


def auto_detect_additional_documents(
    paths: List[str],
    additional_pdf: bool = False,
) -> List[DocumentContent]:
    """
    Auto-detect and read additional documents.

    Args:
        paths: List of file paths
        additional_pdf: Flag to enable PDF processing (for backwards compatibility)

    Returns:
        List of DocumentContent objects
    """
    reader = DocumentReader()
    documents = []

    for path_str in paths:
        path = Path(path_str)

        if not path.exists():
            continue

        if path.is_file():
            if reader.is_supported(path):
                doc = reader.read_document(path)
                if not doc.error:
                    documents.append(doc)

        elif path.is_dir():
            # Recursively find supported documents
            for ext in DocumentReader.get_supported_extensions():
                for file_path in path.rglob(f"*{ext}"):
                    doc = reader.read_document(file_path)
                    if not doc.error:
                        documents.append(doc)

    return documents
